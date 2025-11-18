import os
import re
import html
from io import BytesIO
from urllib.parse import urlparse
from PIL import Image
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from lxml import etree
import shutil


class MediaExtractorService:
    """
    Service for:
        - Extracting images from Word documents
        - Converting tables → LaTeX
        - Converting equations → LaTeX
        - Replacing placeholders inside document
        - Saving updated Word document
        - Ensuring extracted_media folder is created/reset properly
    """

    # --------------------------------------------------------
    # Constructor
    # --------------------------------------------------------
    def __init__(self):
        pass

    # --------------------------------------------------------
    # Ensure destination folder exists and is empty
    # --------------------------------------------------------
    def prepare_destination_folder(self, dest_folder: str):
        """
        If extracted_media folder:
            - does NOT exist → create it
            - exists but NOT empty → clear all files
        """

        if not os.path.exists(dest_folder):
            os.makedirs(dest_folder, exist_ok=True)
            return

        # Folder exists → clear all files
        for filename in os.listdir(dest_folder):
            file_path = os.path.join(dest_folder, filename)
            try:
                os.remove(file_path)
            except:
                shutil.rmtree(file_path, ignore_errors=True)

    # --------------------------------------------------------
    # XML escaping helper
    # --------------------------------------------------------
    def xml_escape(self, text: str) -> str:
        return (
            text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace('"', "&quot;")
            .replace("'", "&apos;")
        )

    # --------------------------------------------------------
    # Save Extracted Image
    # --------------------------------------------------------
    def save_image(self, image_blob, file_basename, counter, destination_folder):
        os.makedirs(destination_folder, exist_ok=True)
        image_name = f"{file_basename}_img_{counter}.png"
        image_path = os.path.join(destination_folder, image_name)
        image = Image.open(BytesIO(image_blob))
        image.save(image_path)
        return image_path

    # --------------------------------------------------------
    # Convert Table → LaTeX
    # --------------------------------------------------------
    def table_to_latex(self, table):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            rows.append(" & ".join(cells) + " \\\\")
        num_cols = max(len(row.cells) for row in table.rows)
        column_format = "|".join(["c"] * num_cols)
        latex = (
            "\\begin{tabular}{" + f"|{column_format}|" + "}\n"
            "\\hline\n" + "\n\\hline\n".join(rows) +
            "\n\\hline\n\\end{tabular}"
        )
        return latex

    # --------------------------------------------------------
    # Extract plain text from OMML
    # --------------------------------------------------------
    def extract_text_from_omml(self, xml_fragment: str) -> str:
        xml_fragment = re.sub(r"<.*?>", "", xml_fragment)
        xml_fragment = html.unescape(xml_fragment)
        return xml_fragment.strip()

    # --------------------------------------------------------
    # Convert OMML Equation → LaTeX
    # --------------------------------------------------------
    def omml_to_latex(self, omml_xml: str) -> str:
        omml_xml = re.sub(r"\s+", " ", omml_xml)

        # Fractions
        omml_xml = re.sub(
            r"<m:f>.*?<m:num>(.*?)</m:num>.*?<m:den>(.*?)</m:den>.*?</m:f>",
            lambda m: f"\\frac{{{self.extract_text_from_omml(m.group(1))}}}{{{self.extract_text_from_omml(m.group(2))}}}",
            omml_xml,
        )

        # Superscripts
        omml_xml = re.sub(
            r"<m:sSup>.*?<m:e>(.*?)</m:e>.*?<m:sup>(.*?)</m:sup>.*?</m:sSup>",
            lambda m: f"{self.extract_text_from_omml(m.group(1))}^{{{self.extract_text_from_omml(m.group(2))}}}",
            omml_xml,
        )

        # Subscripts
        omml_xml = re.sub(
            r"<m:sSub>.*?<m:e>(.*?)</m:e>.*?<m:sub>(.*?)</m:sub>.*?</m:sSub>",
            lambda m: f"{self.extract_text_from_omml(m.group(1))}_{{{self.extract_text_from_omml(m.group(2))}}}",
            omml_xml,
        )

        # Root
        omml_xml = re.sub(
            r"<m:rad>.*?<m:e>(.*?)</m:e>.*?</m:rad>",
            lambda m: f"\\sqrt{{{self.extract_text_from_omml(m.group(1))}}}",
            omml_xml,
        )

        clean_text = self.extract_text_from_omml(omml_xml)
        return f"\\({clean_text}\\)"

    # --------------------------------------------------------
    # MAIN PROCESS FUNCTION
    # --------------------------------------------------------
    def process_document(self, file_url: str, dest_folder: str) -> str:
        """
        Input:
            file_url → path/URL of Word doc
            dest_folder → extracted_media folder

        Returns:
            updated_doc_path → saved .docx (with media removed and replaced)
        """

        # Ensure extracted_media folder is ready
        self.prepare_destination_folder(dest_folder)

        # Load document
        doc = Document(file_url)
        basename = os.path.splitext(os.path.basename(file_url))[0]
        file_directory = os.path.dirname(file_url)

        img_counter = 0
        tbl_counter = 0
        eqn_counter = 0

        body_elements = list(doc.element.body)

        for element in body_elements:
            tag = element.tag

            # ----------------------------------------------------
            # TABLES
            # ----------------------------------------------------
            if tag.endswith("tbl"):
                tbl_counter += 1
                for tbl in doc.tables:
                    if tbl._element is element:
                        latex_code = self.table_to_latex(tbl)
                        safe_latex = self.xml_escape(latex_code)
                        parent = element.getparent()
                        idx = parent.index(element)
                        placeholder = parse_xml(
                            f'<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                            f'<w:r><w:t>{safe_latex}</w:t></w:r></w:p>'
                        )
                        parent.insert(idx + 1, placeholder)
                        parent.remove(element)
                        break

            # ----------------------------------------------------
            # PARAGRAPHS (images + equations)
            # ----------------------------------------------------
            elif tag.endswith("p"):
                paragraph = next((p for p in doc.paragraphs if p._element is element), None)
                if not paragraph:
                    continue

                # Extract images
                for run in paragraph.runs:
                    blips = run.element.xpath(".//a:blip")
                    if blips:
                        img_counter += 1
                        embed_rid = blips[0].get(qn("r:embed"))
                        image_part = doc.part.related_parts[embed_rid]
                        image_data = image_part.blob

                        img_path = self.save_image(image_data, basename, img_counter, dest_folder)
                        run.text = f"[Image: {img_path}]"

                        for blip in blips:
                            blip.getparent().remove(blip)

                # Extract OMML equations
                math_elems = paragraph._element.xpath(".//m:oMath | .//m:oMathPara")
                for math_elem in math_elems:
                    eqn_counter += 1
                    eqn_xml = etree.tostring(math_elem, encoding="unicode")
                    latex_code = self.omml_to_latex(eqn_xml)
                    safe_latex = self.xml_escape(latex_code)

                    math_elem.getparent().replace(
                        math_elem,
                        parse_xml(
                            f'<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                            f'<w:t>{safe_latex}</w:t></w:r>'
                        ),
                    )

        # --------------------------------------------------------
        # SAVE UPDATED DOCX
        # --------------------------------------------------------
        updated_path = os.path.join(file_directory, f"{basename}_updated.docx")
        doc.save(updated_path)

        return updated_path
