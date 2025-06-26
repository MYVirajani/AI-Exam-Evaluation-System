# # # # src/scripts/embed_lecture_materials.py

# # # import os
# # # import glob
# # # from langchain_community.document_loaders import PyPDFLoader
# # # from langchain.text_splitter import RecursiveCharacterTextSplitter
# # # from langchain_community.vectorstores.pgvector import PGVector
# # # from langchain_openai import OpenAIEmbeddings
# # # from config.settings import config

# # # def embed_lecture_materials():
# # #     pdf_dir = "src/data/lecture_materials"
# # #     embedding_model = OpenAIEmbeddings(
# # #         api_key=config.openai.api_key,
# # #         model=config.openai.embedding_model
# # #     )

# # #     vector_db = PGVector(
# # #         collection_name=config.database.lecture_collection,
# # #         connection_string=config.database.connection_string,
# # #         embedding_function=embedding_model,
# # #     )

# # #     pdf_files = glob.glob(os.path.join(pdf_dir, "*.pdf"))
# # #     if not pdf_files:
# # #         print("No PDFs found in", pdf_dir)
# # #         return

# # #     all_docs = []
# # #     for file in pdf_files:
# # #         loader = PyPDFLoader(file)
# # #         docs = loader.load()
# # #         splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
# # #         split_docs = splitter.split_documents(docs)
# # #         all_docs.extend(split_docs)

# # #     vector_db.add_documents(all_docs)
# # #     print(f"Embedded {len(all_docs)} chunks from {len(pdf_files)} PDFs into vector DB.")

# # # if __name__ == "__main__":
# # #     embed_lecture_materials()


# # import argparse
# # import os
# # import sys
# # sys.path.append(os.path.abspath(os.path.join(os.path.dirname(_file_), "../../")))

# # from src.controller.lecture_material_embed_controller import embed_lecture_materials

# # if __name__ == "__main__":
# #     parser = argparse.ArgumentParser()
# #     parser.add_argument("--dir", required=True, help="Directory of lecture materials")
# #     args = parser.parse_args()

# #     try:
# #         embed_lecture_materials(args.dir)
# #         print("Lecture materials embedded successfully.")
# #     except Exception as e:
# #         print(f"Error: {e}")

# """
# Run example:
#   python -m src.scripts.embed_lecture_materials \
#          --provider OpenAI \
#          --model gpt-4o \
#          --embedder text-embedding-3-small \
#          --root data/Lecture_Material
# """

# import argparse, pathlib, logging
# from docx import Document
# import pdfplumber

# from src.utils.token_chunker import chunk_text
# from src.models.lecture_chunk import LectureChunk
# from src.services.model_answer_extractor import ModelAnswerExtractor     # only to pick provider structs
# from src.services.embedding.openai_embedder  import OpenAIEmbedder
# from src.services.embedding.gemini_embedder  import GeminiEmbedder
# from src.services.database_services.lecture_material_embedding_db import (
#     LectureMaterialEmbeddingDB,
# )

# logging.basicConfig(level=logging.INFO)
# log = logging.getLogger(__name__)

# # ---------------- helpers --------------------------------------------------
# def read_text(path: pathlib.Path) -> str:
#     if path.suffix.lower() == ".docx":
#         doc = Document(path)
#         return "\n".join(p.text for p in doc.paragraphs)
#     if path.suffix.lower() == ".pdf":
#         with pdfplumber.open(path) as pdf:
#             return "\n".join(p.extract_text() or "" for p in pdf.pages)
#     raise ValueError(f"Unsupported file type: {path.name}")

# def module_from_path(path: pathlib.Path) -> str:
#     """
#     Folder structure: data/Lecture_Material/<MODULE_CODE>/filename.pdf
#     If you don't use sub-folders per module, fall back to filename stem.
#     """
#     parts = path.relative_to(root).parts
#     return parts[0] if len(parts) > 1 else path.stem.split("_")[0]

# # ---------------- main -----------------------------------------------------
# def main():
#     ap = argparse.ArgumentParser()
#     ap.add_argument("--provider", required=True, choices=["OpenAI", "GoogleGemini"])
#     ap.add_argument("--model",    required=True)
#     ap.add_argument("--embedder", default="text-embedding-3-small")
#     ap.add_argument("--root",     default="data/Lecture_Material")
#     ap.add_argument("--max_tokens", type=int, default=1000)
#     ap.add_argument("--overlap",    type=int, default=200)
#     args = ap.parse_args()

#     global root
#     root = pathlib.Path(args.root)

#     # choose embedder
#     embedder = (
#         OpenAIEmbedder(args.embedder)
#         if args.provider == "OpenAI"
#         else GeminiEmbedder(model_name=args.embedder)
#     )
#     vec_db = LectureMaterialEmbeddingDB(embedder)

#     for file in root.rglob("*"):
#         if file.suffix.lower() not in {".pdf", ".docx"}:
#             continue
#         log.info("→ %s", file.relative_to(root))

#         module_code = module_from_path(file)
#         raw = read_text(file)
#         chunks = [
#             LectureChunk(module_code, file.name, idx, txt)
#             for idx, txt in enumerate(
#                 chunk_text(raw, max_tokens=args.max_tokens, overlap=args.overlap)
#             )
#             if txt.strip()
#         ]
#         vec_db.save_chunks(chunks)

#     vec_db.close()
#     print(" Lecture materials embedded.")

# if __name__ == "__main__":
#     main()

"""
Embed lecture notes for a single module.

Example (PowerShell one-liner):
  python -m src.scripts.embed_lecture_materials `
    --provider GoogleGemini `
    --model gemini-1.5-pro `
    --embedder embedding-001 `
    --root data/Lecture_Material `
    --module EE6250
"""

import argparse, pathlib, logging, sys
from docx import Document
import pdfplumber
from dotenv import load_dotenv

from src.utils.token_chunker import chunk_text
from src.models.lecture_chunk import LectureChunk
from src.services.embedding.openai_embedder  import OpenAIEmbedder
from src.services.embedding.gemini_embedder  import GeminiEmbedder
from src.services.database_services.lecture_material_embedding_db import (
    LectureMaterialEmbeddingDB,
)

load_dotenv()
logging.basicConfig(level=logging.INFO)
log = logging.getLogger(__name__)

# ──────────────────────────────────────────────────────────────────────
def read_text(path: pathlib.Path) -> str:
    if path.suffix.lower() == ".docx":
        doc = Document(path)
        return "\n".join(p.text for p in doc.paragraphs)
    if path.suffix.lower() == ".pdf":
        with pdfplumber.open(path) as pdf:
            return "\n".join(p.extract_text() or "" for p in pdf.pages)
    raise ValueError(f"Unsupported file type: {path.name}")

# ──────────────────────────────────────────────────────────────────────
def main() -> None:
    ap = argparse.ArgumentParser(description="Embed lecture materials for one module")
    ap.add_argument("--provider", required=True, choices=["OpenAI", "GoogleGemini"])
    ap.add_argument("--model",    required=True, help="LLM model name (unused here)")
    ap.add_argument("--embedder", default="text-embedding-3-small")
    ap.add_argument("--root",     default="data/Lecture_Material")
    ap.add_argument("--module",   help="Module code to ingest, e.g. EE6250")
    ap.add_argument("--max_tokens", type=int, default=1000)
    ap.add_argument("--overlap",    type=int, default=200)
    args = ap.parse_args()

    # ask interactively if --module omitted
    if not args.module:
        print("👋  Available modules under", args.root)
        for p in pathlib.Path(args.root).iterdir():
            if p.is_dir():
                print(" •", p.name)
        args.module = input("\nEnter module code to embed: ").strip()

    module_dir = pathlib.Path(args.root) / args.module
    if not module_dir.exists():
        log.error("Folder %s does not exist.", module_dir)
        sys.exit(1)

    # choose embedder
    embedder = (
        OpenAIEmbedder(args.embedder)
        if args.provider == "OpenAI"
        else GeminiEmbedder(model_name=args.embedder)
    )
    vec_db = LectureMaterialEmbeddingDB(embedder)

    # scan & embed
    files = list(module_dir.rglob("*"))
    if not any(f.suffix.lower() in {".pdf", ".docx"} for f in files):
        log.warning("No .pdf or .docx files found in %s", module_dir)
        sys.exit(0)

    for f in files:
        if f.suffix.lower() not in {".pdf", ".docx"}:
            continue

        log.info("📄 %s", f.relative_to(module_dir))
        raw = read_text(f)

        chunks = [
            LectureChunk(args.module, f.name, idx, txt)
            for idx, txt in enumerate(
                chunk_text(raw, max_tokens=args.max_tokens, overlap=args.overlap)
            )
            if txt.strip()
        ]
        log.info("   → %d chunks", len(chunks))
        vec_db.save_chunks(chunks)

    vec_db.close()
    print(f"✅ Finished embedding lecture material for {args.module}.")

if __name__ == "__main__":
    main()
