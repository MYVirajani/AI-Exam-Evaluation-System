import os
import json
import logging
from psycopg2 import sql
from docx import Document

from .base_vector_db_service import BaseVectorDBService
from src.services.database_services.evaluation_model_db import EvaluationModelService

from src.services.embedding.openai_embedder import OpenAIEmbedder
from src.services.embedding.gemini_embedder import GeminiEmbedder

logger = logging.getLogger(__name__)
logger.setLevel(logging.INFO)


def build_embedder(provider: str, embedding_model: str):
    provider = provider.lower()
    if provider == "openai":
        return OpenAIEmbedder(embedding_model=embedding_model)
    elif provider == "gemini":
        return GeminiEmbedder(embedding_model=embedding_model)
    raise ValueError(
        f"❌ Unsupported provider '{provider}' for embedding model: {embedding_model}"
    )


class LectureMaterialVectorDBService(BaseVectorDBService):
    def __init__(self, model_id: str):
        # Load model config from Evaluation_Model table
        self.model_db = EvaluationModelService()
        model_config = self.model_db.get_model_config(model_id)

        if not model_config:
            raise ValueError(f"Model ID not found in Evaluation_Model: {model_id}")

        self.model_id = model_id
        self.provider = model_config["provider"]
        self.chat_model = model_config.get("chat_model")
        self.embedding_model = model_config.get("embedding_model")
        self.temperature = model_config.get("temperature")
        self.model_name = model_config.get("model_name", "").lower()

        logger.info(f"Loaded model config: {model_config}")

        embedder = build_embedder(provider=self.provider, embedding_model=self.embedding_model)

        super().__init__(embedder)

        self.suffix = self.provider.lower()
        self.table_name = f"lecture_material_embeddings_{self.suffix}"

        self.ensure_table_exists()

    def ensure_table_exists(self):
        table_identifier = sql.Identifier(self.table_name)
        vector_dim = sql.SQL(str(self.embedder.get_embedding_dimension()))

        create_table_query = sql.SQL("""
            CREATE EXTENSION IF NOT EXISTS vector;
            CREATE EXTENSION IF NOT EXISTS "uuid-ossp";

            CREATE TABLE IF NOT EXISTS {table_name} (
                id UUID PRIMARY KEY DEFAULT uuid_generate_v4(),
                lecturer_id TEXT NOT NULL,
                module_id TEXT NOT NULL,
                lecture_material_id TEXT NOT NULL,
                file_path TEXT NOT NULL,
                content TEXT,
                embedding VECTOR({dim}),
                model_name TEXT NOT NULL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                UNIQUE (lecturer_id, module_id, lecture_material_id)
            );
        """).format(
            table_name=table_identifier,
            dim=vector_dim
        )

        self.cursor.execute(create_table_query)
        self.conn.commit()
        logger.info(f"Table ready: {self.table_name}")

    def already_exists(self, lecturer_id: str, module_id: str, lecture_material_id: str) -> bool:
        check_query = sql.SQL("""
            SELECT 1 FROM {table_name}
            WHERE lecturer_id = %s AND module_id = %s AND lecture_material_id = %s
            LIMIT 1;
        """).format(table_name=sql.Identifier(self.table_name))

        self.cursor.execute(check_query, (lecturer_id, module_id, lecture_material_id))
        return self.cursor.fetchone() is not None

    def insert_embedding(self, lecturer_id, module_id, lecture_material_id, file_path, content, embedding):
        if self.already_exists(lecturer_id, module_id, lecture_material_id):
            logger.info(f"Skipping duplicate: {lecture_material_id}")
            return

        insert_query = sql.SQL("""
            INSERT INTO {table_name} 
            (lecturer_id, module_id, lecture_material_id, file_path, content, embedding, model_name)
            VALUES (%s, %s, %s, %s, %s, %s, %s);
        """).format(table_name=sql.Identifier(self.table_name))

        # Some embedders return lists/tuples — pass as-is; psycopg2 will adapt
        self.cursor.execute(
            insert_query,
            (lecturer_id, module_id, lecture_material_id, file_path, content, embedding, self.model_name)
        )

        logger.info(f"Inserted embedding for {lecture_material_id}")

    def bulk_insert_embeddings(self, lecturer_id, module_id, file_path, lecture_material_id, contents, embeddings):
        if len(contents) != len(embeddings):
            raise ValueError("Mismatch between contents and embeddings!")

        if self.already_exists(lecturer_id, module_id, lecture_material_id):
            logger.info(f"Skipping bulk insert for {lecture_material_id} (already embedded)")
            return

        for content, emb in zip(contents, embeddings):
            self.insert_embedding(lecturer_id, module_id, lecture_material_id, file_path, content, emb)

        self.commit()
        logger.info(f"Bulk embeddings inserted for {lecture_material_id}")

    # ----------------------------------------------------------------------
    # NEW FUNCTION:
    # Generate embeddings for the concatenated content of a DOCX and store it
    # ----------------------------------------------------------------------
    def generate_and_store_embeddings(self, lecturer_id: str, module_id: str, lecture_material_id: str, file_path: str):
        """
        Reads the DOCX (file_path), concatenates all textual content (paragraphs, table cells),
        generates embeddings using the selected embedder, and inserts them into the vector table.
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found for embeddings: {file_path}")

        # Read docx text (paragraphs + table cells)
        try:
            doc = Document(file_path)
        except Exception as e:
            logger.error(f"Failed to open DOCX for embedding: {e}")
            raise

        parts = []
        for p in doc.paragraphs:
            text = p.text.strip()
            if text:
                parts.append(text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        text = p.text.strip()
                        if text:
                            parts.append(text)

        # Concatenate into a single document content. If this is too large for your embedder,
        # consider chunking. For now we create one concatenated string.
        concatenated_content = "\n\n".join(parts).strip()
        if not concatenated_content:
            logger.warning(f"No textual content found in DOCX: {file_path}")
            concatenated_content = ""

        # Generate embeddings using embedder. Try common method names for compatibility.
        try:
            # preferred: embed_texts / embed_documents / embed
            if hasattr(self.embedder, "embed_texts"):
                embeddings = self.embedder.embed_texts([concatenated_content])
            elif hasattr(self.embedder, "embed_documents"):
                embeddings = self.embedder.embed_documents([concatenated_content])
            elif hasattr(self.embedder, "embed"):
                embeddings = self.embedder.embed([concatenated_content])
            elif hasattr(self.embedder, "get_embeddings"):
                embeddings = self.embedder.get_embeddings([concatenated_content])
            else:
                raise AttributeError("Embedder does not expose a known embed method")
        except Exception as e:
            logger.error(f"Failed to generate embeddings: {e}")
            raise

        # Expect embeddings as list-of-vectors; pull first
        if not embeddings or len(embeddings) == 0:
            logger.error("Embedder returned no embeddings")
            raise ValueError("Embedder returned no embeddings")

        embedding_vector = embeddings[0]

        # Insert embedding into the DB (single insert)
        try:
            self.insert_embedding(
                lecturer_id=lecturer_id,
                module_id=module_id,
                lecture_material_id=lecture_material_id,
                file_path=file_path,
                content=concatenated_content,
                embedding=embedding_vector
            )
            # commit
            self.commit()
            logger.info(f"Embedding saved for lecture_material_id={lecture_material_id}")
        except Exception as e:
            self.conn.rollback()
            logger.error(f"Failed to insert embedding into DB: {e}")
            raise
