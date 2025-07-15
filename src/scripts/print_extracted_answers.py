# # from src.services.answer_extractor import AnswerExtractor
# # from docx import Document
# # import sys
# # import os
# # sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '../../')))


# # def load_docx_text(docx_path):
# #     doc = Document(docx_path)
# #     return "\n".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])

# # if __name__ == "__main__":
# #     # 🔁 Provide your actual .docx path here
# #     docx_path = "data/student_answer1.docx"

# #     raw_text = load_docx_text(docx_path)
# #     extractor = AnswerExtractor()

# #     answers = extractor.extract_answers_with_llm(raw_text)

# #     print("\nExtracted Answers:\n" + "-" * 50)
# #     for ans in answers:
# #         question_label = f"{ans.question_id})" if not ans.sub_question_id else f"{ans.question_id}) {ans.sub_question_id})"
# #         print(question_label)
# #         print(f"Answer: {ans.answer_text}\n")


# import sys
# import os
# from docx import Document

# # ✅ Fix import path so Python can find the src folder
# sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "../../")))

# from src.services.answer_extractor import AnswerExtractor

# def load_docx_text(docx_path: str) -> str:
#     """Extract full text from a DOCX file."""
#     doc = Document(docx_path)
#     return "\n".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])

# if __name__ == "__main__":
#     docx_path = "data/Answer_Scripts/student_answer1.docx"  # ✅ Ensure this file exists

#     raw_text = load_docx_text(docx_path)
#     extractor = AnswerExtractor()
#     answers = extractor.extract_answers_with_llm(raw_text)

#     print("\nExtracted Answers:\n" + "-" * 50)
#     for ans in answers:
#         q = ans.question_id
#         sub = ans.sub_question_id
#         label = f"{q})" if not sub else f"{q}) {sub})"
#         print(label)
#         print(f"Answer: {ans.answer_text}\n")


import sys
import os
from docx import Document

# ✅ Fix import path so Python can find the src folder
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "../../")))

from src.services.answer_extractor import AnswerExtractor

def load_docx_text(docx_path: str) -> str:
    """Extract full text from a DOCX file."""
    doc = Document(docx_path)
    return "\n".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])

if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="Extract student answers from DOCX using LLMs")
    parser.add_argument("--provider", required=True, choices=["OpenAI", "GoogleGemini"], help="LLM provider")
    parser.add_argument("--model", required=True, help="Model name (e.g., gpt-4o, gemini-pro)")
    parser.add_argument("--file", default="data/Answer_Scripts/EE6250_EG-2020-4247.docx", help="Path to DOCX answer script")

    args = parser.parse_args()

    raw_text = load_docx_text(args.file)

    extractor = AnswerExtractor(
        selected_provider=args.provider,
        selected_model=args.model
    )

    answers = extractor.extract_answers_with_llm(raw_text)

    print(f"\nExtracted Answers using {args.provider} ({args.model}):\n" + "-" * 60)
    for ans in answers:
        label = f"{ans.question_id})" if not ans.sub_question_id else f"{ans.question_id}) {ans.sub_question_id})"
        print(label)
        print(f"Answer: {ans.answer_text}\n")
