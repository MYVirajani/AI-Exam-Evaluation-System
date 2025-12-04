#!/usr/bin/env python3
import os
import logging
from dotenv import load_dotenv
from docx import Document
from typing import List

from src.services.database_services.lecture_material_db_service import LectureMaterialDBService
from src.services.database_services.lecture_material_vector_db_service import LectureMaterialVectorDBService
from src.services.extractors.media_extractor_service import MediaExtractorService
from src.services.summary.image_summarise_service import ImageSummarizer
from src.services.extractors.content_extractor_service import ContentExtractorService  

load_dotenv()
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


# ---------------------------------------------------------
# Replace image tags in DOCX
# ---------------------------------------------------------
def replace_image_tags_in_docx(docx_path: str, media_items: list) -> list:
    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"DOCX not found: {docx_path}")

    doc = Document(docx_path)
    unmatched = []

    for media in media_items:
        media['matched'] = False

    def process_paragraphs(paragraphs):
        for p in paragraphs:
            text = p.text
            if not text:
                continue
            for media in media_items:
                pattern = f"[Image: {media['media_url']}]"
                if pattern in text:
                    new_text = text.replace(pattern, media.get('media_summary') or "")
                    p.text = new_text
                    media['matched'] = True
                    text = new_text

    process_paragraphs(doc.paragraphs)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                process_paragraphs(cell.paragraphs)

    for media in media_items:
        if not media.get('matched', False):
            unmatched.append(media)

    if unmatched:
        doc.add_page_break()
        doc.add_paragraph("Image Summaries:")
        for media in unmatched:
            summary = media.get('media_summary') or ""
            p = doc.add_paragraph()
            p.add_run(f"{media['media_url']}: ").bold = True
            p.add_run(summary)

    doc.save(docx_path)
    return unmatched


# ---------------------------------------------------------
# Main pipeline: handle multiple lessons & models
# ---------------------------------------------------------
def process_lecture_materials(lesson_ids: List[str], model_ids: List[str]):
    db = LectureMaterialDBService()
    extractor = MediaExtractorService()
    content_extractor = ContentExtractorService()  # For PDF, PPTX, DOCX, TXT

    for model_id in model_ids:
        summarizer = ImageSummarizer(model_id)

        # Init vector DB per model
        try:
            vector_service = LectureMaterialVectorDBService(model_id=model_id)
            logger.info(f"✅ Vector service initialized for model_id: {model_id}")
        except Exception as e:
            logger.error(f"❌ Failed to initialize vector service for model {model_id}: {e}")
            vector_service = None

        for lesson_id in lesson_ids:
            logger.info(f"🔍 Processing lesson_id={lesson_id} with model_id={model_id}")

            lecture_materials = db.fetch_lecture_materials_by_lesson(lesson_id)
            if not lecture_materials:
                logger.warning(f"No lecture materials found for lesson {lesson_id}")
                continue

            # -------------------------
            # PROCESS EACH MATERIAL
            # -------------------------
            for material in lecture_materials:
                (
                    lecture_material_id,
                    _,
                    file_name,
                    file_url,
                    extracted_file_url,
                    created_on,
                    updated_on,
                    description
                ) = material

                logger.info(f"\n📄 Processing File: {file_url} (lecture_material_id={lecture_material_id})")

                if not os.path.exists(file_url):
                    logger.error(f"❌ File not found: {file_url}")
                    continue

                parent_folder = os.path.dirname(file_url)
                extracted_dest = os.path.join(parent_folder, "extracted_media")

                # Extract media (images)
                try:
                    updated_doc_path, extracted_media_paths = extractor.process_document(file_url, extracted_dest)
                    logger.info(f"✅ Updated DOCX Created: {updated_doc_path}")
                except Exception as e:
                    logger.error(f"❌ Failed to extract media from {file_url}: {e}")
                    continue

                # Save extracted media + summaries
                if extracted_media_paths:
                    logger.info(f"🖼 Saving {len(extracted_media_paths)} extracted media items...")
                    for media_path in extracted_media_paths:
                        try:
                            summary_text = summarizer.summarize_image(
                                image_path=media_path,
                                mode="model",
                                domain="Engineering"
                            ) or None

                            db.insert_media(
                                model_id=model_id,
                                lecture_material_id=lecture_material_id,
                                media_url=media_path,
                                media_summary=summary_text
                            )
                            logger.info(f"   → Inserted media: {media_path}")

                        except Exception as e:
                            logger.error(f"❌ Failed inserting media {media_path}: {e}")

                # Update extracted_file_url
                try:
                    db.update_extracted_file_path(lecture_material_id, updated_doc_path)
                except Exception as e:
                    logger.error(f"❌ Failed to update extracted_file_url: {e}")

                # Replace [Image: ...] tags
                try:
                    if updated_doc_path and os.path.exists(updated_doc_path):
                        full_data = db.get_full_material_data_by_lesson_ids([lesson_id])
                        current_entry = next(
                            (entry for entry in full_data if str(entry["lecture_material_id"]) == str(lecture_material_id)),
                            None
                        )
                        if current_entry:
                            media_items = current_entry.get("media", [])
                            if media_items:
                                unmatched = replace_image_tags_in_docx(updated_doc_path, media_items)
                                if unmatched:
                                    logger.info(f"→ {len(unmatched)} media summaries appended")
                        else:
                            logger.warning(f"No full data entry found for {lecture_material_id}")
                except Exception as e:
                    logger.error(f"❌ Failed to replace image tags: {e}")

            # -------------------------
            # Generate embeddings
            # -------------------------
            logger.info(f"\n📌 Generating embeddings for lesson_id={lesson_id}")

            updated_materials = db.fetch_lecture_materials_by_lesson(lesson_id)
            for material in updated_materials:
                (
                    lecture_material_id,
                    _,
                    file_name,
                    file_url,
                    extracted_file_url,
                    created_on,
                    updated_on,
                    description
                ) = material

                final_path = extracted_file_url if extracted_file_url else file_url
                if not os.path.exists(final_path):
                    logger.error(f"❌ Final file not found for embedding: {final_path}")
                    continue

                try:
                    full_text = content_extractor.extract_text(final_path)
                    if not full_text.strip():
                        logger.warning(f"⚠️ No text extracted for {lecture_material_id}, skipping embedding")
                        continue
                except Exception as e:
                    logger.error(f"❌ Failed extracting text from {final_path}: {e}")
                    continue

                # Save embeddings
                if vector_service:
                    try:
                        vector_service.save_lecture_material(
                            lecture_material_id=str(lecture_material_id),
                            full_content=full_text
                        )
                        logger.info(f"✅ Embeddings saved for lecture_material_id={lecture_material_id}")
                    except Exception as e:
                        logger.error(f"❌ Failed saving embeddings: {e}")

        # Close vector service for this model
        if vector_service:
            try:
                vector_service.close()
            except:
                pass

    # Close DB
    try:
        db.close()
    except:
        pass

    logger.info("\n🎉 All lessons & models processed successfully!\n")


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="Process lecture materials for multiple lessons and models")
    parser.add_argument(
        "--lesson_ids",
        required=True,
        nargs='+',
        help="List of lesson IDs to process"
    )
    parser.add_argument(
        "--model_ids",
        required=False,
        nargs='+',
        default=["eaa81306-f9e3-4c96-901d-3b7a80a3f4ac"],
        help="List of model IDs to use for summarizer & embeddings"
    )

    args = parser.parse_args()

    process_lecture_materials(args.lesson_ids, args.model_ids)
