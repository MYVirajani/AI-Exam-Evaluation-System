import os
import html
import re
from io import BytesIO
from PIL import Image
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from lxml import etree
from docx2pdf import convert  # pip install docx2pdf

# -----------------------------
# CONFIGURATION
# -----------------------------
ROOT_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
DATA_DIR = os.path.join(ROOT_DIR, "data")
INPUT_DIR = os.path.join(DATA_DIR, "raw_answer_scripts_docs")
EXTRACT_DIR = os.path.join(DATA_DIR, "extracted_media")
OUTPUT_DIR = os.path.join(DATA_DIR, "updated_answer_scripts_docs")
PDF_DIR = os.path.join(DATA_DIR, "Answer_scripts")

os.makedirs(EXTRACT_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)
os.makedirs(PDF_DIR, exist_ok=True)

# -----------------------------
# HELPER FUNCTIONS
# -----------------------------
def xml_escape(text: str) -> str:
    """Escape XML special characters (&, <, >, ", ')"""
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
        .replace("'", "&apos;")
    )


def save_image(image_blob, doc_basename, counter):
    """Save image bytes to file and return relative path."""
    image_name = f"{doc_basename}_img_{counter}.png"
    image_path = os.path.join(EXTRACT_DIR, image_name)
    image = Image.open(BytesIO(image_blob))
    image.save(image_path)
    return image_path


def table_to_latex(table):
    """Convert a Word table to LaTeX tabular code."""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append(" & ".join(cells) + " \\\\")
    num_cols = max(len(row.cells) for row in table.rows)
    column_format = "|".join(["c"] * num_cols)
    latex = (
        "\\begin{tabular}{" + f"|{column_format}|" + "}\n"
        "\\hline\n" + "\n\\hline\n".join(rows) +
        "\n\\hline\n\\end{tabular}"
    )
    return latex


def extract_text_from_omml(xml_fragment: str) -> str:
    """Extract readable text from OMML XML (remove tags, keep symbols)."""
    xml_fragment = re.sub(r"<.*?>", "", xml_fragment)
    xml_fragment = html.unescape(xml_fragment)
    return xml_fragment.strip()


def omml_to_latex(omml_xml: str) -> str:
    """
    Convert Word OMML (Office Math Markup Language) equations to simplified LaTeX.
    This works for most typical engineering/science expressions.
    """
    omml_xml = re.sub(r"\s+", " ", omml_xml)

    # Fractions
    omml_xml = re.sub(
        r"<m:f>.*?<m:num>(.*?)</m:num>.*?<m:den>(.*?)</m:den>.*?</m:f>",
        lambda m: f"\\frac{{{extract_text_from_omml(m.group(1))}}}{{{extract_text_from_omml(m.group(2))}}}",
        omml_xml,
    )

    # Superscripts
    omml_xml = re.sub(
        r"<m:sSup>.*?<m:e>(.*?)</m:e>.*?<m:sup>(.*?)</m:sup>.*?</m:sSup>",
        lambda m: f"{extract_text_from_omml(m.group(1))}^{{{extract_text_from_omml(m.group(2))}}}",
        omml_xml,
    )

    # Subscripts
    omml_xml = re.sub(
        r"<m:sSub>.*?<m:e>(.*?)</m:e>.*?<m:sub>(.*?)</m:sub>.*?</m:sSub>",
        lambda m: f"{extract_text_from_omml(m.group(1))}_{{{extract_text_from_omml(m.group(2))}}}",
        omml_xml,
    )

    # Roots
    omml_xml = re.sub(
        r"<m:rad>.*?<m:deg>(.*?)</m:deg>.*?<m:e>(.*?)</m:e>.*?</m:rad>",
        lambda m: f"\\sqrt[{extract_text_from_omml(m.group(1))}]{{{extract_text_from_omml(m.group(2))}}}",
        omml_xml,
    )
    omml_xml = re.sub(
        r"<m:rad>.*?<m:e>(.*?)</m:e>.*?</m:rad>",
        lambda m: f"\\sqrt{{{extract_text_from_omml(m.group(1))}}}",
        omml_xml,
    )

    # Strip remaining tags and wrap in inline math
    clean_text = extract_text_from_omml(omml_xml)
    return f"\\({clean_text}\\)"


# -----------------------------
# CORE PROCESSING FUNCTION
# -----------------------------
def process_docx(docx_path):
    """Extract images, tables, and equations; convert to LaTeX placeholders."""
    doc = Document(docx_path)
    basename = os.path.splitext(os.path.basename(docx_path))[0]
    print(f"\n📄 Processing: {basename}.docx")

    img_counter = 0
    tbl_counter = 0
    eqn_counter = 0

    body_elements = list(doc.element.body)

    for element in body_elements:
        tag = element.tag

        # Handle Tables
        if tag.endswith("tbl"):
            tbl_counter += 1
            for tbl in doc.tables:
                if tbl._element is element:
                    latex_code = table_to_latex(tbl)
                    safe_latex = xml_escape(latex_code)
                    parent = element.getparent()
                    idx = parent.index(element)
                    placeholder = parse_xml(
                        f'<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                        f'<w:r><w:t>{safe_latex}</w:t></w:r></w:p>'
                    )
                    parent.insert(idx + 1, placeholder)
                    parent.remove(element)
                    break

        # Handle Paragraphs
        elif tag.endswith("p"):
            paragraph = next((p for p in doc.paragraphs if p._element is element), None)
            if not paragraph:
                continue

            # Convert images
            for run in paragraph.runs:
                blips = run.element.xpath(".//a:blip")
                if blips:
                    img_counter += 1
                    embed_rid = blips[0].get(qn("r:embed"))
                    image_part = doc.part.related_parts[embed_rid]
                    image_data = image_part.blob
                    img_path = save_image(image_data, basename, img_counter)
                    run.text = f"[Image: {img_path}]"
                    for blip in blips:
                        blip.getparent().remove(blip)

            # Convert equations (OMML)
            math_elems = paragraph._element.xpath(".//m:oMath | .//m:oMathPara")
            for math_elem in math_elems:
                eqn_counter += 1
                # FIX: etree.tostring() instead of .xml
                eqn_xml = etree.tostring(math_elem, encoding="unicode")
                latex_code = omml_to_latex(eqn_xml)
                safe_latex = xml_escape(latex_code)
                # Replace the math XML with LaTeX text node
                math_elem.getparent().replace(
                    math_elem,
                    parse_xml(
                        f'<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                        f'<w:t>{safe_latex}</w:t></w:r>'
                    ),
                )

    # Save updated document
    updated_path = os.path.join(OUTPUT_DIR, f"{basename}.docx")
    doc.save(updated_path)
    print(f"✅ Saved updated Word: {updated_path}")
    print(f"   Extracted {img_counter} images, replaced {tbl_counter} tables, converted {eqn_counter} equations.")

    # Convert to PDF
    pdf_path = os.path.join(PDF_DIR, f"{basename}.pdf")
    try:
        convert(updated_path, pdf_path)
        print(f"📘 Converted to PDF: {pdf_path}")
    except Exception as e:
        print(f"⚠️ PDF conversion failed for {basename}: {e}")

    return updated_path, pdf_path


# -----------------------------
# MAIN SCRIPT
# -----------------------------
def main():
    docx_files = [f for f in os.listdir(INPUT_DIR) if f.lower().endswith(".docx")]
    if not docx_files:
        print("⚠️ No .docx files found in input folder.")
        return

    for file in docx_files:
        process_docx(os.path.join(INPUT_DIR, file))

    print("\n🎉 All documents processed successfully!")
    print(f"📂 Extracted images in: {EXTRACT_DIR}")
    print(f"📄 Updated Word docs in: {OUTPUT_DIR}")
    print(f"📘 PDFs in: {PDF_DIR}")


if __name__ == "__main__":
    main()
