import os
from io import BytesIO
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx2pdf import convert  # pip install docx2pdf


def save_image(image_blob, folder_path, doc_basename, counter):
    """Save image bytes to the same folder as the original file."""
    image_name = f"{doc_basename}_img_{counter}.png"
    image_path = os.path.join(folder_path, image_name)
    image = Image.open(BytesIO(image_blob))
    image.save(image_path)
    return image_path


def save_table_as_image(table, folder_path, doc_basename, counter):
    """Save table content as an image file with text drawn on it."""
    rows = []
    for row in table.rows:
        cols = [cell.text.strip() for cell in row.cells]
        rows.append(" | ".join(cols))
    text = "\n".join(rows) or "(Empty Table)"

    img_name = f"{doc_basename}_tbl_{counter}.png"
    img_path = os.path.join(folder_path, img_name)

    font = ImageFont.load_default()
    lines = text.splitlines()
    width = int(max(font.getlength(line) for line in lines) + 20)
    height = int(15 * len(lines) + 20)
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    y = 10
    for line in lines:
        draw.text((10, y), line, fill="black", font=font)
        y += 15
    img.save(img_path)
    return img_path


def process_docx(file_path):
    """
    Process a .docx file in any location:
      - Extracts images and tables.
      - Saves them in the same folder.
      - Replaces them with [Image: path] / [Table: path] references.
      - Saves updated .docx and PDF in same folder.
      - Returns updated file absolute path.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"❌ File not found: {file_path}")

    if not file_path.lower().endswith(".docx"):
        raise ValueError("❌ Only .docx files are supported.")

    folder_path = os.path.dirname(file_path)
    basename = os.path.splitext(os.path.basename(file_path))[0]

    print(f"\n📄 Processing file: {file_path}")

    doc = Document(file_path)
    img_counter = 0
    tbl_counter = 0

    # Iterate through elements in body order
    body_elements = list(doc.element.body)

    for element in body_elements:
        tag = element.tag
        if tag.endswith("tbl"):  # Table
            tbl_counter += 1
            for tbl in doc.tables:
                if tbl._element is element:
                    tbl_path = save_table_as_image(tbl, folder_path, basename, tbl_counter)
                    parent = element.getparent()
                    idx = parent.index(element)
                    placeholder = parse_xml(
                        f'<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                        f'<w:r><w:t>[Table: {os.path.basename(tbl_path)}]</w:t></w:r></w:p>'
                    )
                    parent.insert(idx + 1, placeholder)
                    parent.remove(element)
                    break

        elif tag.endswith("p"):  # Paragraph (check for image)
            paragraph = None
            for p in doc.paragraphs:
                if p._element is element:
                    paragraph = p
                    break
            if paragraph:
                for run in paragraph.runs:
                    blips = run.element.xpath(".//a:blip")
                    if blips:
                        img_counter += 1
                        embed_rid = blips[0].get(qn("r:embed"))
                        image_part = doc.part.related_parts[embed_rid]
                        image_data = image_part.blob
                        img_path = save_image(image_data, folder_path, basename, img_counter)
                        run.text = f"[Image: {os.path.basename(img_path)}]"
                        for blip in blips:
                            blip.getparent().remove(blip)

    # Save updated document in same folder
    updated_path = os.path.join(folder_path, f"updated_{basename}.docx")
    doc.save(updated_path)
    print(f"✅ Saved updated Word file: {updated_path}")

    # Convert to PDF (also saved in same folder)
    pdf_path = os.path.join(folder_path, f"{basename}.pdf")
    try:
        convert(updated_path, pdf_path)
        print(f"📘 Converted to PDF: {pdf_path}")
    except Exception as e:
        print(f"⚠️ PDF conversion failed: {e}")

    print(f"   Extracted {img_counter} images, {tbl_counter} tables.")
    return os.path.abspath(updated_path)


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="Process a single DOCX file and extract its media.")
    parser.add_argument("file_path", help="Path to the DOCX file to process")
    args = parser.parse_args()

    try:
        updated_url = process_docx(args.file_path)
        print(f"\n🎉 Processing complete!\n📍 Updated file URL: {updated_url}")
    except Exception as e:
        print(f"❌ Error: {e}")
