import sys
import os
import time
import pathlib
import json
import re
from docx import Document
from pprint import pprint
import psycopg2
from psycopg2.extras import RealDictCursor
from dotenv import load_dotenv
import pdfplumber
import argparse

sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "../../")))

from src.services.answer_extractor import AnswerExtractor
from src.services.database_services.student_answer_db import StudentAnswerService

load_dotenv()

def clean_json_string(json_str):
    """Clean JSON string by removing/escaping problematic control characters."""
    if not json_str:
        return json_str
    
    # Remove or replace common problematic control characters
    # Keep only allowed whitespace characters (space, tab, newline, carriage return)
    cleaned = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', json_str)
    
    # Additional cleaning - escape unescaped quotes within strings
    # This is a basic approach - you might need more sophisticated handling
    return cleaned

def safe_json_parse(json_str, max_retries=3):
    """Safely parse JSON with error handling and cleaning."""
    for attempt in range(max_retries):
        try:
            # First attempt - try parsing as-is
            if attempt == 0:
                return json.loads(json_str)
            
            # Second attempt - clean the string
            elif attempt == 1:
                cleaned_str = clean_json_string(json_str)
                print(f"🔧 Attempt {attempt + 1}: Trying with cleaned JSON string")
                return json.loads(cleaned_str)
            
            # Third attempt - more aggressive cleaning
            else:
                # Try to extract just the JSON part if it's embedded in other text
                json_match = re.search(r'\{.*\}|\[.*\]', json_str, re.DOTALL)
                if json_match:
                    json_part = json_match.group()
                    cleaned_part = clean_json_string(json_part)
                    print(f"🔧 Attempt {attempt + 1}: Trying with extracted and cleaned JSON")
                    return json.loads(cleaned_part)
                else:
                    raise ValueError("No valid JSON structure found")
                    
        except json.JSONDecodeError as e:
            print(f"❌ JSON parse attempt {attempt + 1} failed: {e}")
            if attempt == max_retries - 1:
                print("📝 Problematic JSON content (first 500 chars):")
                print(repr(json_str[:500]))
                raise e
            continue
    
    return None

def get_provider_suffix(provider):
    """Convert provider name to database table suffix."""
    print(f"🐛 DEBUG: Converting provider '{provider}' to suffix...")
    
    # Standardize the mapping to ensure consistent table naming
    provider_map = {
        "OpenAI": "openai",
        "GoogleGemini": "gemini",
        "Gemini": "gemini",  # Alternative name
        "Google": "gemini"   # Another alternative
    }
    
    suffix = provider_map.get(provider, provider.lower())
    print(f"🐛 DEBUG: Provider '{provider}' -> Suffix '{suffix}'")
    return suffix

def get_database_connection():
    """Get database connection using environment variables."""
    try:
        conn = psycopg2.connect(
            host=os.getenv('POSTGRES_HOST'),
            port=os.getenv('POSTGRES_PORT'),
            database=os.getenv('POSTGRES_DB'),
            user=os.getenv('POSTGRES_USER'),
            password=os.getenv('POSTGRES_PASSWORD')
        )
        return conn
    except Exception as e:
        print(f"Failed to connect to database: {e}")
        sys.exit(1)

def get_selected_submissions_from_db(selected_submission_ids=None, assessment_id=None):
    """
    Retrieve specific submissions from database with full mapping data.
    Filter by submission IDs and/or assessment ID.
    """
    conn = get_database_connection()
    try:
        with conn.cursor(cursor_factory=RealDictCursor) as cur:
            if selected_submission_ids and assessment_id:
                # Filter by both submission IDs and assessment ID
                submission_ids_tuple = tuple(selected_submission_ids)
                placeholders = ','.join(['%s'] * len(submission_ids_tuple))
                
                cur.execute(f"""
                    SELECT s.submission_id, s.assessment_id, s.student_id, s.file_url, 
                           s.submission_start_at,
                           st.registration_number, st.user_id as student_user_id,
                           a.assessment_id, a.created_on as assessment_created_on,
                           m.module_code, m.module_name
                    FROM "Submission" s
                    JOIN "Student" st ON s.student_id = st.user_id
                    JOIN "Assessment" a ON s.assessment_id = a.assessment_id
                    JOIN "Module" m ON a.module_id = m.module_id
                    WHERE s.submission_id IN ({placeholders}) AND s.assessment_id = %s
                    ORDER BY s.submission_start_at ASC
                """, submission_ids_tuple + (assessment_id,))
                
            elif selected_submission_ids:
                # Filter by submission IDs only
                submission_ids_tuple = tuple(selected_submission_ids)
                placeholders = ','.join(['%s'] * len(submission_ids_tuple))
                
                cur.execute(f"""
                    SELECT s.submission_id, s.assessment_id, s.student_id, s.file_url, 
                           s.submission_start_at,
                           st.registration_number, st.user_id as student_user_id,
                           a.assessment_id, a.created_on as assessment_created_on,
                           m.module_code, m.module_name
                    FROM "Submission" s
                    JOIN "Student" st ON s.student_id = st.user_id
                    JOIN "Assessment" a ON s.assessment_id = a.assessment_id
                    JOIN "Module" m ON a.module_id = m.module_id
                    WHERE s.submission_id IN ({placeholders})
                    ORDER BY s.submission_start_at ASC
                """, submission_ids_tuple)
                
            elif assessment_id:
                # Filter by assessment ID only
                cur.execute("""
                    SELECT s.submission_id, s.assessment_id, s.student_id, s.file_url, 
                           s.submission_start_at,
                           st.registration_number, st.user_id as student_user_id,
                           a.assessment_id, a.created_on as assessment_created_on,
                           m.module_code, m.module_name
                    FROM "Submission" s
                    JOIN "Student" st ON s.student_id = st.user_id
                    JOIN "Assessment" a ON s.assessment_id = a.assessment_id
                    JOIN "Module" m ON a.module_id = m.module_id
                    WHERE s.assessment_id = %s
                    ORDER BY s.submission_start_at ASC
                """, (assessment_id,))
            else:
                # Get all submissions (fallback)
                cur.execute("""
                    SELECT s.submission_id, s.assessment_id, s.student_id, s.file_url, 
                           s.submission_start_at,
                           st.registration_number, st.user_id as student_user_id,
                           a.assessment_id, a.created_on as assessment_created_on,
                           m.module_code, m.module_name
                    FROM "Submission" s
                    JOIN "Student" st ON s.student_id = st.user_id
                    JOIN "Assessment" a ON s.assessment_id = a.assessment_id
                    JOIN "Module" m ON a.module_id = m.module_id
                    ORDER BY s.submission_start_at ASC
                """)
            
            submissions = cur.fetchall()
            print(f"Found {len(submissions)} submissions matching criteria")
            return [dict(sub) for sub in submissions]
            
    except Exception as e:
        print(f"Database error: {e}")
        return []
    finally:
        conn.close()

def check_submission_already_extracted(submission_id, provider_suffix):
    """Check if a specific submission has already been extracted."""
    conn = get_database_connection()
    try:
        with conn.cursor() as cur:
            answers_table = f"student_answers_{provider_suffix}"
            
            # Check if the table exists first
            cur.execute("""
                SELECT EXISTS (
                    SELECT FROM information_schema.tables 
                    WHERE table_name = %s
                );
            """, (answers_table,))
            
            table_exists = cur.fetchone()[0]
            if not table_exists:
                return False
            
            # Check if submission_id column exists
            cur.execute("""
                SELECT EXISTS (
                    SELECT FROM information_schema.columns 
                    WHERE table_name = %s AND column_name = 'submission_id'
                );
            """, (answers_table,))
            
            column_exists = cur.fetchone()[0]
            if not column_exists:
                return False
            
            # Check if this submission has been processed
            cur.execute(f"""
                SELECT EXISTS (
                    SELECT 1 FROM "{answers_table}" 
                    WHERE submission_id = %s
                );
            """, (submission_id,))
            
            return cur.fetchone()[0]
            
    except Exception as e:
        print(f"Error checking extraction status for submission {submission_id}: {e}")
        return False
    finally:
        conn.close()

def resolve_file_path(file_path: str) -> pathlib.Path:
    """Resolve database file path to actual file location."""
    project_root = pathlib.Path.cwd()
    # Try parent directory first (where data folder should be)
    full_path = project_root.parent / file_path
    
    if not full_path.exists():
        # Try alternative path in project directory
        alternative_path = project_root / file_path
        if alternative_path.exists():
            full_path = alternative_path
        else:
            raise FileNotFoundError(f"File not found at: {full_path} or {alternative_path}")
    
    return full_path

def load_docx_text(docx_path: str) -> str:
    doc = Document(docx_path)
    return "\n".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])

def get_provider_suffix(provider: str) -> str:
    """Map provider names to database suffixes"""
    mapping = {
        "DeepSeek": "deepseek",
        "GoogleGemini": "gemini",
        "OpenAI": "openai"
    }
    return mapping.get(provider, provider.lower())

def extract_and_save(docx_path: str, extractor: AnswerExtractor, provider: str):
    filename = os.path.basename(docx_path)
    print(f"\n📄 Processing: {filename}")
    print(f"🔗 Submission ID: {submission_id}")
    print(f"👤 Student: {student_registration_number}")
    print(f"📚 Module: {module_code}")
    print(f"🆔 Assessment: {assessment_id}")

    try:
        # Check if already extracted
        if check_submission_already_extracted(submission_id, provider_suffix):
            print(f"⚠️ Submission {submission_id} already extracted, skipping...")
            return True

        # Load file content
        full_path = resolve_file_path(file_path)
        raw_text = load_file_text(str(full_path))
        
        # Enhanced extraction with robust error handling
        answers = None
        max_extraction_attempts = 3
        
        for attempt in range(max_extraction_attempts):
            try:
                print(f"🔄 Extraction attempt {attempt + 1}/{max_extraction_attempts}")
                answers = extractor.extract_answers_with_llm(raw_text)
                
                if answers:
                    print(f"✅ Successfully extracted {len(answers)} answers on attempt {attempt + 1}")
                    break
                else:
                    print(f"⚠️ No answers extracted on attempt {attempt + 1}")
                    
            except json.JSONDecodeError as json_error:
                print(f"❌ JSON parsing error on attempt {attempt + 1}: {json_error}")
                
                if "Invalid control character" in str(json_error):
                    print("🔧 Detected control character issue - this might be due to LLM response formatting")
                    
                    # If this is not the last attempt, wait and retry
                    if attempt < max_extraction_attempts - 1:
                        print(f"⏳ Waiting 10 seconds before retry...")
                        time.sleep(10)
                        continue
                    else:
                        print("💥 All extraction attempts failed due to JSON parsing errors")
                        answers = None
                        break
                else:
                    # For other JSON errors, re-raise immediately
                    raise json_error
                    
            except Exception as other_error:
                print(f"❌ Other error on attempt {attempt + 1}: {other_error}")
                if attempt == max_extraction_attempts - 1:
                    raise other_error
                else:
                    print(f"⏳ Waiting 10 seconds before retry...")
                    time.sleep(10)
                    continue

        if not answers:
            print("❌ No answers extracted after all attempts.")
            return False

        # Override extracted metadata with database-mapped data
        year = assessment_created_on.year
        month = assessment_created_on.strftime('%B')
        
        print(f"🔄 Overriding extracted metadata with database values:")
        print(f"   Student Index: {student_registration_number}")
        print(f"   Module Code: {module_code}")
        print(f"   Year: {year}")
        print(f"   Month: {month}")
        
        # Update all answers with database-mapped metadata
        for answer in answers:
            answer.student_index = student_registration_number
            answer.module_code = module_code.upper()  # Normalize to uppercase
            answer.exam_year = year
            answer.exam_month = month

        # Preview the result with database-mapped data
        print("\nAnswers with database-mapped metadata:")
        pprint([
            {
                "question": ans.full_question_id, 
                "answer": ans.answer_text[:100] + "..." if len(ans.answer_text) > 100 else ans.answer_text,
                "student_index": ans.student_index,
                "module_code": ans.module_code,
                "year": ans.exam_year,
                "month": ans.exam_month
            }
            for ans in answers[:3]  # Show first 3 for preview
        ])

        # Save to database - use the mapped provider suffix from extractor
        first = answers[0]
        provider_suffix = extractor.provider_suffix
        db = StudentAnswerService(provider_suffix=provider_suffix)
        db.initialize_table()
        
        # Save answers with submission tracking
        db.save_answers_with_submission_tracking(
            student_index=student_registration_number,
            module_code=module_code.upper(),
            year=year,
            month=month,
            answers=answers,
            submission_id=submission_id,
            assessment_id=assessment_id
        )
        db.close()

        print(f"✅ Saved answers for {student_registration_number} | {module_code} | {year}-{month} | Submission: {submission_id}")
        return True
        
    except Exception as e:
        print(f"❌ Failed to process {filename}: {e}")

if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="Extract and save student answers using LLMs")
    parser.add_argument("--provider", required=True, choices=["OpenAI", "GoogleGemini", "DeepSeek"], help="LLM provider")
    parser.add_argument("--model", required=True, help="Model name (e.g., gpt-4o, gemini-2.0-flash, deepseek-r1:7b)")
    parser.add_argument("--folder", required=True, help="Single DOCX file or folder containing DOCX files")

    args = parser.parse_args()
    extractor = AnswerExtractor(selected_provider=args.provider, selected_model=args.model)

    if os.path.isfile(args.folder) and args.folder.endswith(".docx"):
        # Single file mode
        extract_and_save(args.folder, extractor, args.provider)
    elif os.path.isdir(args.folder):
        # Folder mode
        for filename in os.listdir(args.folder):
            if filename.lower().endswith(".docx"):
                filepath = os.path.join(args.folder, filename)
                extract_and_save(filepath, extractor, args.provider)

                # Delay to respect Gemini rate limits (15 requests/min)
                if args.provider == "GoogleGemini":
                    time.sleep(10)
                # Add delay for DeepSeek if needed
                elif args.provider == "DeepSeek":
                    time.sleep(5)  # Adjust as needed for DeepSeek rate limits
    else:
        print("❌ Invalid --folder path. Must be either a .docx file or a directory.")