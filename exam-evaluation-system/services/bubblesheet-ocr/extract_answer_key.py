# # D:\FYP Mid\AI-Exam-Evaluation-System\exam-evaluation-system\services\bubblesheet-ocr\extract_answer_key.py
# """
# Answer Key Extractor for Bubble Sheet Assessments
# Extracts correct answers from PDF, Word, or Excel files
# """

# import re
# import json
# import sys
# from pathlib import Path
# from typing import List, Dict, Optional

# try:
#     import PyPDF2
# except ImportError:
#     PyPDF2 = None

# try:
#     from docx import Document
# except ImportError:
#     Document = None

# try:
#     import pandas as pd
# except ImportError:
#     pd = None


# class AnswerKeyExtractor:
#     """Extract answer keys from Word, PDF, or Excel files"""
    
#     def __init__(self):
#         # Common patterns for answer key format
#         self.patterns = [
#             r'(?:Q|Question|q)?\s*(\d+)[.:\s)]+([A-Da-d])',  # Q1. A or 1: A
#             r'(\d+)\s*[-–—]\s*([A-Da-d])',  # 1 - A
#             r'(\d+)\s*[=:]\s*([A-Da-d])',  # 1 = A or 1: A
#             r'(\d+)\)\s*([A-Da-d])',  # 1) A
#             r'(\d+)\s+([A-Da-d])\b',  # 1 A
#         ]
    
#     def extract_from_pdf(self, file_path: str) -> List[Dict[str, any]]:
#         """Extract answer key from PDF"""
#         if PyPDF2 is None:
#             raise ImportError("PyPDF2 is required for PDF extraction. Install: pip install PyPDF2")
        
#         answers = []
        
#         try:
#             with open(file_path, 'rb') as file:
#                 pdf_reader = PyPDF2.PdfReader(file)
#                 text = ""
                
#                 for page in pdf_reader.pages:
#                     text += page.extract_text()
                
#                 answers = self._parse_text(text)
                
#         except Exception as e:
#             print(f"Error reading PDF: {e}", file=sys.stderr)
            
#         return answers
    
#     def extract_from_docx(self, file_path: str) -> List[Dict[str, any]]:
#         """Extract answer key from Word document"""
#         if Document is None:
#             raise ImportError("python-docx is required for Word extraction. Install: pip install python-docx")
        
#         answers = []
        
#         try:
#             doc = Document(file_path)
#             text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
#             # Also check tables
#             for table in doc.tables:
#                 for row in table.rows:
#                     text += "\n" + "\t".join([cell.text for cell in row.cells])
            
#             answers = self._parse_text(text)
            
#         except Exception as e:
#             print(f"Error reading DOCX: {e}", file=sys.stderr)
            
#         return answers
    
#     def extract_from_excel(self, file_path: str) -> List[Dict[str, any]]:
#         """Extract answer key from Excel file"""
#         if pd is None:
#             raise ImportError("pandas and openpyxl are required for Excel extraction. Install: pip install pandas openpyxl")
        
#         answers = []
        
#         try:
#             df = pd.read_excel(file_path)
            
#             # Try to find question and answer columns
#             question_col = None
#             answer_col = None
            
#             # Look for columns with names like 'question', 'q', 'no', 'number'
#             for col in df.columns:
#                 col_lower = str(col).lower()
#                 if any(term in col_lower for term in ['question', 'q', 'no', 'number', '#']):
#                     question_col = col
#                 elif any(term in col_lower for term in ['answer', 'correct', 'key', 'option']):
#                     answer_col = col
            
#             if question_col and answer_col:
#                 for idx, row in df.iterrows():
#                     q_num = row[question_col]
#                     ans = row[answer_col]
                    
#                     # Clean and validate
#                     try:
#                         q_num = int(float(q_num))
#                         ans = str(ans).strip().upper()
                        
#                         if ans in ['A', 'B', 'C', 'D']:
#                             answers.append({
#                                 'question_number': q_num,
#                                 'correct_option': ans
#                             })
#                     except (ValueError, TypeError):
#                         continue
#             else:
#                 # Fallback: try first two columns
#                 for idx, row in df.iterrows():
#                     try:
#                         q_num = int(float(row.iloc[0]))
#                         ans = str(row.iloc[1]).strip().upper()
                        
#                         if ans in ['A', 'B', 'C', 'D']:
#                             answers.append({
#                                 'question_number': q_num,
#                                 'correct_option': ans
#                             })
#                     except (ValueError, TypeError, IndexError):
#                         continue
                        
#         except Exception as e:
#             print(f"Error reading Excel: {e}", file=sys.stderr)
            
#         return answers
    
#     def _parse_text(self, text: str) -> List[Dict[str, any]]:
#         """Parse text to extract question-answer pairs"""
#         answers = []
#         found_pairs = {}  # Track by question number to avoid duplicates
        
#         for pattern in self.patterns:
#             matches = re.finditer(pattern, text, re.IGNORECASE | re.MULTILINE)
            
#             for match in matches:
#                 q_num = int(match.group(1))
#                 ans = match.group(2).upper()
                
#                 if ans in ['A', 'B', 'C', 'D']:
#                     # Only keep first occurrence of each question number
#                     if q_num not in found_pairs:
#                         found_pairs[q_num] = ans
        
#         # Convert to list and sort
#         answers = [
#             {'question_number': q_num, 'correct_option': ans}
#             for q_num, ans in sorted(found_pairs.items())
#         ]
        
#         return answers
    
#     def extract_answer_key(self, file_path: str) -> Optional[List[Dict[str, any]]]:
#         """Main method to extract answer key based on file type"""
#         file_path = Path(file_path)
        
#         if not file_path.exists():
#             print(f"Error: File not found: {file_path}", file=sys.stderr)
#             return None
        
#         extension = file_path.suffix.lower()
        
#         try:
#             if extension == '.pdf':
#                 answers = self.extract_from_pdf(str(file_path))
#             elif extension in ['.docx', '.doc']:
#                 answers = self.extract_from_docx(str(file_path))
#             elif extension in ['.xlsx', '.xls']:
#                 answers = self.extract_from_excel(str(file_path))
#             else:
#                 print(f"Error: Unsupported file type: {extension}", file=sys.stderr)
#                 return None
#         except ImportError as e:
#             print(f"Error: {e}", file=sys.stderr)
#             return None
        
#         if not answers:
#             print("Warning: No answer key found in the document", file=sys.stderr)
#             return None
        
#         print(f"✓ Extracted {len(answers)} answers", file=sys.stderr)
#         return answers
    
#     def validate_answer_key(self, answers: List[Dict[str, any]], 
#                            expected_count: Optional[int] = None) -> bool:
#         """Validate extracted answer key"""
#         if not answers:
#             return False
        
#         # Check for sequential question numbers
#         question_numbers = [a['question_number'] for a in answers]
        
#         if question_numbers != sorted(question_numbers):
#             print("Warning: Question numbers are not in order", file=sys.stderr)
        
#         # Check for duplicates
#         if len(question_numbers) != len(set(question_numbers)):
#             print("Error: Duplicate question numbers found", file=sys.stderr)
#             return False
        
#         # Check expected count
#         if expected_count and len(answers) != expected_count:
#             print(f"Warning: Expected {expected_count} answers but found {len(answers)}", file=sys.stderr)
        
#         # Validate all answers are A, B, C, or D
#         for ans in answers:
#             if ans['correct_option'] not in ['A', 'B', 'C', 'D']:
#                 print(f"Error: Invalid answer '{ans['correct_option']}' for Q{ans['question_number']}", file=sys.stderr)
#                 return False
        
#         return True


# def main():
#     """CLI interface for answer key extraction"""
#     if len(sys.argv) < 2:
#         print("Usage: python3 extract_answer_key.py <file_path> [--json]")
#         print("\nSupported formats: PDF, Word (.docx), Excel (.xlsx)")
#         sys.exit(1)
    
#     file_path = sys.argv[1]
#     output_json = '--json' in sys.argv
    
#     extractor = AnswerKeyExtractor()
#     answers = extractor.extract_answer_key(file_path)
    
#     if not answers:
#         sys.exit(1)
    
#     if not extractor.validate_answer_key(answers):
#         print("\nWarning: Validation issues detected. Please review the extracted answers.", file=sys.stderr)
    
#     # Output results
#     if output_json:
#         # Output JSON format for Node.js to parse
#         print(json.dumps(answers, indent=2))
#     else:
#         # Human-readable format
#         print("\nExtracted Answer Key:")
#         print("-" * 40)
#         for ans in answers:
#             print(f"Q{ans['question_number']}: {ans['correct_option']}")
#         print("-" * 40)
#         print(f"Total: {len(answers)} answers")


# if __name__ == "__main__":
#     main()

"""
Answer Key Extractor for Bubble Sheet Assessments
Extracts correct answers from PDF, Word, or Excel files
"""

import re
import json
import sys
import warnings
from pathlib import Path
from typing import List, Dict, Optional

# Suppress warnings
warnings.filterwarnings('ignore')

# Check for required packages
MISSING_PACKAGES = []

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None
    MISSING_PACKAGES.append('PyPDF2')

try:
    from docx import Document
except ImportError:
    Document = None
    MISSING_PACKAGES.append('python-docx')

try:
    import pandas as pd
except ImportError:
    pd = None
    MISSING_PACKAGES.append('pandas')

try:
    import openpyxl
except ImportError:
    if pd is not None:
        MISSING_PACKAGES.append('openpyxl')

if MISSING_PACKAGES:
    print(f"⚠️ WARNING: Missing packages: {', '.join(MISSING_PACKAGES)}", file=sys.stderr)
    print(f"Install with: pip install {' '.join(MISSING_PACKAGES)}", file=sys.stderr)


class AnswerKeyExtractor:
    """Extract answer keys from Word, PDF, or Excel files"""
    
    def __init__(self):
        # Common patterns for answer key format
        self.patterns = [
            r'(?:Q|Question|q)?\s*(\d+)[.:\s)]+([A-Da-d])',  # Q1. A or 1: A
            r'(\d+)\s*[-–—]\s*([A-Da-d])',  # 1 - A
            r'(\d+)\s*[=:]\s*([A-Da-d])',  # 1 = A or 1: A
            r'(\d+)\)\s*([A-Da-d])',  # 1) A
            r'(\d+)\s+([A-Da-d])\b',  # 1 A
        ]
    
    def extract_from_pdf(self, file_path: str) -> List[Dict[str, any]]:
        """Extract answer key from PDF"""
        if PyPDF2 is None:
            raise ImportError("PyPDF2 is required for PDF extraction. Install: pip install PyPDF2")
        
        answers = []
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                
                print(f"📄 PDF has {len(pdf_reader.pages)} pages", file=sys.stderr)
                
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    page_text = page.extract_text()
                    text += page_text
                    print(f"   Page {page_num}: {len(page_text)} characters", file=sys.stderr)
                
                print(f"📝 Total extracted text: {len(text)} characters", file=sys.stderr)
                answers = self._parse_text(text)
                
        except Exception as e:
            print(f"❌ Error reading PDF: {e}", file=sys.stderr)
            raise
            
        return answers
    
    def extract_from_docx(self, file_path: str) -> List[Dict[str, any]]:
        """Extract answer key from Word document"""
        if Document is None:
            raise ImportError("python-docx is required for Word extraction. Install: pip install python-docx")
        
        answers = []
        
        try:
            doc = Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            print(f"📄 Document has {len(doc.paragraphs)} paragraphs", file=sys.stderr)
            
            # Also check tables
            if doc.tables:
                print(f"📊 Document has {len(doc.tables)} tables", file=sys.stderr)
                for table in doc.tables:
                    for row in table.rows:
                        text += "\n" + "\t".join([cell.text for cell in row.cells])
            
            print(f"📝 Total extracted text: {len(text)} characters", file=sys.stderr)
            answers = self._parse_text(text)
            
        except Exception as e:
            print(f"❌ Error reading DOCX: {e}", file=sys.stderr)
            raise
            
        return answers
    
    def extract_from_excel(self, file_path: str) -> List[Dict[str, any]]:
        """Extract answer key from Excel file"""
        if pd is None:
            raise ImportError("pandas is required for Excel extraction. Install: pip install pandas openpyxl")
        
        answers = []
        
        try:
            # Read Excel file
            df = pd.read_excel(file_path, engine='openpyxl')
            
            print(f"📊 Excel shape: {df.shape[0]} rows x {df.shape[1]} columns", file=sys.stderr)
            print(f"   Columns: {list(df.columns)}", file=sys.stderr)
            
            # Try to find question and answer columns
            question_col = None
            answer_col = None
            
            # Look for columns with names like 'question', 'q', 'no', 'number'
            for col in df.columns:
                col_lower = str(col).lower()
                if any(term in col_lower for term in ['question', 'q', 'no', 'number', '#']):
                    if question_col is None:
                        question_col = col
                        print(f"   Found question column: {col}", file=sys.stderr)
                elif any(term in col_lower for term in ['answer', 'correct', 'key', 'option']):
                    if answer_col is None:
                        answer_col = col
                        print(f"   Found answer column: {col}", file=sys.stderr)
            
            if question_col and answer_col:
                print(f"✅ Using columns: {question_col} -> {answer_col}", file=sys.stderr)
                for idx, row in df.iterrows():
                    q_num = row[question_col]
                    ans = row[answer_col]
                    
                    # Clean and validate
                    try:
                        q_num = int(float(q_num))
                        ans = str(ans).strip().upper()
                        
                        if ans in ['A', 'B', 'C', 'D']:
                            answers.append({
                                'question_number': q_num,
                                'correct_option': ans
                            })
                    except (ValueError, TypeError):
                        continue
            else:
                # Fallback: try first two columns
                print(f"⚠️ Using fallback: first two columns", file=sys.stderr)
                for idx, row in df.iterrows():
                    try:
                        q_num = int(float(row.iloc[0]))
                        ans = str(row.iloc[1]).strip().upper()
                        
                        if ans in ['A', 'B', 'C', 'D']:
                            answers.append({
                                'question_number': q_num,
                                'correct_option': ans
                            })
                    except (ValueError, TypeError, IndexError):
                        continue
                        
        except Exception as e:
            print(f"❌ Error reading Excel: {e}", file=sys.stderr)
            raise
            
        return answers
    
    def _parse_text(self, text: str) -> List[Dict[str, any]]:
        """Parse text to extract question-answer pairs"""
        answers = []
        found_pairs = {}  # Track by question number to avoid duplicates
        
        print(f"🔍 Parsing text with {len(self.patterns)} patterns...", file=sys.stderr)
        
        for pattern_idx, pattern in enumerate(self.patterns):
            matches = list(re.finditer(pattern, text, re.IGNORECASE | re.MULTILINE))
            
            if matches:
                print(f"   Pattern {pattern_idx + 1}: Found {len(matches)} matches", file=sys.stderr)
            
            for match in matches:
                q_num = int(match.group(1))
                ans = match.group(2).upper()
                
                if ans in ['A', 'B', 'C', 'D']:
                    # Only keep first occurrence of each question number
                    if q_num not in found_pairs:
                        found_pairs[q_num] = ans
        
        # Convert to list and sort
        answers = [
            {'question_number': q_num, 'correct_option': ans}
            for q_num, ans in sorted(found_pairs.items())
        ]
        
        print(f"✅ Found {len(answers)} unique question-answer pairs", file=sys.stderr)
        
        return answers
    
    def extract_answer_key(self, file_path: str) -> Optional[List[Dict[str, any]]]:
        """Main method to extract answer key based on file type"""
        file_path = Path(file_path)
        
        if not file_path.exists():
            print(f"❌ Error: File not found: {file_path}", file=sys.stderr)
            return None
        
        extension = file_path.suffix.lower()
        print(f"📁 Processing file: {file_path.name} ({extension})", file=sys.stderr)
        
        try:
            if extension == '.pdf':
                if PyPDF2 is None:
                    raise ImportError("PyPDF2 not installed")
                answers = self.extract_from_pdf(str(file_path))
            elif extension in ['.docx', '.doc']:
                if Document is None:
                    raise ImportError("python-docx not installed")
                answers = self.extract_from_docx(str(file_path))
            elif extension in ['.xlsx', '.xls', '.xlsm']:
                if pd is None:
                    raise ImportError("pandas not installed")
                answers = self.extract_from_excel(str(file_path))
            else:
                print(f"❌ Error: Unsupported file type: {extension}", file=sys.stderr)
                return None
        except ImportError as e:
            print(f"❌ Error: {e}", file=sys.stderr)
            print(f"Install required packages: pip install {' '.join(MISSING_PACKAGES)}", file=sys.stderr)
            raise
        except Exception as e:
            print(f"❌ Error processing file: {e}", file=sys.stderr)
            raise
        
        if not answers:
            print("⚠️ Warning: No answer key found in the document", file=sys.stderr)
            return None
        
        print(f"✅ Successfully extracted {len(answers)} answers", file=sys.stderr)
        return answers
    
    def validate_answer_key(self, answers: List[Dict[str, any]], 
                           expected_count: Optional[int] = None) -> bool:
        """Validate extracted answer key"""
        if not answers:
            print("❌ Validation failed: No answers", file=sys.stderr)
            return False
        
        # Check for sequential question numbers
        question_numbers = [a['question_number'] for a in answers]
        
        if question_numbers != sorted(question_numbers):
            print("⚠️ Warning: Question numbers are not in order", file=sys.stderr)
        
        # Check for duplicates
        if len(question_numbers) != len(set(question_numbers)):
            print("❌ Error: Duplicate question numbers found", file=sys.stderr)
            return False
        
        # Check for gaps
        expected_sequence = list(range(1, len(answers) + 1))
        if question_numbers != expected_sequence:
            missing = set(expected_sequence) - set(question_numbers)
            if missing:
                print(f"⚠️ Warning: Missing question numbers: {sorted(missing)}", file=sys.stderr)
        
        # Check expected count
        if expected_count and len(answers) != expected_count:
            print(f"⚠️ Warning: Expected {expected_count} answers but found {len(answers)}", file=sys.stderr)
        
        # Validate all answers are A, B, C, or D
        for ans in answers:
            if ans['correct_option'] not in ['A', 'B', 'C', 'D']:
                print(f"❌ Error: Invalid answer '{ans['correct_option']}' for Q{ans['question_number']}", file=sys.stderr)
                return False
        
        print(f"✅ Validation passed: {len(answers)} valid answers", file=sys.stderr)
        return True


def main():
    """CLI interface for answer key extraction"""
    if len(sys.argv) < 2:
        print("Usage: python extract_answer_key.py <file_path> [--json]")
        print("\nSupported formats: PDF, Word (.docx), Excel (.xlsx)")
        print("\nExamples:")
        print("  python extract_answer_key.py answer_key.pdf")
        print("  python extract_answer_key.py answer_key.xlsx --json")
        sys.exit(1)
    
    file_path = sys.argv[1]
    output_json = '--json' in sys.argv
    
    extractor = AnswerKeyExtractor()
    
    try:
        answers = extractor.extract_answer_key(file_path)
    except Exception as e:
        print(f"\n❌ Failed to extract answer key: {e}", file=sys.stderr)
        sys.exit(1)
    
    if not answers:
        print("\n❌ No answers found in the document", file=sys.stderr)
        sys.exit(1)
    
    if not extractor.validate_answer_key(answers):
        print("\n⚠️ Warning: Validation issues detected. Please review the extracted answers.", file=sys.stderr)
    
    # Output results
    if output_json:
        # Output JSON format for Node.js to parse
        print(json.dumps(answers, indent=2))
    else:
        # Human-readable format
        print("\n" + "="*50)
        print("EXTRACTED ANSWER KEY")
        print("="*50)
        for ans in answers:
            print(f"Q{ans['question_number']:3d}: {ans['correct_option']}")
        print("="*50)
        print(f"Total: {len(answers)} answers")
        print()


if __name__ == "__main__":
    main()