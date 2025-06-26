# # import sys
# # import os
# # from docx import Document

# # # Set path to root so we can import correctly
# # sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "../../")))

# # from src.services.answer_extractor import AnswerExtractor
# # from src.services.answer_extractor import StudentAnswerService

# # def load_docx_text(docx_path: str) -> str:
# #     doc = Document(docx_path)
# #     return "\n".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])

# # if __name__ == "__main__":
# #     import argparse
# #     parser = argparse.ArgumentParser()
# #     parser.add_argument("--file", required=True, help="Path to student answer DOCX")
# #     parser.add_argument("--provider", required=True, choices=["OpenAI", "GoogleGemini"])
# #     parser.add_argument("--model", required=True, help="e.g. gpt-4o or gemini-pro")
# #     parser.add_argument("--student_index", required=True, help="Student Index Number")

# #     args = parser.parse_args()

# #     # Load DOCX
# #     raw_text = load_docx_text(args.file)

# #     # Extract answers
# #     extractor = AnswerExtractor(args.provider, args.model)
# #     answers = extractor.extract_answers_with_llm(raw_text, student_index=args.student_index)

# #     # Save to DB
# #     db = StudentAnswerService()
# #     db.initialize_table()
# #     db.save_student_answers(args.student_index, answers)
# #     db.close()

# #     print(f"\n✅ Successfully extracted and saved answers for: {args.student_index}")


# import sys
# import os
# from docx import Document

# # Ensure Python can locate the src package
# sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "../../")))

# from src.services.answer_extractor import AnswerExtractor
# from src.services.database_services.student_answer_db import StudentAnswerService


# def load_docx_text(docx_path: str) -> str:
#     """Extract full text from a DOCX file."""
#     doc = Document(docx_path)
#     return "\n".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])

# if __name__ == "__main__":
#     import argparse

#     parser = argparse.ArgumentParser(description="Extract and save student answers using LLMs")
#     parser.add_argument("--provider", required=True, choices=["OpenAI", "GoogleGemini"], help="LLM provider")
#     parser.add_argument("--model", required=True, help="Model name (e.g., gpt-4o, gemini-pro)")
#     parser.add_argument("--file", default="data/Answer_Scripts/student_answer1.docx", help="Path to DOCX answer script")
#     parser.add_argument("--index", required=True, help="Student index number")

#     args = parser.parse_args()

#     # 1. Extract raw text from the docx
#     raw_text = load_docx_text(args.file)

#     # 2. Extract answers using LLM
#     extractor = AnswerExtractor(
#         selected_provider=args.provider,
#         selected_model=args.model
#     )
#     answers = extractor.extract_answers_with_llm(raw_text)

#     # 3. Save to DB
#     db = StudentAnswerService()

#     db.initialize_table()
#     db.save_student_answers(student_index=args.index, answers=answers)
#     db.close()

#     print(f"\nSuccessfully extracted and saved answers for student index: {args.index}")

import sys
import os
from docx import Document

# Ensure Python can locate the src package
# sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "../../")))

# from src.services.answer_extractor import AnswerExtractor
# from src.services.database_services.student_answer_db import StudentAnswerService

# def load_docx_text(docx_path: str) -> str:
#     doc = Document(docx_path)
#     return "\n".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])

# if __name__ == "__main__":
#     import argparse

#     parser = argparse.ArgumentParser(description="Extract and save student answers using LLMs")
#     parser.add_argument("--provider", required=True, choices=["OpenAI", "GoogleGemini"], help="LLM provider")
#     parser.add_argument("--model", required=True, help="Model name (e.g., gpt-4o, gemini-pro)")
#     parser.add_argument("--file", default="data/Answer_Scripts/student_answer1.docx", help="Path to DOCX answer script")

#     args = parser.parse_args()

#     # 1. Load answer script text
#     raw_text = load_docx_text(args.file)

#     # 2. Extract answers (and metadata) using LLM
#     extractor = AnswerExtractor(selected_provider=args.provider, selected_model=args.model)
#     answers = extractor.extract_answers_with_llm(raw_text)

#     if not answers:
#         print(" No answers extracted.")
#         sys.exit(1)

#     # 3. Save answers to DB
#     first = answers[0]
#     db = StudentAnswerService()
#     db.initialize_table()
#     db.save_answers(
#         student_index=first.student_index,
#         module_code=first.module_code,
#         year=first.exam_year,
#         month=first.exam_month,
#         answers=answers
#     )
#     db.close()

#     print(f" Saved answers for {first.student_index} | {first.module_code} | {first.exam_year}-{first.exam_month}")

import sys
import os
from docx import Document
from pprint import pprint

sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "../../")))

from src.services.answer_extractor import AnswerExtractor
from src.services.database_services.student_answer_db import StudentAnswerService

def load_docx_text(docx_path: str) -> str:
    doc = Document(docx_path)
    return "\n".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])

if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="Extract and save student answers using LLMs")
    parser.add_argument("--provider", required=True, choices=["OpenAI", "GoogleGemini"], help="LLM provider")
    parser.add_argument("--model", required=True, help="Model name (e.g., gpt-4o, gemini-pro)")
    parser.add_argument("--file", default="data/Answer_Scripts/EE6250_EG-2020-4247.docx", help="Path to DOCX answer script")

    args = parser.parse_args()

    # 1. Load answer script text
    raw_text = load_docx_text(args.file)

    # 2. Extract answers (and metadata) using LLM
    extractor = AnswerExtractor(selected_provider=args.provider, selected_model=args.model)
    answers = extractor.extract_answers_with_llm(raw_text)

    if not answers:
        print("❌ No answers extracted.")
        sys.exit(1)

    # ✅ Print extracted answers for verification
    print("\n🧾 Extracted Answers (before saving to DB):\n")
    pprint([
        {
            "question": ans.full_question_id,
            "answer": ans.answer_text
        }
        for ans in answers
    ])

    # 3. Save answers to DB
    first = answers[0]
    db = StudentAnswerService()
    db.initialize_table()
    db.save_answers(
        student_index=first.student_index,
        module_code=first.module_code,
        year=first.exam_year,
        month=first.exam_month,
        answers=answers
    )
    db.close()

    print(f"\n✅ Saved answers for {first.student_index} | {first.module_code} | {first.exam_year}-{first.exam_month}")
