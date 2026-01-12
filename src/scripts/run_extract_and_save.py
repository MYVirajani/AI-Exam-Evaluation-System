

# # # # # # # # # # import sys
# # # # # # # # # # import os
# # # # # # # # # # from docx import Document
# # # # # # # # # # from pprint import pprint

# # # # # # # # # # sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "../../")))

# # # # # # # # # # from src.services.answer_extractor import AnswerExtractor
# # # # # # # # # # from src.services.database_services.student_answer_db import StudentAnswerService

# # # # # # # # # # def load_docx_text(docx_path: str) -> str:
# # # # # # # # # #     doc = Document(docx_path)
# # # # # # # # # #     return "\n".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])

# # # # # # # # # # if __name__ == "__main__":
# # # # # # # # # #     import argparse
# # # # # # # # # #     parser = argparse.ArgumentParser(description="Extract and save ALL student answers using LLMs")
# # # # # # # # # #     parser.add_argument("--provider", required=True, choices=["OpenAI", "GoogleGemini"], help="LLM provider")
# # # # # # # # # #     parser.add_argument("--model", required=True, help="Model name (e.g., gpt-4o, gemini-1.5-flash)")
# # # # # # # # # #     parser.add_argument("--folder", default="data/Answer_Scripts", help="Folder with DOCX files")

# # # # # # # # # #     args = parser.parse_args()

# # # # # # # # # #     extractor = AnswerExtractor(selected_provider=args.provider, selected_model=args.model)

# # # # # # # # # #     for filename in os.listdir(args.folder):
# # # # # # # # # #         if filename.lower().endswith(".docx"):
# # # # # # # # # #             filepath = os.path.join(args.folder, filename)
# # # # # # # # # #             print(f"\n📄 Processing: {filename}")

# # # # # # # # # #             try:
# # # # # # # # # #                 raw_text = load_docx_text(filepath)
# # # # # # # # # #                 answers = extractor.extract_answers_with_llm(raw_text)

# # # # # # # # # #                 if not answers:
# # # # # # # # # #                     print("❌ No answers extracted.")
# # # # # # # # # #                     continue

# # # # # # # # # #                 # Print preview
# # # # # # # # # #                 pprint([
# # # # # # # # # #                     {"question": ans.full_question_id, "answer": ans.answer_text}
# # # # # # # # # #                     for ans in answers
# # # # # # # # # #                 ])

# # # # # # # # # #                 first = answers[0]
# # # # # # # # # #                 db = StudentAnswerService(provider_suffix=args.provider)
# # # # # # # # # #                 db.initialize_table()
# # # # # # # # # #                 db.save_answers(
# # # # # # # # # #                     student_index=first.student_index,
# # # # # # # # # #                     module_code=first.module_code,
# # # # # # # # # #                     year=first.exam_year,
# # # # # # # # # #                     month=first.exam_month,
# # # # # # # # # #                     answers=answers
# # # # # # # # # #                 )
# # # # # # # # # #                 db.close()
# # # # # # # # # #                 print(f"✅ Saved answers for {first.student_index} | {first.module_code} | {first.exam_year}-{first.exam_month}")
# # # # # # # # # #             except Exception as e:
# # # # # # # # # #                 print(f"❌ Failed to process {filename}: {e}")

# # # # # # # # # import sys
# # # # # # # # # import os
# # # # # # # # # import time
# # # # # # # # # from docx import Document
# # # # # # # # # from pprint import pprint

# # # # # # # # # sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "../../")))

# # # # # # # # # from src.services.answer_extractor import AnswerExtractor
# # # # # # # # # from src.services.database_services.student_answer_db import StudentAnswerService

# # # # # # # # # def load_docx_text(docx_path: str) -> str:
# # # # # # # # #     doc = Document(docx_path)
# # # # # # # # #     return "\n".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])

# # # # # # # # # def extract_and_save(docx_path: str, extractor: AnswerExtractor, provider: str):
# # # # # # # # #     filename = os.path.basename(docx_path)
# # # # # # # # #     print(f"\n📄 Processing: {filename}")

# # # # # # # # #     try:
# # # # # # # # #         raw_text = load_docx_text(docx_path)
# # # # # # # # #         answers = extractor.extract_answers_with_llm(raw_text)

# # # # # # # # #         if not answers:
# # # # # # # # #             print("❌ No answers extracted.")
# # # # # # # # #             return

# # # # # # # # #         # Preview the result
# # # # # # # # #         pprint([
# # # # # # # # #             {"question": ans.full_question_id, "answer": ans.answer_text}
# # # # # # # # #             for ans in answers
# # # # # # # # #         ])

# # # # # # # # #         # Save to database
# # # # # # # # #         first = answers[0]
# # # # # # # # #         db = StudentAnswerService(provider_suffix=provider)
# # # # # # # # #         db.initialize_table()
# # # # # # # # #         db.save_answers(
# # # # # # # # #             student_index=first.student_index,
# # # # # # # # #             module_code=first.module_code,
# # # # # # # # #             year=first.exam_year,
# # # # # # # # #             month=first.exam_month,
# # # # # # # # #             answers=answers
# # # # # # # # #         )
# # # # # # # # #         db.close()

# # # # # # # # #         print(f"✅ Saved answers for {first.student_index} | {first.module_code} | {first.exam_year}-{first.exam_month}")
# # # # # # # # #     except Exception as e:
# # # # # # # # #         print(f"❌ Failed to process {filename}: {e}")

# # # # # # # # # if __name__ == "__main__":
# # # # # # # # #     import argparse

# # # # # # # # #     parser = argparse.ArgumentParser(description="Extract and save student answers using LLMs")
# # # # # # # # #     parser.add_argument("--provider", required=True, choices=["OpenAI", "GoogleGemini"], help="LLM provider")
# # # # # # # # #     parser.add_argument("--model", required=True, help="Model name (e.g., gpt-4o, gemini-2.0-flash)")
# # # # # # # # #     parser.add_argument("--folder", required=True, help="Single DOCX file or folder containing DOCX files")

# # # # # # # # #     args = parser.parse_args()
# # # # # # # # #     extractor = AnswerExtractor(selected_provider=args.provider, selected_model=args.model)

# # # # # # # # #     if os.path.isfile(args.folder) and args.folder.endswith(".docx"):
# # # # # # # # #         # Single file mode
# # # # # # # # #         extract_and_save(args.folder, extractor, args.provider)
# # # # # # # # #     elif os.path.isdir(args.folder):
# # # # # # # # #         # Folder mode
# # # # # # # # #         for filename in os.listdir(args.folder):
# # # # # # # # #             if filename.lower().endswith(".docx"):
# # # # # # # # #                 filepath = os.path.join(args.folder, filename)
# # # # # # # # #                 extract_and_save(filepath, extractor, args.provider)

# # # # # # # # #                 # Delay to respect Gemini rate limits (15 requests/min)
# # # # # # # # #                 if args.provider == "GoogleGemini":
# # # # # # # # #                     time.sleep(10)
# # # # # # # # #     else:
# # # # # # # # #         print("❌ Invalid --folder path. Must be either a .docx file or a directory.")


# # # # # # # # # import sys
# # # # # # # # # import os
# # # # # # # # # import time
# # # # # # # # # import pathlib
# # # # # # # # # from docx import Document
# # # # # # # # # from pprint import pprint
# # # # # # # # # import psycopg2
# # # # # # # # # from psycopg2.extras import RealDictCursor
# # # # # # # # # from dotenv import load_dotenv
# # # # # # # # # import pdfplumber

# # # # # # # # # sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "../../")))

# # # # # # # # # from src.services.answer_extractor import AnswerExtractor
# # # # # # # # # from src.services.database_services.student_answer_db import StudentAnswerService

# # # # # # # # # load_dotenv()

# # # # # # # # # def get_database_connection():
# # # # # # # # #     """Get database connection using environment variables."""
# # # # # # # # #     try:
# # # # # # # # #         conn = psycopg2.connect(
# # # # # # # # #             host=os.getenv('POSTGRES_HOST'),
# # # # # # # # #             port=os.getenv('POSTGRES_PORT'),
# # # # # # # # #             database=os.getenv('POSTGRES_DB'),
# # # # # # # # #             user=os.getenv('POSTGRES_USER'),
# # # # # # # # #             password=os.getenv('POSTGRES_PASSWORD')
# # # # # # # # #         )
# # # # # # # # #         return conn
# # # # # # # # #     except Exception as e:
# # # # # # # # #         print(f"Failed to connect to database: {e}")
# # # # # # # # #         sys.exit(1)

# # # # # # # # # def get_submissions_from_db():
# # # # # # # # #     """Retrieve all submission file paths from database using new schema."""
# # # # # # # # #     conn = get_database_connection()
# # # # # # # # #     try:
# # # # # # # # #         with conn.cursor(cursor_factory=RealDictCursor) as cur:
# # # # # # # # #             cur.execute("""
# # # # # # # # #                 SELECT submission_id, assessment_id, student_id, file_url, submission_start_at
# # # # # # # # #                 FROM "Submission" 
# # # # # # # # #                 ORDER BY assessment_id, submission_start_at ASC
# # # # # # # # #             """)
            
# # # # # # # # #             submissions = cur.fetchall()
# # # # # # # # #             print(f"Found {len(submissions)} submissions in database")
# # # # # # # # #             return submissions
            
# # # # # # # # #     except Exception as e:
# # # # # # # # #         print(f"Database error: {e}")
# # # # # # # # #         return []
# # # # # # # # #     finally:
# # # # # # # # #         conn.close()

# # # # # # # # # def resolve_file_path(file_path: str) -> pathlib.Path:
# # # # # # # # #     """Resolve database file path to actual file location."""
# # # # # # # # #     project_root = pathlib.Path.cwd()
# # # # # # # # #     # Try parent directory first (where data folder should be)
# # # # # # # # #     full_path = project_root.parent / file_path
    
# # # # # # # # #     if not full_path.exists():
# # # # # # # # #         # Try alternative path in project directory
# # # # # # # # #         alternative_path = project_root / file_path
# # # # # # # # #         if alternative_path.exists():
# # # # # # # # #             full_path = alternative_path
# # # # # # # # #         else:
# # # # # # # # #             raise FileNotFoundError(f"File not found at: {full_path} or {alternative_path}")
    
# # # # # # # # #     return full_path

# # # # # # # # # def load_docx_text(docx_path: str) -> str:
# # # # # # # # #     doc = Document(docx_path)
# # # # # # # # #     return "\n".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])

# # # # # # # # # def load_pdf_text(pdf_path: str) -> str:
# # # # # # # # #     with pdfplumber.open(pdf_path) as pdf:
# # # # # # # # #         return "\n".join(page.extract_text() or "" for page in pdf.pages)

# # # # # # # # # def load_file_text(file_path: str) -> str:
# # # # # # # # #     """Load text from file based on extension."""
# # # # # # # # #     path_obj = pathlib.Path(file_path)
# # # # # # # # #     if path_obj.suffix.lower() == '.docx':
# # # # # # # # #         return load_docx_text(file_path)
# # # # # # # # #     elif path_obj.suffix.lower() == '.pdf':
# # # # # # # # #         return load_pdf_text(file_path)
# # # # # # # # #     else:
# # # # # # # # #         raise ValueError(f"Unsupported file format: {path_obj.suffix}")

# # # # # # # # # def extract_and_save(docx_path: str, extractor: AnswerExtractor, provider: str):
# # # # # # # # #     filename = os.path.basename(docx_path)
# # # # # # # # #     print(f"\n📄 Processing: {filename}")

# # # # # # # # #     try:
# # # # # # # # #         raw_text = load_docx_text(docx_path)
# # # # # # # # #         answers = extractor.extract_answers_with_llm(raw_text)

# # # # # # # # #         if not answers:
# # # # # # # # #             print("❌ No answers extracted.")
# # # # # # # # #             return

# # # # # # # # #         # Preview the result
# # # # # # # # #         pprint([
# # # # # # # # #             {"question": ans.full_question_id, "answer": ans.answer_text}
# # # # # # # # #             for ans in answers
# # # # # # # # #         ])

# # # # # # # # #         # Save to database
# # # # # # # # #         first = answers[0]
# # # # # # # # #         db = StudentAnswerService(provider_suffix=provider)
# # # # # # # # #         db.initialize_table()
# # # # # # # # #         db.save_answers(
# # # # # # # # #             student_index=first.student_index,
# # # # # # # # #             module_code=first.module_code,
# # # # # # # # #             year=first.exam_year,
# # # # # # # # #             month=first.exam_month,
# # # # # # # # #             answers=answers
# # # # # # # # #         )
# # # # # # # # #         db.close()

# # # # # # # # #         print(f"✅ Saved answers for {first.student_index} | {first.module_code} | {first.exam_year}-{first.exam_month}")
# # # # # # # # #     except Exception as e:
# # # # # # # # #         print(f"❌ Failed to process {filename}: {e}")

# # # # # # # # # if __name__ == "__main__":
# # # # # # # # #     import argparse

# # # # # # # # #     parser = argparse.ArgumentParser(description="Extract and save student answers using LLMs")
# # # # # # # # #     parser.add_argument("--provider", required=True, choices=["OpenAI", "GoogleGemini"], help="LLM provider")
# # # # # # # # #     parser.add_argument("--model", required=True, help="Model name (e.g., gpt-4o, gemini-2.0-flash)")
# # # # # # # # #     parser.add_argument("--folder", help="Single DOCX file or folder containing DOCX files")
# # # # # # # # #     parser.add_argument("--from-db", action="store_true", help="Process files from database instead of folder")

# # # # # # # # #     args = parser.parse_args()
# # # # # # # # #     extractor = AnswerExtractor(selected_provider=args.provider, selected_model=args.model)

# # # # # # # # #     if args.from_db:
# # # # # # # # #         # Database mode - process all files from database
# # # # # # # # #         submissions = get_submissions_from_db()
        
# # # # # # # # #         if not submissions:
# # # # # # # # #             print("No submissions found in database.")
# # # # # # # # #             sys.exit(1)
        
# # # # # # # # #         # Process all files from database
# # # # # # # # #         db = StudentAnswerService(provider_suffix=args.provider)
# # # # # # # # #         db.initialize_table()
        
# # # # # # # # #         for submission in submissions:
# # # # # # # # #             try:
# # # # # # # # #                 file_path = submission['file_url']
# # # # # # # # #                 assessment_id = submission['assessment_id']
# # # # # # # # #                 student_id = submission['student_id']
                
# # # # # # # # #                 print(f"\n📄 Processing: {pathlib.Path(file_path).name} (Assessment: {assessment_id}, Student: {student_id})")
                
# # # # # # # # #                 # Resolve file path and load text
# # # # # # # # #                 full_path = resolve_file_path(file_path)
# # # # # # # # #                 raw_text = load_file_text(str(full_path))
                
# # # # # # # # #                 # Extract answers using LLM
# # # # # # # # #                 answers = extractor.extract_answers_with_llm(raw_text)

# # # # # # # # #                 if not answers:
# # # # # # # # #                     print("❌ No answers extracted.")
# # # # # # # # #                     continue

# # # # # # # # #                 # Print extracted answers for verification
# # # # # # # # #                 print("\nExtracted Answers (before saving to DB):\n")
# # # # # # # # #                 pprint([
# # # # # # # # #                     {
# # # # # # # # #                         "question": ans.full_question_id,
# # # # # # # # #                         "answer": ans.answer_text
# # # # # # # # #                     }
# # # # # # # # #                     for ans in answers
# # # # # # # # #                 ])

# # # # # # # # #                 # Save answers to DB
# # # # # # # # #                 first = answers[0]
# # # # # # # # #                 db.save_answers(
# # # # # # # # #                     student_index=first.student_index,
# # # # # # # # #                     module_code=first.module_code,
# # # # # # # # #                     year=first.exam_year,
# # # # # # # # #                     month=first.exam_month,
# # # # # # # # #                     answers=answers
# # # # # # # # #                 )

# # # # # # # # #                 print(f"✅ Saved answers for {first.student_index} | {first.module_code} | {first.exam_year}-{first.exam_month}")
                
# # # # # # # # #                 # Delay to respect Gemini rate limits (15 requests/min)
# # # # # # # # #                 if args.provider == "GoogleGemini":
# # # # # # # # #                     time.sleep(10)
                
# # # # # # # # #             except Exception as e:
# # # # # # # # #                 print(f"❌ Error processing {pathlib.Path(submission['file_url']).name}: {e}")
# # # # # # # # #                 continue
        
# # # # # # # # #         db.close()
        
# # # # # # # # #     elif args.folder:
# # # # # # # # #         # Original folder/file mode
# # # # # # # # #         if os.path.isfile(args.folder) and args.folder.endswith(".docx"):
# # # # # # # # #             # Single file mode
# # # # # # # # #             extract_and_save(args.folder, extractor, args.provider)
# # # # # # # # #         elif os.path.isdir(args.folder):
# # # # # # # # #             # Folder mode
# # # # # # # # #             for filename in os.listdir(args.folder):
# # # # # # # # #                 if filename.lower().endswith(".docx"):
# # # # # # # # #                     filepath = os.path.join(args.folder, filename)
# # # # # # # # #                     extract_and_save(filepath, extractor, args.provider)

# # # # # # # # #                     # Delay to respect Gemini rate limits (15 requests/min)
# # # # # # # # #                     if args.provider == "GoogleGemini":
# # # # # # # # #                         time.sleep(10)
# # # # # # # # #         else:
# # # # # # # # #             print("❌ Invalid --folder path. Must be either a .docx file or a directory.")
# # # # # # # # #     else:
# # # # # # # # #         print("❌ Please specify either --folder for file/folder mode or --from-db for database mode.")

# # # # # # # # # import sys
# # # # # # # # # import os
# # # # # # # # # import time
# # # # # # # # # import pathlib
# # # # # # # # # from docx import Document
# # # # # # # # # from pprint import pprint
# # # # # # # # # import psycopg2
# # # # # # # # # from psycopg2.extras import RealDictCursor
# # # # # # # # # from dotenv import load_dotenv
# # # # # # # # # import pdfplumber

# # # # # # # # # sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "../../")))

# # # # # # # # # from src.services.answer_extractor import AnswerExtractor
# # # # # # # # # from src.services.database_services.student_answer_db import StudentAnswerService

# # # # # # # # # load_dotenv()

# # # # # # # # # def get_database_connection():
# # # # # # # # #     """Get database connection using environment variables."""
# # # # # # # # #     try:
# # # # # # # # #         conn = psycopg2.connect(
# # # # # # # # #             host=os.getenv('POSTGRES_HOST'),
# # # # # # # # #             port=os.getenv('POSTGRES_PORT'),
# # # # # # # # #             database=os.getenv('POSTGRES_DB'),
# # # # # # # # #             user=os.getenv('POSTGRES_USER'),
# # # # # # # # #             password=os.getenv('POSTGRES_PASSWORD')
# # # # # # # # #         )
# # # # # # # # #         return conn
# # # # # # # # #     except Exception as e:
# # # # # # # # #         print(f"Failed to connect to database: {e}")
# # # # # # # # #         sys.exit(1)

# # # # # # # # # def get_submissions_from_db():
# # # # # # # # #     """Retrieve all submission file paths from database using new schema."""
# # # # # # # # #     conn = get_database_connection()
# # # # # # # # #     try:
# # # # # # # # #         with conn.cursor(cursor_factory=RealDictCursor) as cur:
# # # # # # # # #             cur.execute("""
# # # # # # # # #                 SELECT submission_id, assessment_id, student_id, file_url, submission_start_at
# # # # # # # # #                 FROM "Submission" 
# # # # # # # # #                 ORDER BY assessment_id, submission_start_at ASC
# # # # # # # # #             """)
            
# # # # # # # # #             submissions = cur.fetchall()
# # # # # # # # #             print(f"Found {len(submissions)} submissions in database")
# # # # # # # # #             return submissions
            
# # # # # # # # #     except Exception as e:
# # # # # # # # #         print(f"Database error: {e}")
# # # # # # # # #         return []
# # # # # # # # #     finally:
# # # # # # # # #         conn.close()

# # # # # # # # # def resolve_file_path(file_path: str) -> pathlib.Path:
# # # # # # # # #     """Resolve database file path to actual file location."""
# # # # # # # # #     project_root = pathlib.Path.cwd()
# # # # # # # # #     # Try parent directory first (where data folder should be)
# # # # # # # # #     full_path = project_root.parent / file_path
    
# # # # # # # # #     if not full_path.exists():
# # # # # # # # #         # Try alternative path in project directory
# # # # # # # # #         alternative_path = project_root / file_path
# # # # # # # # #         if alternative_path.exists():
# # # # # # # # #             full_path = alternative_path
# # # # # # # # #         else:
# # # # # # # # #             raise FileNotFoundError(f"File not found at: {full_path} or {alternative_path}")
    
# # # # # # # # #     return full_path

# # # # # # # # # def load_docx_text(docx_path: str) -> str:
# # # # # # # # #     doc = Document(docx_path)
# # # # # # # # #     return "\n".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])

# # # # # # # # # def load_pdf_text(pdf_path: str) -> str:
# # # # # # # # #     with pdfplumber.open(pdf_path) as pdf:
# # # # # # # # #         return "\n".join(page.extract_text() or "" for page in pdf.pages)

# # # # # # # # # def load_file_text(file_path: str) -> str:
# # # # # # # # #     """Load text from file based on extension."""
# # # # # # # # #     path_obj = pathlib.Path(file_path)
# # # # # # # # #     if path_obj.suffix.lower() == '.docx':
# # # # # # # # #         return load_docx_text(file_path)
# # # # # # # # #     elif path_obj.suffix.lower() == '.pdf':
# # # # # # # # #         return load_pdf_text(file_path)
# # # # # # # # #     else:
# # # # # # # # #         raise ValueError(f"Unsupported file format: {path_obj.suffix}")

# # # # # # # # # def extract_and_save(docx_path: str, extractor: AnswerExtractor, provider: str):
# # # # # # # # #     filename = os.path.basename(docx_path)
# # # # # # # # #     print(f"\n📄 Processing: {filename}")

# # # # # # # # #     try:
# # # # # # # # #         raw_text = load_docx_text(docx_path)
# # # # # # # # #         answers = extractor.extract_answers_with_llm(raw_text)

# # # # # # # # #         if not answers:
# # # # # # # # #             print("❌ No answers extracted.")
# # # # # # # # #             return

# # # # # # # # #         # Preview the result
# # # # # # # # #         pprint([
# # # # # # # # #             {"question": ans.full_question_id, "answer": ans.answer_text}
# # # # # # # # #             for ans in answers
# # # # # # # # #         ])

# # # # # # # # #         # Save to database
# # # # # # # # #         first = answers[0]
# # # # # # # # #         db = StudentAnswerService(provider_suffix=provider)
# # # # # # # # #         db.initialize_table()
# # # # # # # # #         db.save_answers(
# # # # # # # # #             student_index=first.student_index,
# # # # # # # # #             module_code=first.module_code,
# # # # # # # # #             year=first.exam_year,
# # # # # # # # #             month=first.exam_month,
# # # # # # # # #             answers=answers
# # # # # # # # #         )
# # # # # # # # #         db.close()

# # # # # # # # #         print(f"✅ Saved answers for {first.student_index} | {first.module_code} | {first.exam_year}-{first.exam_month}")
# # # # # # # # #     except Exception as e:
# # # # # # # # #         print(f"❌ Failed to process {filename}: {e}")

# # # # # # # # # def main(provider="OpenAI", model="gpt-4o", folder=None, from_db=True):
# # # # # # # # #     """
# # # # # # # # #     Main function that can be called from Flask API.
    
# # # # # # # # #     Args:
# # # # # # # # #         provider: LLM provider ("OpenAI" or "GoogleGemini")
# # # # # # # # #         model: Model name (e.g., "gpt-4o", "gemini-2.0-flash")
# # # # # # # # #         folder: Single file or folder path (optional)
# # # # # # # # #         from_db: Whether to process files from database (default: True)
    
# # # # # # # # #     Returns:
# # # # # # # # #         dict: Result status and message
# # # # # # # # #     """
# # # # # # # # #     try:
# # # # # # # # #         extractor = AnswerExtractor(selected_provider=provider, selected_model=model)
        
# # # # # # # # #         if from_db:
# # # # # # # # #             # Database mode - process all files from database
# # # # # # # # #             submissions = get_submissions_from_db()
            
# # # # # # # # #             if not submissions:
# # # # # # # # #                 return {"status": "error", "message": "No submissions found in database."}
            
# # # # # # # # #             # Process all files from database
# # # # # # # # #             db = StudentAnswerService(provider_suffix=provider)
# # # # # # # # #             db.initialize_table()
            
# # # # # # # # #             processed_count = 0
# # # # # # # # #             error_count = 0
            
# # # # # # # # #             for submission in submissions:
# # # # # # # # #                 try:
# # # # # # # # #                     file_path = submission['file_url']
# # # # # # # # #                     assessment_id = submission['assessment_id']
# # # # # # # # #                     student_id = submission['student_id']
                    
# # # # # # # # #                     print(f"\n📄 Processing: {pathlib.Path(file_path).name} (Assessment: {assessment_id}, Student: {student_id})")
                    
# # # # # # # # #                     # Resolve file path and load text
# # # # # # # # #                     full_path = resolve_file_path(file_path)
# # # # # # # # #                     raw_text = load_file_text(str(full_path))
                    
# # # # # # # # #                     # Extract answers using LLM
# # # # # # # # #                     answers = extractor.extract_answers_with_llm(raw_text)

# # # # # # # # #                     if not answers:
# # # # # # # # #                         print("❌ No answers extracted.")
# # # # # # # # #                         error_count += 1
# # # # # # # # #                         continue

# # # # # # # # #                     # Print extracted answers for verification
# # # # # # # # #                     print("\nExtracted Answers (before saving to DB):\n")
# # # # # # # # #                     pprint([
# # # # # # # # #                         {
# # # # # # # # #                             "question": ans.full_question_id,
# # # # # # # # #                             "answer": ans.answer_text
# # # # # # # # #                         }
# # # # # # # # #                         for ans in answers
# # # # # # # # #                     ])

# # # # # # # # #                     # Save answers to DB
# # # # # # # # #                     first = answers[0]
# # # # # # # # #                     db.save_answers(
# # # # # # # # #                         student_index=first.student_index,
# # # # # # # # #                         module_code=first.module_code,
# # # # # # # # #                         year=first.exam_year,
# # # # # # # # #                         month=first.exam_month,
# # # # # # # # #                         answers=answers
# # # # # # # # #                     )

# # # # # # # # #                     print(f"✅ Saved answers for {first.student_index} | {first.module_code} | {first.exam_year}-{first.exam_month}")
# # # # # # # # #                     processed_count += 1
                    
# # # # # # # # #                     # Delay to respect Gemini rate limits (15 requests/min)
# # # # # # # # #                     if provider == "GoogleGemini":
# # # # # # # # #                         time.sleep(10)
                    
# # # # # # # # #                 except Exception as e:
# # # # # # # # #                     print(f"❌ Error processing {pathlib.Path(submission['file_url']).name}: {e}")
# # # # # # # # #                     error_count += 1
# # # # # # # # #                     continue
            
# # # # # # # # #             db.close()
            
# # # # # # # # #             return {
# # # # # # # # #                 "status": "success",
# # # # # # # # #                 "message": f"Processing complete. {processed_count} files processed successfully, {error_count} errors."
# # # # # # # # #             }
            
# # # # # # # # #         elif folder:
# # # # # # # # #             # Original folder/file mode
# # # # # # # # #             processed_count = 0
# # # # # # # # #             error_count = 0
            
# # # # # # # # #             if os.path.isfile(folder) and folder.endswith(".docx"):
# # # # # # # # #                 # Single file mode
# # # # # # # # #                 try:
# # # # # # # # #                     extract_and_save(folder, extractor, provider)
# # # # # # # # #                     processed_count = 1
# # # # # # # # #                 except Exception as e:
# # # # # # # # #                     error_count = 1
# # # # # # # # #                     print(f"❌ Error processing file: {e}")
                    
# # # # # # # # #             elif os.path.isdir(folder):
# # # # # # # # #                 # Folder mode
# # # # # # # # #                 for filename in os.listdir(folder):
# # # # # # # # #                     if filename.lower().endswith(".docx"):
# # # # # # # # #                         filepath = os.path.join(folder, filename)
# # # # # # # # #                         try:
# # # # # # # # #                             extract_and_save(filepath, extractor, provider)
# # # # # # # # #                             processed_count += 1
# # # # # # # # #                         except Exception as e:
# # # # # # # # #                             error_count += 1
# # # # # # # # #                             print(f"❌ Error processing {filename}: {e}")

# # # # # # # # #                         # Delay to respect Gemini rate limits (15 requests/min)
# # # # # # # # #                         if provider == "GoogleGemini":
# # # # # # # # #                             time.sleep(10)
# # # # # # # # #             else:
# # # # # # # # #                 return {"status": "error", "message": "Invalid folder path. Must be either a .docx file or a directory."}
            
# # # # # # # # #             return {
# # # # # # # # #                 "status": "success",
# # # # # # # # #                 "message": f"Processing complete. {processed_count} files processed successfully, {error_count} errors."
# # # # # # # # #             }
# # # # # # # # #         else:
# # # # # # # # #             return {"status": "error", "message": "Please specify either folder for file/folder mode or use from_db=True for database mode."}
            
# # # # # # # # #     except Exception as e:
# # # # # # # # #         return {"status": "error", "message": f"Unexpected error: {str(e)}"}

# # # # # # # # # if __name__ == "__main__":
# # # # # # # # #     import argparse

# # # # # # # # #     parser = argparse.ArgumentParser(description="Extract and save student answers using LLMs")
# # # # # # # # #     parser.add_argument("--provider", required=True, choices=["OpenAI", "GoogleGemini"], help="LLM provider")
# # # # # # # # #     parser.add_argument("--model", required=True, help="Model name (e.g., gpt-4o, gemini-2.0-flash)")
# # # # # # # # #     parser.add_argument("--folder", help="Single DOCX file or folder containing DOCX files")
# # # # # # # # #     parser.add_argument("--from-db", action="store_true", help="Process files from database instead of folder")

# # # # # # # # #     args = parser.parse_args()
    
# # # # # # # # #     # Call the main function with parsed arguments
# # # # # # # # #     result = main(
# # # # # # # # #         provider=args.provider,
# # # # # # # # #         model=args.model,
# # # # # # # # #         folder=args.folder,
# # # # # # # # #         from_db=args.from_db
# # # # # # # # #     )
    
# # # # # # # # #     print(f"\nFinal Result: {result}")
    
# # # # # # # # #     if result["status"] == "error":
# # # # # # # # #         sys.exit(1)

# # # # # # # # import sys
# # # # # # # # import os
# # # # # # # # import time
# # # # # # # # import pathlib
# # # # # # # # from docx import Document
# # # # # # # # from pprint import pprint
# # # # # # # # import psycopg2
# # # # # # # # from psycopg2.extras import RealDictCursor
# # # # # # # # from dotenv import load_dotenv
# # # # # # # # import pdfplumber

# # # # # # # # sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "../../")))

# # # # # # # # from src.services.answer_extractor import AnswerExtractor
# # # # # # # # from src.services.database_services.student_answer_db import StudentAnswerService

# # # # # # # # load_dotenv()

# # # # # # # # def get_provider_suffix(provider):
# # # # # # # #     """Convert provider name to database table suffix."""
# # # # # # # #     if provider == "OpenAI":
# # # # # # # #         return "openai"
# # # # # # # #     elif provider == "GoogleGemini":
# # # # # # # #         return "gemini"
# # # # # # # #     else:
# # # # # # # #         # Default fallback
# # # # # # # #         return provider.lower()

# # # # # # # # def get_database_connection():
# # # # # # # #     """Get database connection using environment variables."""
# # # # # # # #     try:
# # # # # # # #         conn = psycopg2.connect(
# # # # # # # #             host=os.getenv('POSTGRES_HOST'),
# # # # # # # #             port=os.getenv('POSTGRES_PORT'),
# # # # # # # #             database=os.getenv('POSTGRES_DB'),
# # # # # # # #             user=os.getenv('POSTGRES_USER'),
# # # # # # # #             password=os.getenv('POSTGRES_PASSWORD')
# # # # # # # #         )
# # # # # # # #         return conn
# # # # # # # #     except Exception as e:
# # # # # # # #         print(f"Failed to connect to database: {e}")
# # # # # # # #         sys.exit(1)

# # # # # # # # def get_submissions_from_db():
# # # # # # # #     """Retrieve all submission file paths from database using new schema."""
# # # # # # # #     conn = get_database_connection()
# # # # # # # #     try:
# # # # # # # #         with conn.cursor(cursor_factory=RealDictCursor) as cur:
# # # # # # # #             cur.execute("""
# # # # # # # #                 SELECT submission_id, assessment_id, student_id, file_url, submission_start_at
# # # # # # # #                 FROM "Submission" 
# # # # # # # #                 ORDER BY assessment_id, submission_start_at ASC
# # # # # # # #             """)
            
# # # # # # # #             submissions = cur.fetchall()
# # # # # # # #             print(f"Found {len(submissions)} submissions in database")
# # # # # # # #             return submissions
            
# # # # # # # #     except Exception as e:
# # # # # # # #         print(f"Database error: {e}")
# # # # # # # #         return []
# # # # # # # #     finally:
# # # # # # # #         conn.close()

# # # # # # # # def resolve_file_path(file_path: str) -> pathlib.Path:
# # # # # # # #     """Resolve database file path to actual file location."""
# # # # # # # #     project_root = pathlib.Path.cwd()
# # # # # # # #     # Try parent directory first (where data folder should be)
# # # # # # # #     full_path = project_root.parent / file_path
    
# # # # # # # #     if not full_path.exists():
# # # # # # # #         # Try alternative path in project directory
# # # # # # # #         alternative_path = project_root / file_path
# # # # # # # #         if alternative_path.exists():
# # # # # # # #             full_path = alternative_path
# # # # # # # #         else:
# # # # # # # #             raise FileNotFoundError(f"File not found at: {full_path} or {alternative_path}")
    
# # # # # # # #     return full_path

# # # # # # # # def load_docx_text(docx_path: str) -> str:
# # # # # # # #     doc = Document(docx_path)
# # # # # # # #     return "\n".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])

# # # # # # # # def load_pdf_text(pdf_path: str) -> str:
# # # # # # # #     with pdfplumber.open(pdf_path) as pdf:
# # # # # # # #         return "\n".join(page.extract_text() or "" for page in pdf.pages)

# # # # # # # # def load_file_text(file_path: str) -> str:
# # # # # # # #     """Load text from file based on extension."""
# # # # # # # #     path_obj = pathlib.Path(file_path)
# # # # # # # #     if path_obj.suffix.lower() == '.docx':
# # # # # # # #         return load_docx_text(file_path)
# # # # # # # #     elif path_obj.suffix.lower() == '.pdf':
# # # # # # # #         return load_pdf_text(file_path)
# # # # # # # #     else:
# # # # # # # #         raise ValueError(f"Unsupported file format: {path_obj.suffix}")

# # # # # # # # def extract_and_save(docx_path: str, extractor: AnswerExtractor, provider: str):
# # # # # # # #     filename = os.path.basename(docx_path)
# # # # # # # #     print(f"\n📄 Processing: {filename}")

# # # # # # # #     try:
# # # # # # # #         raw_text = load_docx_text(docx_path)
# # # # # # # #         answers = extractor.extract_answers_with_llm(raw_text)

# # # # # # # #         if not answers:
# # # # # # # #             print("❌ No answers extracted.")
# # # # # # # #             return

# # # # # # # #         # Preview the result
# # # # # # # #         pprint([
# # # # # # # #             {"question": ans.full_question_id, "answer": ans.answer_text}
# # # # # # # #             for ans in answers
# # # # # # # #         ])

# # # # # # # #         # Save to database - use proper suffix
# # # # # # # #         first = answers[0]
# # # # # # # #         provider_suffix = get_provider_suffix(provider)
# # # # # # # #         db = StudentAnswerService(provider_suffix=provider_suffix)
# # # # # # # #         db.initialize_table()
# # # # # # # #         db.save_answers(
# # # # # # # #             student_index=first.student_index,
# # # # # # # #             module_code=first.module_code,
# # # # # # # #             year=first.exam_year,
# # # # # # # #             month=first.exam_month,
# # # # # # # #             answers=answers
# # # # # # # #         )
# # # # # # # #         db.close()

# # # # # # # #         print(f"✅ Saved answers for {first.student_index} | {first.module_code} | {first.exam_year}-{first.exam_month}")
# # # # # # # #     except Exception as e:
# # # # # # # #         print(f"❌ Failed to process {filename}: {e}")

# # # # # # # # def main(provider="OpenAI", model="gpt-4o", folder=None, from_db=True):
# # # # # # # #     """
# # # # # # # #     Main function that can be called from Flask API.
    
# # # # # # # #     Args:
# # # # # # # #         provider: LLM provider ("OpenAI" or "GoogleGemini")
# # # # # # # #         model: Model name (e.g., "gpt-4o", "gemini-2.0-flash")
# # # # # # # #         folder: Single file or folder path (optional)
# # # # # # # #         from_db: Whether to process files from database (default: True)
    
# # # # # # # #     Returns:
# # # # # # # #         dict: Result status and message
# # # # # # # #     """
# # # # # # # #     try:
# # # # # # # #         extractor = AnswerExtractor(selected_provider=provider, selected_model=model)
        
# # # # # # # #         if from_db:
# # # # # # # #             # Database mode - process all files from database
# # # # # # # #             submissions = get_submissions_from_db()
            
# # # # # # # #             if not submissions:
# # # # # # # #                 return {"status": "error", "message": "No submissions found in database."}
            
# # # # # # # #             # Process all files from database - use proper suffix
# # # # # # # #             provider_suffix = get_provider_suffix(provider)
# # # # # # # #             print(f"⚙️ Using table suffix: {provider_suffix}")
# # # # # # # #             db = StudentAnswerService(provider_suffix=provider_suffix)
# # # # # # # #             db.initialize_table()
            
# # # # # # # #             processed_count = 0
# # # # # # # #             error_count = 0
            
# # # # # # # #             for submission in submissions:
# # # # # # # #                 try:
# # # # # # # #                     file_path = submission['file_url']
# # # # # # # #                     assessment_id = submission['assessment_id']
# # # # # # # #                     student_id = submission['student_id']
                    
# # # # # # # #                     print(f"\n📄 Processing: {pathlib.Path(file_path).name} (Assessment: {assessment_id}, Student: {student_id})")
                    
# # # # # # # #                     # Resolve file path and load text
# # # # # # # #                     full_path = resolve_file_path(file_path)
# # # # # # # #                     raw_text = load_file_text(str(full_path))
                    
# # # # # # # #                     # Extract answers using LLM
# # # # # # # #                     answers = extractor.extract_answers_with_llm(raw_text)

# # # # # # # #                     if not answers:
# # # # # # # #                         print("❌ No answers extracted.")
# # # # # # # #                         error_count += 1
# # # # # # # #                         continue

# # # # # # # #                     # Print extracted answers for verification
# # # # # # # #                     print("\nExtracted Answers (before saving to DB):\n")
# # # # # # # #                     pprint([
# # # # # # # #                         {
# # # # # # # #                             "question": ans.full_question_id,
# # # # # # # #                             "answer": ans.answer_text
# # # # # # # #                         }
# # # # # # # #                         for ans in answers
# # # # # # # #                     ])

# # # # # # # #                     # Save answers to DB
# # # # # # # #                     first = answers[0]
# # # # # # # #                     db.save_answers(
# # # # # # # #                         student_index=first.student_index,
# # # # # # # #                         module_code=first.module_code,
# # # # # # # #                         year=first.exam_year,
# # # # # # # #                         month=first.exam_month,
# # # # # # # #                         answers=answers
# # # # # # # #                     )

# # # # # # # #                     print(f"✅ Saved answers for {first.student_index} | {first.module_code} | {first.exam_year}-{first.exam_month}")
# # # # # # # #                     processed_count += 1
                    
# # # # # # # #                     # Delay to respect Gemini rate limits (15 requests/min)
# # # # # # # #                     if provider == "GoogleGemini":
# # # # # # # #                         time.sleep(10)
                    
# # # # # # # #                 except Exception as e:
# # # # # # # #                     print(f"❌ Error processing {pathlib.Path(submission['file_url']).name}: {e}")
# # # # # # # #                     error_count += 1
# # # # # # # #                     continue
            
# # # # # # # #             db.close()
            
# # # # # # # #             return {
# # # # # # # #                 "status": "success",
# # # # # # # #                 "message": f"Processing complete. {processed_count} files processed successfully, {error_count} errors."
# # # # # # # #             }
            
# # # # # # # #         elif folder:
# # # # # # # #             # Original folder/file mode
# # # # # # # #             processed_count = 0
# # # # # # # #             error_count = 0
            
# # # # # # # #             if os.path.isfile(folder) and folder.endswith(".docx"):
# # # # # # # #                 # Single file mode
# # # # # # # #                 try:
# # # # # # # #                     extract_and_save(folder, extractor, provider)
# # # # # # # #                     processed_count = 1
# # # # # # # #                 except Exception as e:
# # # # # # # #                     error_count = 1
# # # # # # # #                     print(f"❌ Error processing file: {e}")
                    
# # # # # # # #             elif os.path.isdir(folder):
# # # # # # # #                 # Folder mode
# # # # # # # #                 for filename in os.listdir(folder):
# # # # # # # #                     if filename.lower().endswith(".docx"):
# # # # # # # #                         filepath = os.path.join(folder, filename)
# # # # # # # #                         try:
# # # # # # # #                             extract_and_save(filepath, extractor, provider)
# # # # # # # #                             processed_count += 1
# # # # # # # #                         except Exception as e:
# # # # # # # #                             error_count += 1
# # # # # # # #                             print(f"❌ Error processing {filename}: {e}")

# # # # # # # #                         # Delay to respect Gemini rate limits (15 requests/min)
# # # # # # # #                         if provider == "GoogleGemini":
# # # # # # # #                             time.sleep(10)
# # # # # # # #             else:
# # # # # # # #                 return {"status": "error", "message": "Invalid folder path. Must be either a .docx file or a directory."}
            
# # # # # # # #             return {
# # # # # # # #                 "status": "success",
# # # # # # # #                 "message": f"Processing complete. {processed_count} files processed successfully, {error_count} errors."
# # # # # # # #             }
# # # # # # # #         else:
# # # # # # # #             return {"status": "error", "message": "Please specify either folder for file/folder mode or use from_db=True for database mode."}
            
# # # # # # # #     except Exception as e:
# # # # # # # #         return {"status": "error", "message": f"Unexpected error: {str(e)}"}

# # # # # # # # if __name__ == "__main__":
# # # # # # # #     import argparse

# # # # # # # #     parser = argparse.ArgumentParser(description="Extract and save student answers using LLMs")
# # # # # # # #     parser.add_argument("--provider", required=True, choices=["OpenAI", "GoogleGemini"], help="LLM provider")
# # # # # # # #     parser.add_argument("--model", required=True, help="Model name (e.g., gpt-4o, gemini-2.0-flash)")
# # # # # # # #     parser.add_argument("--folder", help="Single DOCX file or folder containing DOCX files")
# # # # # # # #     parser.add_argument("--from-db", action="store_true", help="Process files from database instead of folder")

# # # # # # # #     args = parser.parse_args()
    
# # # # # # # #     # Call the main function with parsed arguments
# # # # # # # #     result = main(
# # # # # # # #         provider=args.provider,
# # # # # # # #         model=args.model,
# # # # # # # #         folder=args.folder,
# # # # # # # #         from_db=args.from_db
# # # # # # # #     )
    
# # # # # # # #     print(f"\nFinal Result: {result}")
    
# # # # # # # #     print(f"⚙️ Used provider: {args.provider} -> table suffix: {get_provider_suffix(args.provider)}")
    
# # # # # # # #     if result["status"] == "error":
# # # # # # # #         sys.exit(1)

# # # # # # # import sys
# # # # # # # import os
# # # # # # # import time
# # # # # # # import pathlib
# # # # # # # from docx import Document
# # # # # # # from pprint import pprint
# # # # # # # import psycopg2
# # # # # # # from psycopg2.extras import RealDictCursor
# # # # # # # from dotenv import load_dotenv
# # # # # # # import pdfplumber

# # # # # # # sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "../../")))

# # # # # # # from src.services.answer_extractor import AnswerExtractor
# # # # # # # from src.services.database_services.student_answer_db import StudentAnswerService

# # # # # # # load_dotenv()

# # # # # # # def get_provider_suffix(provider):
# # # # # # #     """Convert provider name to database table suffix."""
# # # # # # #     if provider == "OpenAI":
# # # # # # #         return "openai"
# # # # # # #     elif provider == "GoogleGemini":
# # # # # # #         return "gemini"
# # # # # # #     else:
# # # # # # #         # Default fallback
# # # # # # #         return provider.lower()

# # # # # # # def get_database_connection():
# # # # # # #     """Get database connection using environment variables."""
# # # # # # #     try:
# # # # # # #         conn = psycopg2.connect(
# # # # # # #             host=os.getenv('POSTGRES_HOST'),
# # # # # # #             port=os.getenv('POSTGRES_PORT'),
# # # # # # #             database=os.getenv('POSTGRES_DB'),
# # # # # # #             user=os.getenv('POSTGRES_USER'),
# # # # # # #             password=os.getenv('POSTGRES_PASSWORD')
# # # # # # #         )
# # # # # # #         return conn
# # # # # # #     except Exception as e:
# # # # # # #         print(f"Failed to connect to database: {e}")
# # # # # # #         sys.exit(1)

# # # # # # # def get_submissions_from_db():
# # # # # # #     """Retrieve all submission file paths from database using new schema."""
# # # # # # #     conn = get_database_connection()
# # # # # # #     try:
# # # # # # #         with conn.cursor(cursor_factory=RealDictCursor) as cur:
# # # # # # #             cur.execute("""
# # # # # # #                 SELECT submission_id, assessment_id, student_id, file_url, submission_start_at
# # # # # # #                 FROM "Submission" 
# # # # # # #                 ORDER BY assessment_id, submission_start_at ASC
# # # # # # #             """)
            
# # # # # # #             submissions = cur.fetchall()
# # # # # # #             print(f"Found {len(submissions)} submissions in database")
# # # # # # #             return submissions
            
# # # # # # #     except Exception as e:
# # # # # # #         print(f"Database error: {e}")
# # # # # # #         return []
# # # # # # #     finally:
# # # # # # #         conn.close()

# # # # # # # def resolve_file_path(file_path: str) -> pathlib.Path:
# # # # # # #     """Resolve database file path to actual file location."""
# # # # # # #     project_root = pathlib.Path.cwd()
# # # # # # #     # Try parent directory first (where data folder should be)
# # # # # # #     full_path = project_root.parent / file_path
    
# # # # # # #     if not full_path.exists():
# # # # # # #         # Try alternative path in project directory
# # # # # # #         alternative_path = project_root / file_path
# # # # # # #         if alternative_path.exists():
# # # # # # #             full_path = alternative_path
# # # # # # #         else:
# # # # # # #             raise FileNotFoundError(f"File not found at: {full_path} or {alternative_path}")
    
# # # # # # #     return full_path

# # # # # # # def load_docx_text(docx_path: str) -> str:
# # # # # # #     doc = Document(docx_path)
# # # # # # #     return "\n".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])

# # # # # # # def load_pdf_text(pdf_path: str) -> str:
# # # # # # #     with pdfplumber.open(pdf_path) as pdf:
# # # # # # #         return "\n".join(page.extract_text() or "" for page in pdf.pages)

# # # # # # # def load_file_text(file_path: str) -> str:
# # # # # # #     """Load text from file based on extension."""
# # # # # # #     path_obj = pathlib.Path(file_path)
# # # # # # #     if path_obj.suffix.lower() == '.docx':
# # # # # # #         return load_docx_text(file_path)
# # # # # # #     elif path_obj.suffix.lower() == '.pdf':
# # # # # # #         return load_pdf_text(file_path)
# # # # # # #     else:
# # # # # # #         raise ValueError(f"Unsupported file format: {path_obj.suffix}")

# # # # # # # def extract_and_save(docx_path: str, extractor: AnswerExtractor, provider_suffix: str):
# # # # # # #     """Extract and save answers - now uses provider_suffix directly."""
# # # # # # #     filename = os.path.basename(docx_path)
# # # # # # #     print(f"\n📄 Processing: {filename}")

# # # # # # #     try:
# # # # # # #         raw_text = load_docx_text(docx_path)
# # # # # # #         answers = extractor.extract_answers_with_llm(raw_text)

# # # # # # #         if not answers:
# # # # # # #             print("❌ No answers extracted.")
# # # # # # #             return

# # # # # # #         # Preview the result
# # # # # # #         pprint([
# # # # # # #             {"question": ans.full_question_id, "answer": ans.answer_text}
# # # # # # #             for ans in answers
# # # # # # #         ])

# # # # # # #         # Save to database - use the passed provider_suffix
# # # # # # #         first = answers[0]
# # # # # # #         db = StudentAnswerService(provider_suffix=provider_suffix)
# # # # # # #         db.initialize_table()
# # # # # # #         db.save_answers(
# # # # # # #             student_index=first.student_index,
# # # # # # #             module_code=first.module_code,
# # # # # # #             year=first.exam_year,
# # # # # # #             month=first.exam_month,
# # # # # # #             answers=answers
# # # # # # #         )
# # # # # # #         db.close()

# # # # # # #         print(f"✅ Saved answers for {first.student_index} | {first.module_code} | {first.exam_year}-{first.exam_month}")
# # # # # # #     except Exception as e:
# # # # # # #         print(f"❌ Failed to process {filename}: {e}")

# # # # # # # def main(provider="OpenAI", model="gpt-4o", folder=None, from_db=True):
# # # # # # #     """
# # # # # # #     Main function that can be called from Flask API.
    
# # # # # # #     Args:
# # # # # # #         provider: LLM provider ("OpenAI" or "GoogleGemini")
# # # # # # #         model: Model name (e.g., "gpt-4o", "gemini-2.0-flash")
# # # # # # #         folder: Single file or folder path (optional)
# # # # # # #         from_db: Whether to process files from database (default: True)
    
# # # # # # #     Returns:
# # # # # # #         dict: Result status and message
# # # # # # #     """
# # # # # # #     try:
# # # # # # #         # Convert provider to proper suffix at the start
# # # # # # #         provider_suffix = get_provider_suffix(provider)
# # # # # # #         print(f"⚙️ Using provider: {provider} -> table suffix: {provider_suffix}")
        
# # # # # # #         extractor = AnswerExtractor(selected_provider=provider, selected_model=model)
        
# # # # # # #         if from_db:
# # # # # # #             # Database mode - process all files from database
# # # # # # #             submissions = get_submissions_from_db()
            
# # # # # # #             if not submissions:
# # # # # # #                 return {"status": "error", "message": "No submissions found in database."}
            
# # # # # # #             # Process all files from database - use converted suffix
# # # # # # #             db = StudentAnswerService(provider_suffix=provider_suffix)
# # # # # # #             db.initialize_table()
            
# # # # # # #             processed_count = 0
# # # # # # #             error_count = 0
            
# # # # # # #             for submission in submissions:
# # # # # # #                 try:
# # # # # # #                     file_path = submission['file_url']
# # # # # # #                     assessment_id = submission['assessment_id']
# # # # # # #                     student_id = submission['student_id']
                    
# # # # # # #                     print(f"\n📄 Processing: {pathlib.Path(file_path).name} (Assessment: {assessment_id}, Student: {student_id})")
                    
# # # # # # #                     # Resolve file path and load text
# # # # # # #                     full_path = resolve_file_path(file_path)
# # # # # # #                     raw_text = load_file_text(str(full_path))
                    
# # # # # # #                     # Extract answers using LLM
# # # # # # #                     answers = extractor.extract_answers_with_llm(raw_text)

# # # # # # #                     if not answers:
# # # # # # #                         print("❌ No answers extracted.")
# # # # # # #                         error_count += 1
# # # # # # #                         continue

# # # # # # #                     # Print extracted answers for verification
# # # # # # #                     print("\nExtracted Answers (before saving to DB):\n")
# # # # # # #                     pprint([
# # # # # # #                         {
# # # # # # #                             "question": ans.full_question_id,
# # # # # # #                             "answer": ans.answer_text
# # # # # # #                         }
# # # # # # #                         for ans in answers
# # # # # # #                     ])

# # # # # # #                     # Save answers to DB - reuse the same db connection
# # # # # # #                     first = answers[0]
# # # # # # #                     db.save_answers(
# # # # # # #                         student_index=first.student_index,
# # # # # # #                         module_code=first.module_code,
# # # # # # #                         year=first.exam_year,
# # # # # # #                         month=first.exam_month,
# # # # # # #                         answers=answers
# # # # # # #                     )

# # # # # # #                     print(f"✅ Saved answers for {first.student_index} | {first.module_code} | {first.exam_year}-{first.exam_month}")
# # # # # # #                     processed_count += 1
                    
# # # # # # #                     # Delay to respect Gemini rate limits (15 requests/min)
# # # # # # #                     if provider == "GoogleGemini":
# # # # # # #                         time.sleep(10)
                    
# # # # # # #                 except Exception as e:
# # # # # # #                     print(f"❌ Error processing {pathlib.Path(submission['file_url']).name}: {e}")
# # # # # # #                     error_count += 1
# # # # # # #                     continue
            
# # # # # # #             db.close()
            
# # # # # # #             return {
# # # # # # #                 "status": "success",
# # # # # # #                 "message": f"Processing complete. {processed_count} files processed successfully, {error_count} errors."
# # # # # # #             }
            
# # # # # # #         elif folder:
# # # # # # #             # Original folder/file mode
# # # # # # #             processed_count = 0
# # # # # # #             error_count = 0
            
# # # # # # #             if os.path.isfile(folder) and folder.endswith(".docx"):
# # # # # # #                 # Single file mode - pass provider_suffix instead of provider
# # # # # # #                 try:
# # # # # # #                     extract_and_save(folder, extractor, provider_suffix)
# # # # # # #                     processed_count = 1
# # # # # # #                 except Exception as e:
# # # # # # #                     error_count = 1
# # # # # # #                     print(f"❌ Error processing file: {e}")
                    
# # # # # # #             elif os.path.isdir(folder):
# # # # # # #                 # Folder mode - pass provider_suffix instead of provider
# # # # # # #                 for filename in os.listdir(folder):
# # # # # # #                     if filename.lower().endswith(".docx"):
# # # # # # #                         filepath = os.path.join(folder, filename)
# # # # # # #                         try:
# # # # # # #                             extract_and_save(filepath, extractor, provider_suffix)
# # # # # # #                             processed_count += 1
# # # # # # #                         except Exception as e:
# # # # # # #                             error_count += 1
# # # # # # #                             print(f"❌ Error processing {filename}: {e}")

# # # # # # #                         # Delay to respect Gemini rate limits (15 requests/min)
# # # # # # #                         if provider == "GoogleGemini":
# # # # # # #                             time.sleep(10)
# # # # # # #             else:
# # # # # # #                 return {"status": "error", "message": "Invalid folder path. Must be either a .docx file or a directory."}
            
# # # # # # #             return {
# # # # # # #                 "status": "success",
# # # # # # #                 "message": f"Processing complete. {processed_count} files processed successfully, {error_count} errors."
# # # # # # #             }
# # # # # # #         else:
# # # # # # #             return {"status": "error", "message": "Please specify either folder for file/folder mode or use from_db=True for database mode."}
            
# # # # # # #     except Exception as e:
# # # # # # #         return {"status": "error", "message": f"Unexpected error: {str(e)}"}

# # # # # # # if __name__ == "__main__":
# # # # # # #     import argparse

# # # # # # #     parser = argparse.ArgumentParser(description="Extract and save student answers using LLMs")
# # # # # # #     parser.add_argument("--provider", required=True, choices=["OpenAI", "GoogleGemini"], help="LLM provider")
# # # # # # #     parser.add_argument("--model", required=True, help="Model name (e.g., gpt-4o, gemini-2.0-flash)")
# # # # # # #     parser.add_argument("--folder", help="Single DOCX file or folder containing DOCX files")
# # # # # # #     parser.add_argument("--from-db", action="store_true", help="Process files from database instead of folder")

# # # # # # #     args = parser.parse_args()
    
# # # # # # #     # Call the main function with parsed arguments
# # # # # # #     result = main(
# # # # # # #         provider=args.provider,
# # # # # # #         model=args.model,
# # # # # # #         folder=args.folder,
# # # # # # #         from_db=args.from_db
# # # # # # #     )
    
# # # # # # #     print(f"\nFinal Result: {result}")
    
# # # # # # #     if result["status"] == "error":
# # # # # # #         sys.exit(1)


# # # # # # # import sys
# # # # # # # import os
# # # # # # # import time
# # # # # # # import pathlib
# # # # # # # from docx import Document
# # # # # # # from pprint import pprint
# # # # # # # import psycopg2
# # # # # # # from psycopg2.extras import RealDictCursor
# # # # # # # from dotenv import load_dotenv
# # # # # # # import pdfplumber

# # # # # # # sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "../../")))

# # # # # # # from src.services.answer_extractor import AnswerExtractor
# # # # # # # from src.services.database_services.student_answer_db import StudentAnswerService

# # # # # # # load_dotenv()

# # # # # # # def get_provider_suffix(provider):
# # # # # # #     """Convert provider name to database table suffix."""
# # # # # # #     print(f"🐛 DEBUG: Converting provider '{provider}' to suffix...")
# # # # # # #     if provider == "OpenAI":
# # # # # # #         suffix = "openai"
# # # # # # #     elif provider == "GoogleGemini":
# # # # # # #         suffix = "gemini"
# # # # # # #     else:
# # # # # # #         # Default fallback
# # # # # # #         suffix = provider.lower()
# # # # # # #     print(f"🐛 DEBUG: Provider '{provider}' -> Suffix '{suffix}'")
# # # # # # #     return suffix

# # # # # # # def get_database_connection():
# # # # # # #     """Get database connection using environment variables."""
# # # # # # #     try:
# # # # # # #         conn = psycopg2.connect(
# # # # # # #             host=os.getenv('POSTGRES_HOST'),
# # # # # # #             port=os.getenv('POSTGRES_PORT'),
# # # # # # #             database=os.getenv('POSTGRES_DB'),
# # # # # # #             user=os.getenv('POSTGRES_USER'),
# # # # # # #             password=os.getenv('POSTGRES_PASSWORD')
# # # # # # #         )
# # # # # # #         return conn
# # # # # # #     except Exception as e:
# # # # # # #         print(f"Failed to connect to database: {e}")
# # # # # # #         sys.exit(1)

# # # # # # # def get_submissions_from_db():
# # # # # # #     """Retrieve all submission file paths from database using new schema."""
# # # # # # #     conn = get_database_connection()
# # # # # # #     try:
# # # # # # #         with conn.cursor(cursor_factory=RealDictCursor) as cur:
# # # # # # #             cur.execute("""
# # # # # # #                 SELECT submission_id, assessment_id, student_id, file_url, submission_start_at
# # # # # # #                 FROM "Submission" 
# # # # # # #                 ORDER BY assessment_id, submission_start_at ASC
# # # # # # #             """)
            
# # # # # # #             submissions = cur.fetchall()
# # # # # # #             print(f"Found {len(submissions)} submissions in database")
# # # # # # #             return submissions
            
# # # # # # #     except Exception as e:
# # # # # # #         print(f"Database error: {e}")
# # # # # # #         return []
# # # # # # #     finally:
# # # # # # #         conn.close()

# # # # # # # def resolve_file_path(file_path: str) -> pathlib.Path:
# # # # # # #     """Resolve database file path to actual file location."""
# # # # # # #     project_root = pathlib.Path.cwd()
# # # # # # #     # Try parent directory first (where data folder should be)
# # # # # # #     full_path = project_root.parent / file_path
    
# # # # # # #     if not full_path.exists():
# # # # # # #         # Try alternative path in project directory
# # # # # # #         alternative_path = project_root / file_path
# # # # # # #         if alternative_path.exists():
# # # # # # #             full_path = alternative_path
# # # # # # #         else:
# # # # # # #             raise FileNotFoundError(f"File not found at: {full_path} or {alternative_path}")
    
# # # # # # #     return full_path

# # # # # # # def load_docx_text(docx_path: str) -> str:
# # # # # # #     doc = Document(docx_path)
# # # # # # #     return "\n".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])

# # # # # # # def load_pdf_text(pdf_path: str) -> str:
# # # # # # #     with pdfplumber.open(pdf_path) as pdf:
# # # # # # #         return "\n".join(page.extract_text() or "" for page in pdf.pages)

# # # # # # # def load_file_text(file_path: str) -> str:
# # # # # # #     """Load text from file based on extension."""
# # # # # # #     path_obj = pathlib.Path(file_path)
# # # # # # #     if path_obj.suffix.lower() == '.docx':
# # # # # # #         return load_docx_text(file_path)
# # # # # # #     elif path_obj.suffix.lower() == '.pdf':
# # # # # # #         return load_pdf_text(file_path)
# # # # # # #     else:
# # # # # # #         raise ValueError(f"Unsupported file format: {path_obj.suffix}")

# # # # # # # def extract_and_save(docx_path: str, extractor: AnswerExtractor, provider_suffix: str):
# # # # # # #     """Extract and save answers - now uses provider_suffix directly."""
# # # # # # #     filename = os.path.basename(docx_path)
# # # # # # #     print(f"\n📄 Processing: {filename}")
# # # # # # #     print(f"🐛 DEBUG: Using provider_suffix: {provider_suffix}")

# # # # # # #     try:
# # # # # # #         raw_text = load_docx_text(docx_path)
        
# # # # # # #         # DEBUG: Check extractor provider before using
# # # # # # #         if hasattr(extractor, 'current_provider'):
# # # # # # #             print(f"🐛 DEBUG: Extractor is using provider: {extractor.current_provider}")
        
# # # # # # #         answers = extractor.extract_answers_with_llm(raw_text)

# # # # # # #         if not answers:
# # # # # # #             print("❌ No answers extracted.")
# # # # # # #             return

# # # # # # #         # Preview the result
# # # # # # #         pprint([
# # # # # # #             {"question": ans.full_question_id, "answer": ans.answer_text}
# # # # # # #             for ans in answers
# # # # # # #         ])

# # # # # # #         # Save to database - use the passed provider_suffix
# # # # # # #         first = answers[0]
# # # # # # #         print(f"🐛 DEBUG: Creating StudentAnswerService with suffix: {provider_suffix}")
# # # # # # #         db = StudentAnswerService(provider_suffix=provider_suffix)
# # # # # # #         db.initialize_table()
# # # # # # #         db.save_answers(
# # # # # # #             student_index=first.student_index,
# # # # # # #             module_code=first.module_code,
# # # # # # #             year=first.exam_year,
# # # # # # #             month=first.exam_month,
# # # # # # #             answers=answers
# # # # # # #         )
# # # # # # #         db.close()

# # # # # # #         print(f"✅ Saved answers for {first.student_index} | {first.module_code} | {first.exam_year}-{first.exam_month}")
# # # # # # #     except Exception as e:
# # # # # # #         print(f"❌ Failed to process {filename}: {e}")
# # # # # # #         raise  # Re-raise to see full traceback

# # # # # # # def main(provider="OpenAI", model="gpt-4o", folder=None, from_db=True):
# # # # # # #     """
# # # # # # #     Main function that can be called from Flask API.
    
# # # # # # #     Args:
# # # # # # #         provider: LLM provider ("OpenAI" or "GoogleGemini")
# # # # # # #         model: Model name (e.g., "gpt-4o", "gemini-2.0-flash")
# # # # # # #         folder: Single file or folder path (optional)
# # # # # # #         from_db: Whether to process files from database (default: True)
    
# # # # # # #     Returns:
# # # # # # #         dict: Result status and message
# # # # # # #     """
# # # # # # #     try:
# # # # # # #         # DEBUG: Print received parameters
# # # # # # #         print(f"🐛 DEBUG: main() called with:")
# # # # # # #         print(f"   provider = '{provider}' (type: {type(provider)})")
# # # # # # #         print(f"   model = '{model}'")
# # # # # # #         print(f"   folder = {folder}")
# # # # # # #         print(f"   from_db = {from_db}")
        
# # # # # # #         # Convert provider to proper suffix at the start
# # # # # # #         provider_suffix = get_provider_suffix(provider)
# # # # # # #         print(f"⚙️ Using provider: {provider} -> table suffix: {provider_suffix}")
        
# # # # # # #         # DEBUG: Create extractor with explicit logging
# # # # # # #         print(f"🐛 DEBUG: Creating AnswerExtractor(selected_provider='{provider}', selected_model='{model}')")
# # # # # # #         extractor = AnswerExtractor(selected_provider=provider, selected_model=model)
        
# # # # # # #         # DEBUG: Inspect the created extractor
# # # # # # #         if hasattr(extractor, 'current_provider'):
# # # # # # #             print(f"🐛 DEBUG: AnswerExtractor.current_provider = {extractor.current_provider}")
# # # # # # #         if hasattr(extractor, 'provider'):
# # # # # # #             print(f"🐛 DEBUG: AnswerExtractor.provider = {extractor.provider}")
# # # # # # #         if hasattr(extractor, 'selected_provider'):
# # # # # # #             print(f"🐛 DEBUG: AnswerExtractor.selected_provider = {extractor.selected_provider}")
        
# # # # # # #         if from_db:
# # # # # # #             # Database mode - process all files from database
# # # # # # #             submissions = get_submissions_from_db()
            
# # # # # # #             if not submissions:
# # # # # # #                 return {"status": "error", "message": "No submissions found in database."}
            
# # # # # # #             # Process all files from database - use converted suffix
# # # # # # #             print(f"🐛 DEBUG: Creating StudentAnswerService with suffix: {provider_suffix}")
# # # # # # #             db = StudentAnswerService(provider_suffix=provider_suffix)
# # # # # # #             db.initialize_table()
            
# # # # # # #             processed_count = 0
# # # # # # #             error_count = 0
            
# # # # # # #             for submission in submissions:
# # # # # # #                 try:
# # # # # # #                     file_path = submission['file_url']
# # # # # # #                     assessment_id = submission['assessment_id']
# # # # # # #                     student_id = submission['student_id']
                    
# # # # # # #                     print(f"\n📄 Processing: {pathlib.Path(file_path).name} (Assessment: {assessment_id}, Student: {student_id})")
                    
# # # # # # #                     # Resolve file path and load text
# # # # # # #                     full_path = resolve_file_path(file_path)
# # # # # # #                     raw_text = load_file_text(str(full_path))
                    
# # # # # # #                     # DEBUG: Check extractor before using
# # # # # # #                     print(f"🐛 DEBUG: About to extract with provider suffix: {provider_suffix}")
                    
# # # # # # #                     # Extract answers using LLM
# # # # # # #                     answers = extractor.extract_answers_with_llm(raw_text)

# # # # # # #                     if not answers:
# # # # # # #                         print("❌ No answers extracted.")
# # # # # # #                         error_count += 1
# # # # # # #                         continue

# # # # # # #                     # Print extracted answers for verification
# # # # # # #                     print("\nExtracted Answers (before saving to DB):\n")
# # # # # # #                     pprint([
# # # # # # #                         {
# # # # # # #                             "question": ans.full_question_id,
# # # # # # #                             "answer": ans.answer_text
# # # # # # #                         }
# # # # # # #                         for ans in answers
# # # # # # #                     ])

# # # # # # #                     # Save answers to DB - reuse the same db connection
# # # # # # #                     first = answers[0]
# # # # # # #                     db.save_answers(
# # # # # # #                         student_index=first.student_index,
# # # # # # #                         module_code=first.module_code,
# # # # # # #                         year=first.exam_year,
# # # # # # #                         month=first.exam_month,
# # # # # # #                         answers=answers
# # # # # # #                     )

# # # # # # #                     print(f"✅ Saved answers for {first.student_index} | {first.module_code} | {first.exam_year}-{first.exam_month}")
# # # # # # #                     processed_count += 1
                    
# # # # # # #                     # Delay to respect Gemini rate limits (15 requests/min)
# # # # # # #                     if provider == "GoogleGemini":
# # # # # # #                         time.sleep(10)
                    
# # # # # # #                 except Exception as e:
# # # # # # #                     print(f"❌ Error processing {pathlib.Path(submission['file_url']).name}: {e}")
# # # # # # #                     error_count += 1
# # # # # # #                     continue
            
# # # # # # #             db.close()
            
# # # # # # #             return {
# # # # # # #                 "status": "success",
# # # # # # #                 "message": f"Processing complete. {processed_count} files processed successfully, {error_count} errors."
# # # # # # #             }
            
# # # # # # #         elif folder:
# # # # # # #             # Original folder/file mode
# # # # # # #             processed_count = 0
# # # # # # #             error_count = 0
            
# # # # # # #             if os.path.isfile(folder) and folder.endswith(".docx"):
# # # # # # #                 # Single file mode - pass provider_suffix instead of provider
# # # # # # #                 try:
# # # # # # #                     extract_and_save(folder, extractor, provider_suffix)
# # # # # # #                     processed_count = 1
# # # # # # #                 except Exception as e:
# # # # # # #                     error_count = 1
# # # # # # #                     print(f"❌ Error processing file: {e}")
                    
# # # # # # #             elif os.path.isdir(folder):
# # # # # # #                 # Folder mode - pass provider_suffix instead of provider
# # # # # # #                 for filename in os.listdir(folder):
# # # # # # #                     if filename.lower().endswith(".docx"):
# # # # # # #                         filepath = os.path.join(folder, filename)
# # # # # # #                         try:
# # # # # # #                             extract_and_save(filepath, extractor, provider_suffix)
# # # # # # #                             processed_count += 1
# # # # # # #                         except Exception as e:
# # # # # # #                             error_count += 1
# # # # # # #                             print(f"❌ Error processing {filename}: {e}")

# # # # # # #                         # Delay to respect Gemini rate limits (15 requests/min)
# # # # # # #                         if provider == "GoogleGemini":
# # # # # # #                             time.sleep(10)
# # # # # # #             else:
# # # # # # #                 return {"status": "error", "message": "Invalid folder path. Must be either a .docx file or a directory."}
            
# # # # # # #             return {
# # # # # # #                 "status": "success",
# # # # # # #                 "message": f"Processing complete. {processed_count} files processed successfully, {error_count} errors."
# # # # # # #             }
# # # # # # #         else:
# # # # # # #             return {"status": "error", "message": "Please specify either folder for file/folder mode or use from_db=True for database mode."}
            
# # # # # # #     except Exception as e:
# # # # # # #         print(f"🐛 DEBUG: Exception in main(): {e}")
# # # # # # #         import traceback
# # # # # # #         traceback.print_exc()
# # # # # # #         return {"status": "error", "message": f"Unexpected error: {str(e)}"}

# # # # # # # if __name__ == "__main__":
# # # # # # #     import argparse

# # # # # # #     parser = argparse.ArgumentParser(description="Extract and save student answers using LLMs")
# # # # # # #     parser.add_argument("--provider", required=True, choices=["OpenAI", "GoogleGemini"], help="LLM provider")
# # # # # # #     parser.add_argument("--model", required=True, help="Model name (e.g., gpt-4o, gemini-2.0-flash)")
# # # # # # #     parser.add_argument("--folder", help="Single DOCX file or folder containing DOCX files")
# # # # # # #     parser.add_argument("--from-db", action="store_true", help="Process files from database instead of folder")

# # # # # # #     args = parser.parse_args()
    
# # # # # # #     print(f"🐛 DEBUG: Command line args:")
# # # # # # #     print(f"   args.provider = '{args.provider}'")
# # # # # # #     print(f"   args.model = '{args.model}'")
# # # # # # #     print(f"   args.from_db = {args.from_db}")
    
# # # # # # #     # Call the main function with parsed arguments
# # # # # # #     result = main(
# # # # # # #         provider=args.provider,
# # # # # # #         model=args.model,
# # # # # # #         folder=args.folder,
# # # # # # #         from_db=args.from_db
# # # # # # #     )
    
# # # # # # #     print(f"\nFinal Result: {result}")
    
# # # # # # #     if result["status"] == "error":
# # # # # # #         sys.exit(1)

# # # # # # import sys
# # # # # # import os
# # # # # # import time
# # # # # # import pathlib
# # # # # # from docx import Document
# # # # # # from pprint import pprint
# # # # # # import psycopg2
# # # # # # from psycopg2.extras import RealDictCursor
# # # # # # from dotenv import load_dotenv
# # # # # # import pdfplumber

# # # # # # sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "../../")))

# # # # # # from src.services.answer_extractor import AnswerExtractor
# # # # # # from src.services.database_services.student_answer_db import StudentAnswerService

# # # # # # load_dotenv()

# # # # # # def get_provider_suffix(provider):
# # # # # #     """Convert provider name to database table suffix."""
# # # # # #     print(f"🐛 DEBUG: Converting provider '{provider}' to suffix...")
    
# # # # # #     # Standardize the mapping to ensure consistent table naming
# # # # # #     provider_map = {
# # # # # #         "OpenAI": "openai",
# # # # # #         "GoogleGemini": "gemini",
# # # # # #         "Gemini": "gemini",  # Alternative name
# # # # # #         "Google": "gemini"   # Another alternative
# # # # # #     }
    
# # # # # #     suffix = provider_map.get(provider, provider.lower())
# # # # # #     print(f"🐛 DEBUG: Provider '{provider}' -> Suffix '{suffix}'")
# # # # # #     return suffix

# # # # # # def get_database_connection():
# # # # # #     """Get database connection using environment variables."""
# # # # # #     try:
# # # # # #         conn = psycopg2.connect(
# # # # # #             host=os.getenv('POSTGRES_HOST'),
# # # # # #             port=os.getenv('POSTGRES_PORT'),
# # # # # #             database=os.getenv('POSTGRES_DB'),
# # # # # #             user=os.getenv('POSTGRES_USER'),
# # # # # #             password=os.getenv('POSTGRES_PASSWORD')
# # # # # #         )
# # # # # #         return conn
# # # # # #     except Exception as e:
# # # # # #         print(f"Failed to connect to database: {e}")
# # # # # #         sys.exit(1)

# # # # # # def get_submissions_from_db():
# # # # # #     """Retrieve all submission file paths from database using new schema."""
# # # # # #     conn = get_database_connection()
# # # # # #     try:
# # # # # #         with conn.cursor(cursor_factory=RealDictCursor) as cur:
# # # # # #             cur.execute("""
# # # # # #                 SELECT submission_id, assessment_id, student_id, file_url, submission_start_at
# # # # # #                 FROM "Submission" 
# # # # # #                 ORDER BY assessment_id, submission_start_at ASC
# # # # # #             """)
            
# # # # # #             submissions = cur.fetchall()
# # # # # #             print(f"Found {len(submissions)} submissions in database")
# # # # # #             return submissions
            
# # # # # #     except Exception as e:
# # # # # #         print(f"Database error: {e}")
# # # # # #         return []
# # # # # #     finally:
# # # # # #         conn.close()

# # # # # # def resolve_file_path(file_path: str) -> pathlib.Path:
# # # # # #     """Resolve database file path to actual file location."""
# # # # # #     project_root = pathlib.Path.cwd()
# # # # # #     # Try parent directory first (where data folder should be)
# # # # # #     full_path = project_root.parent / file_path
    
# # # # # #     if not full_path.exists():
# # # # # #         # Try alternative path in project directory
# # # # # #         alternative_path = project_root / file_path
# # # # # #         if alternative_path.exists():
# # # # # #             full_path = alternative_path
# # # # # #         else:
# # # # # #             raise FileNotFoundError(f"File not found at: {full_path} or {alternative_path}")
    
# # # # # #     return full_path

# # # # # # def load_docx_text(docx_path: str) -> str:
# # # # # #     doc = Document(docx_path)
# # # # # #     return "\n".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])

# # # # # # def load_pdf_text(pdf_path: str) -> str:
# # # # # #     with pdfplumber.open(pdf_path) as pdf:
# # # # # #         return "\n".join(page.extract_text() or "" for page in pdf.pages)

# # # # # # def load_file_text(file_path: str) -> str:
# # # # # #     """Load text from file based on extension."""
# # # # # #     path_obj = pathlib.Path(file_path)
# # # # # #     if path_obj.suffix.lower() == '.docx':
# # # # # #         return load_docx_text(file_path)
# # # # # #     elif path_obj.suffix.lower() == '.pdf':
# # # # # #         return load_pdf_text(file_path)
# # # # # #     else:
# # # # # #         raise ValueError(f"Unsupported file format: {path_obj.suffix}")

# # # # # # def extract_and_save(docx_path: str, extractor: AnswerExtractor, provider_suffix: str):
# # # # # #     """Extract and save answers - now uses provider_suffix directly."""
# # # # # #     filename = os.path.basename(docx_path)
# # # # # #     print(f"\n📄 Processing: {filename}")
# # # # # #     print(f"🐛 DEBUG: Using provider_suffix: {provider_suffix}")

# # # # # #     try:
# # # # # #         raw_text = load_docx_text(docx_path)
        
# # # # # #         # DEBUG: Check extractor provider before using
# # # # # #         if hasattr(extractor, 'current_provider'):
# # # # # #             print(f"🐛 DEBUG: Extractor is using provider: {extractor.current_provider}")
        
# # # # # #         answers = extractor.extract_answers_with_llm(raw_text)

# # # # # #         if not answers:
# # # # # #             print("❌ No answers extracted.")
# # # # # #             return

# # # # # #         # Preview the result
# # # # # #         pprint([
# # # # # #             {"question": ans.full_question_id, "answer": ans.answer_text}
# # # # # #             for ans in answers
# # # # # #         ])

# # # # # #         # Save to database - use the passed provider_suffix
# # # # # #         first = answers[0]
# # # # # #         print(f"🐛 DEBUG: Creating StudentAnswerService with suffix: {provider_suffix}")
# # # # # #         db = StudentAnswerService(provider_suffix=provider_suffix)
# # # # # #         db.initialize_table()
# # # # # #         db.save_answers(
# # # # # #             student_index=first.student_index,
# # # # # #             module_code=first.module_code,
# # # # # #             year=first.exam_year,
# # # # # #             month=first.exam_month,
# # # # # #             answers=answers
# # # # # #         )
# # # # # #         db.close()

# # # # # #         print(f"✅ Saved answers for {first.student_index} | {first.module_code} | {first.exam_year}-{first.exam_month}")
# # # # # #     except Exception as e:
# # # # # #         print(f"❌ Failed to process {filename}: {e}")
# # # # # #         raise  # Re-raise to see full traceback

# # # # # # def main(provider="OpenAI", model="gpt-4o", folder=None, from_db=True):
# # # # # #     """
# # # # # #     Main function that can be called from Flask API.
    
# # # # # #     Args:
# # # # # #         provider: LLM provider ("OpenAI" or "GoogleGemini")
# # # # # #         model: Model name (e.g., "gpt-4o", "gemini-2.0-flash")
# # # # # #         folder: Single file or folder path (optional)
# # # # # #         from_db: Whether to process files from database (default: True)
    
# # # # # #     Returns:
# # # # # #         dict: Result status and message
# # # # # #     """
# # # # # #     try:
# # # # # #         # DEBUG: Print received parameters
# # # # # #         print(f"🐛 DEBUG: main() called with:")
# # # # # #         print(f"   provider = '{provider}' (type: {type(provider)})")
# # # # # #         print(f"   model = '{model}'")
# # # # # #         print(f"   folder = {folder}")
# # # # # #         print(f"   from_db = {from_db}")
        
# # # # # #         # Convert provider to proper suffix at the start and validate
# # # # # #         provider_suffix = get_provider_suffix(provider)
# # # # # #         print(f"⚙️ Using provider: {provider} -> table suffix: {provider_suffix}")
        
# # # # # #         # Validate the provider suffix
# # # # # #         if provider_suffix not in ["openai", "gemini"]:
# # # # # #             raise ValueError(f"Unsupported provider suffix: {provider_suffix}. Expected 'openai' or 'gemini'")
        
# # # # # #         # DEBUG: Create extractor with explicit logging
# # # # # #         print(f"🐛 DEBUG: Creating AnswerExtractor(selected_provider='{provider}', selected_model='{model}')")
# # # # # #         extractor = AnswerExtractor(selected_provider=provider, selected_model=model)
        
# # # # # #         # DEBUG: Inspect the created extractor
# # # # # #         if hasattr(extractor, 'current_provider'):
# # # # # #             print(f"🐛 DEBUG: AnswerExtractor.current_provider = {extractor.current_provider}")
# # # # # #         if hasattr(extractor, 'provider'):
# # # # # #             print(f"🐛 DEBUG: AnswerExtractor.provider = {extractor.provider}")
# # # # # #         if hasattr(extractor, 'selected_provider'):
# # # # # #             print(f"🐛 DEBUG: AnswerExtractor.selected_provider = {extractor.selected_provider}")
        
# # # # # #         # Create database service once with the correct suffix
# # # # # #         print(f"🐛 DEBUG: Creating StudentAnswerService with suffix: '{provider_suffix}'")
# # # # # #         db_service = StudentAnswerService(provider_suffix=provider_suffix)
        
# # # # # #         # Initialize the correct table (this should create student_answers_gemini for Gemini)
# # # # # #         print(f"🐛 DEBUG: Initializing table for suffix: '{provider_suffix}'")
# # # # # #         db_service.initialize_table()
        
# # # # # #         if from_db:
# # # # # #             # Database mode - process all files from database
# # # # # #             submissions = get_submissions_from_db()
            
# # # # # #             if not submissions:
# # # # # #                 db_service.close()
# # # # # #                 return {"status": "error", "message": "No submissions found in database."}
            
# # # # # #             processed_count = 0
# # # # # #             error_count = 0
            
# # # # # #             for submission in submissions:
# # # # # #                 try:
# # # # # #                     file_path = submission['file_url']
# # # # # #                     assessment_id = submission['assessment_id']
# # # # # #                     student_id = submission['student_id']
                    
# # # # # #                     print(f"\n📄 Processing: {pathlib.Path(file_path).name} (Assessment: {assessment_id}, Student: {student_id})")
                    
# # # # # #                     # Resolve file path and load text
# # # # # #                     full_path = resolve_file_path(file_path)
# # # # # #                     raw_text = load_file_text(str(full_path))
                    
# # # # # #                     # DEBUG: Confirm we're using the right provider
# # # # # #                     print(f"🐛 DEBUG: About to extract with provider: {provider} (suffix: {provider_suffix})")
                    
# # # # # #                     # Extract answers using LLM
# # # # # #                     answers = extractor.extract_answers_with_llm(raw_text)

# # # # # #                     if not answers:
# # # # # #                         print("❌ No answers extracted.")
# # # # # #                         error_count += 1
# # # # # #                         continue

# # # # # #                     # Print extracted answers for verification
# # # # # #                     print("\nExtracted Answers (before saving to DB):\n")
# # # # # #                     pprint([
# # # # # #                         {
# # # # # #                             "question": ans.full_question_id,
# # # # # #                             "answer": ans.answer_text
# # # # # #                         }
# # # # # #                         for ans in answers
# # # # # #                     ])

# # # # # #                     # Save answers to DB using the shared db_service
# # # # # #                     first = answers[0]
                    
# # # # # #                     # DEBUG: Confirm table being used
# # # # # #                     table_name = f"student_answers_{provider_suffix}"
# # # # # #                     print(f"🐛 DEBUG: Saving to table: {table_name}")
                    
# # # # # #                     db_service.save_answers(
# # # # # #                         student_index=first.student_index,
# # # # # #                         module_code=first.module_code,
# # # # # #                         year=first.exam_year,
# # # # # #                         month=first.exam_month,
# # # # # #                         answers=answers
# # # # # #                     )

# # # # # #                     print(f"✅ Saved answers for {first.student_index} | {first.module_code} | {first.exam_year}-{first.exam_month} to {table_name}")
# # # # # #                     processed_count += 1
                    
# # # # # #                     # Delay to respect Gemini rate limits (15 requests/min)
# # # # # #                     if provider == "GoogleGemini":
# # # # # #                         print("⏱️ Waiting 10 seconds for Gemini rate limit...")
# # # # # #                         time.sleep(10)
                    
# # # # # #                 except Exception as e:
# # # # # #                     print(f"❌ Error processing {pathlib.Path(submission['file_url']).name}: {e}")
# # # # # #                     error_count += 1
# # # # # #                     continue
            
# # # # # #             db_service.close()
            
# # # # # #             return {
# # # # # #                 "status": "success",
# # # # # #                 "message": f"Processing complete. {processed_count} files processed successfully, {error_count} errors. Data saved to table: student_answers_{provider_suffix}",
# # # # # #                 "details": {
# # # # # #                     "processed": processed_count,
# # # # # #                     "errors": error_count,
# # # # # #                     "table": f"student_answers_{provider_suffix}"
# # # # # #                 }
# # # # # #             }
            
# # # # # #         elif folder:
# # # # # #             # Original folder/file mode
# # # # # #             processed_count = 0
# # # # # #             error_count = 0
            
# # # # # #             if os.path.isfile(folder) and folder.endswith(".docx"):
# # # # # #                 # Single file mode
# # # # # #                 try:
# # # # # #                     extract_and_save(folder, extractor, provider_suffix)
# # # # # #                     processed_count = 1
# # # # # #                 except Exception as e:
# # # # # #                     error_count = 1
# # # # # #                     print(f"❌ Error processing file: {e}")
                    
# # # # # #             elif os.path.isdir(folder):
# # # # # #                 # Folder mode
# # # # # #                 for filename in os.listdir(folder):
# # # # # #                     if filename.lower().endswith(".docx"):
# # # # # #                         filepath = os.path.join(folder, filename)
# # # # # #                         try:
# # # # # #                             extract_and_save(filepath, extractor, provider_suffix)
# # # # # #                             processed_count += 1
# # # # # #                         except Exception as e:
# # # # # #                             error_count += 1
# # # # # #                             print(f"❌ Error processing {filename}: {e}")

# # # # # #                         # Delay to respect Gemini rate limits (15 requests/min)
# # # # # #                         if provider == "GoogleGemini":
# # # # # #                             print("⏱️ Waiting 10 seconds for Gemini rate limit...")
# # # # # #                             time.sleep(10)
# # # # # #             else:
# # # # # #                 db_service.close()
# # # # # #                 return {"status": "error", "message": "Invalid folder path. Must be either a .docx file or a directory."}
            
# # # # # #             db_service.close()
            
# # # # # #             return {
# # # # # #                 "status": "success",
# # # # # #                 "message": f"Processing complete. {processed_count} files processed successfully, {error_count} errors. Data saved to table: student_answers_{provider_suffix}",
# # # # # #                 "details": {
# # # # # #                     "processed": processed_count,
# # # # # #                     "errors": error_count,
# # # # # #                     "table": f"student_answers_{provider_suffix}"
# # # # # #                 }
# # # # # #             }
# # # # # #         else:
# # # # # #             db_service.close()
# # # # # #             return {"status": "error", "message": "Please specify either folder for file/folder mode or use from_db=True for database mode."}
            
# # # # # #     except Exception as e:
# # # # # #         print(f"🐛 DEBUG: Exception in main(): {e}")
# # # # # #         import traceback
# # # # # #         traceback.print_exc()
# # # # # #         return {"status": "error", "message": f"Unexpected error: {str(e)}"}

# # # # # # if __name__ == "__main__":
# # # # # #     import argparse

# # # # # #     parser = argparse.ArgumentParser(description="Extract and save student answers using LLMs")
# # # # # #     parser.add_argument("--provider", required=True, choices=["OpenAI", "GoogleGemini"], help="LLM provider")
# # # # # #     parser.add_argument("--model", required=True, help="Model name (e.g., gpt-4o, gemini-2.0-flash)")
# # # # # #     parser.add_argument("--folder", help="Single DOCX file or folder containing DOCX files")
# # # # # #     parser.add_argument("--from-db", action="store_true", help="Process files from database instead of folder")

# # # # # #     args = parser.parse_args()
    
# # # # # #     print(f"🐛 DEBUG: Command line args:")
# # # # # #     print(f"   args.provider = '{args.provider}'")
# # # # # #     print(f"   args.model = '{args.model}'")
# # # # # #     print(f"   args.from_db = {args.from_db}")
    
# # # # # #     # Call the main function with parsed arguments
# # # # # #     result = main(
# # # # # #         provider=args.provider,
# # # # # #         model=args.model,
# # # # # #         folder=args.folder,
# # # # # #         from_db=args.from_db
# # # # # #     )
    
# # # # # #     print(f"\nFinal Result: {result}")
    
# # # # # #     if result["status"] == "error":
# # # # # #         sys.exit(1)


# # # # # import sys
# # # # # import os
# # # # # import time
# # # # # import pathlib
# # # # # from docx import Document
# # # # # from pprint import pprint
# # # # # import psycopg2
# # # # # from psycopg2.extras import RealDictCursor
# # # # # from dotenv import load_dotenv
# # # # # import pdfplumber
# # # # # import argparse

# # # # # sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "../../")))

# # # # # from src.services.answer_extractor import AnswerExtractor
# # # # # from src.services.database_services.student_answer_db import StudentAnswerService

# # # # # load_dotenv()

# # # # # def get_provider_suffix(provider):
# # # # #     """Convert provider name to database table suffix."""
# # # # #     print(f"🐛 DEBUG: Converting provider '{provider}' to suffix...")
    
# # # # #     # Standardize the mapping to ensure consistent table naming
# # # # #     provider_map = {
# # # # #         "OpenAI": "openai",
# # # # #         "GoogleGemini": "gemini",
# # # # #         "Gemini": "gemini",  # Alternative name
# # # # #         "Google": "gemini"   # Another alternative
# # # # #     }
    
# # # # #     suffix = provider_map.get(provider, provider.lower())
# # # # #     print(f"🐛 DEBUG: Provider '{provider}' -> Suffix '{suffix}'")
# # # # #     return suffix

# # # # # def get_database_connection():
# # # # #     """Get database connection using environment variables."""
# # # # #     try:
# # # # #         conn = psycopg2.connect(
# # # # #             host=os.getenv('POSTGRES_HOST'),
# # # # #             port=os.getenv('POSTGRES_PORT'),
# # # # #             database=os.getenv('POSTGRES_DB'),
# # # # #             user=os.getenv('POSTGRES_USER'),
# # # # #             password=os.getenv('POSTGRES_PASSWORD')
# # # # #         )
# # # # #         return conn
# # # # #     except Exception as e:
# # # # #         print(f"Failed to connect to database: {e}")
# # # # #         sys.exit(1)

# # # # # def get_submissions_from_db():
# # # # #     """Retrieve all submission file paths from database using new schema."""
# # # # #     conn = get_database_connection()
# # # # #     try:
# # # # #         with conn.cursor(cursor_factory=RealDictCursor) as cur:
# # # # #             cur.execute("""
# # # # #                 SELECT submission_id, assessment_id, student_id, file_url, submission_start_at
# # # # #                 FROM "Submission" 
# # # # #                 ORDER BY assessment_id, submission_start_at ASC
# # # # #             """)
            
# # # # #             submissions = cur.fetchall()
# # # # #             print(f"Found {len(submissions)} submissions in database")
# # # # #             return submissions
            
# # # # #     except Exception as e:
# # # # #         print(f"Database error: {e}")
# # # # #         return []
# # # # #     finally:
# # # # #         conn.close()

# # # # # def resolve_file_path(file_path: str) -> pathlib.Path:
# # # # #     """Resolve database file path to actual file location."""
# # # # #     project_root = pathlib.Path.cwd()
# # # # #     # Try parent directory first (where data folder should be)
# # # # #     full_path = project_root.parent / file_path
    
# # # # #     if not full_path.exists():
# # # # #         # Try alternative path in project directory
# # # # #         alternative_path = project_root / file_path
# # # # #         if alternative_path.exists():
# # # # #             full_path = alternative_path
# # # # #         else:
# # # # #             raise FileNotFoundError(f"File not found at: {full_path} or {alternative_path}")
    
# # # # #     return full_path

# # # # # def load_docx_text(docx_path: str) -> str:
# # # # #     doc = Document(docx_path)
# # # # #     return "\n".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])

# # # # # def load_pdf_text(pdf_path: str) -> str:
# # # # #     with pdfplumber.open(pdf_path) as pdf:
# # # # #         return "\n".join(page.extract_text() or "" for page in pdf.pages)

# # # # # def load_file_text(file_path: str) -> str:
# # # # #     """Load text from file based on extension."""
# # # # #     path_obj = pathlib.Path(file_path)
# # # # #     if path_obj.suffix.lower() == '.docx':
# # # # #         return load_docx_text(file_path)
# # # # #     elif path_obj.suffix.lower() == '.pdf':
# # # # #         return load_pdf_text(file_path)
# # # # #     else:
# # # # #         raise ValueError(f"Unsupported file format: {path_obj.suffix}")

# # # # # def extract_and_save(docx_path: str, extractor: AnswerExtractor, provider_suffix: str):
# # # # #     """Extract and save answers - now uses provider_suffix directly."""
# # # # #     filename = os.path.basename(docx_path)
# # # # #     print(f"\n📄 Processing: {filename}")
# # # # #     print(f"🐛 DEBUG: Using provider_suffix: {provider_suffix}")

# # # # #     try:
# # # # #         raw_text = load_docx_text(docx_path)
        
# # # # #         # DEBUG: Check extractor provider before using
# # # # #         if hasattr(extractor, 'current_provider'):
# # # # #             print(f"🐛 DEBUG: Extractor is using provider: {extractor.current_provider}")
        
# # # # #         answers = extractor.extract_answers_with_llm(raw_text)

# # # # #         if not answers:
# # # # #             print("❌ No answers extracted.")
# # # # #             return

# # # # #         # Preview the result
# # # # #         pprint([
# # # # #             {"question": ans.full_question_id, "answer": ans.answer_text}
# # # # #             for ans in answers
# # # # #         ])

# # # # #         # Save to database - use the passed provider_suffix
# # # # #         first = answers[0]
# # # # #         print(f"🐛 DEBUG: Creating StudentAnswerService with suffix: {provider_suffix}")
# # # # #         db = StudentAnswerService(provider_suffix=provider_suffix)
# # # # #         db.initialize_table()
# # # # #         db.save_answers(
# # # # #             student_index=first.student_index,
# # # # #             module_code=first.module_code,
# # # # #             year=first.exam_year,
# # # # #             month=first.exam_month,
# # # # #             answers=answers
# # # # #         )
# # # # #         db.close()

# # # # #         print(f"✅ Saved answers for {first.student_index} | {first.module_code} | {first.exam_year}-{first.exam_month}")
# # # # #     except Exception as e:
# # # # #         print(f"❌ Failed to process {filename}: {e}")
# # # # #         raise  # Re-raise to see full traceback

# # # # # def main(provider="OpenAI", model="gpt-4o", folder=None, from_db=True):
# # # # #     """
# # # # #     Main function that can be called from Flask API.
    
# # # # #     Args:
# # # # #         provider: LLM provider ("OpenAI" or "GoogleGemini")
# # # # #         model: Model name (e.g., "gpt-4o", "gemini-2.0-flash")
# # # # #         folder: Single file or folder path (optional)
# # # # #         from_db: Whether to process files from database (default: True)
    
# # # # #     Returns:
# # # # #         dict: Result status and message
# # # # #     """
# # # # #     try:
# # # # #         # DEBUG: Print received parameters
# # # # #         print(f"🐛 DEBUG: main() called with:")
# # # # #         print(f"   provider = '{provider}' (type: {type(provider)})")
# # # # #         print(f"   model = '{model}'")
# # # # #         print(f"   folder = {folder}")
# # # # #         print(f"   from_db = {from_db}")
        
# # # # #         # Convert provider to proper suffix at the start and validate
# # # # #         provider_suffix = get_provider_suffix(provider)
# # # # #         print(f"⚙️ Using provider: {provider} -> table suffix: {provider_suffix}")
        
# # # # #         # Validate the provider suffix
# # # # #         if provider_suffix not in ["openai", "gemini"]:
# # # # #             raise ValueError(f"Unsupported provider suffix: {provider_suffix}. Expected 'openai' or 'gemini'")
        
# # # # #         # DEBUG: Create extractor with explicit logging
# # # # #         print(f"🐛 DEBUG: Creating AnswerExtractor(selected_provider='{provider}', selected_model='{model}')")
# # # # #         extractor = AnswerExtractor(selected_provider=provider, selected_model=model)
        
# # # # #         # DEBUG: Inspect the created extractor
# # # # #         if hasattr(extractor, 'current_provider'):
# # # # #             print(f"🐛 DEBUG: AnswerExtractor.current_provider = {extractor.current_provider}")
# # # # #         if hasattr(extractor, 'provider'):
# # # # #             print(f"🐛 DEBUG: AnswerExtractor.provider = {extractor.provider}")
# # # # #         if hasattr(extractor, 'selected_provider'):
# # # # #             print(f"🐛 DEBUG: AnswerExtractor.selected_provider = {extractor.selected_provider}")
        
# # # # #         # Create database service once with the correct suffix
# # # # #         print(f"🐛 DEBUG: Creating StudentAnswerService with suffix: '{provider_suffix}'")
# # # # #         db_service = StudentAnswerService(provider_suffix=provider_suffix)
        
# # # # #         # Initialize the correct table (this should create student_answers_gemini for Gemini)
# # # # #         print(f"🐛 DEBUG: Initializing table for suffix: '{provider_suffix}'")
# # # # #         db_service.initialize_table()
        
# # # # #         if from_db:
# # # # #             # Database mode - process all files from database
# # # # #             submissions = get_submissions_from_db()
            
# # # # #             if not submissions:
# # # # #                 db_service.close()
# # # # #                 return {"status": "error", "message": "No submissions found in database."}
            
# # # # #             processed_count = 0
# # # # #             error_count = 0
            
# # # # #             for submission in submissions:
# # # # #                 try:
# # # # #                     file_path = submission['file_url']
# # # # #                     assessment_id = submission['assessment_id']
# # # # #                     student_id = submission['student_id']
                    
# # # # #                     print(f"\n📄 Processing: {pathlib.Path(file_path).name} (Assessment: {assessment_id}, Student: {student_id})")
                    
# # # # #                     # Resolve file path and load text
# # # # #                     full_path = resolve_file_path(file_path)
# # # # #                     raw_text = load_file_text(str(full_path))
                    
# # # # #                     # DEBUG: Confirm we're using the right provider
# # # # #                     print(f"🐛 DEBUG: About to extract with provider: {provider} (suffix: {provider_suffix})")
                    
# # # # #                     # Extract answers using LLM
# # # # #                     answers = extractor.extract_answers_with_llm(raw_text)

# # # # #                     if not answers:
# # # # #                         print("❌ No answers extracted.")
# # # # #                         error_count += 1
# # # # #                         continue

# # # # #                     # Print extracted answers for verification
# # # # #                     print("\nExtracted Answers (before saving to DB):\n")
# # # # #                     pprint([
# # # # #                         {
# # # # #                             "question": ans.full_question_id,
# # # # #                             "answer": ans.answer_text
# # # # #                         }
# # # # #                         for ans in answers
# # # # #                     ])

# # # # #                     # Save answers to DB using the shared db_service
# # # # #                     first = answers[0]
                    
# # # # #                     # DEBUG: Confirm table being used
# # # # #                     table_name = f"student_answers_{provider_suffix}"
# # # # #                     print(f"🐛 DEBUG: Saving to table: {table_name}")
                    
# # # # #                     db_service.save_answers(
# # # # #                         student_index=first.student_index,
# # # # #                         module_code=first.module_code,
# # # # #                         year=first.exam_year,
# # # # #                         month=first.exam_month,
# # # # #                         answers=answers
# # # # #                     )

# # # # #                     print(f"✅ Saved answers for {first.student_index} | {first.module_code} | {first.exam_year}-{first.exam_month} to {table_name}")
# # # # #                     processed_count += 1
                    
# # # # #                     # Delay to respect Gemini rate limits (15 requests/min)
# # # # #                     if provider == "GoogleGemini":
# # # # #                         print("⏱️ Waiting 10 seconds for Gemini rate limit...")
# # # # #                         time.sleep(10)
                    
# # # # #                 except Exception as e:
# # # # #                     print(f"❌ Error processing {pathlib.Path(submission['file_url']).name}: {e}")
# # # # #                     error_count += 1
# # # # #                     continue
            
# # # # #             db_service.close()
            
# # # # #             return {
# # # # #                 "status": "success",
# # # # #                 "message": f"Processing complete. {processed_count} files processed successfully, {error_count} errors. Data saved to table: student_answers_{provider_suffix}",
# # # # #                 "details": {
# # # # #                     "processed": processed_count,
# # # # #                     "errors": error_count,
# # # # #                     "table": f"student_answers_{provider_suffix}"
# # # # #                 }
# # # # #             }
            
# # # # #         elif folder:
# # # # #             # Original folder/file mode
# # # # #             processed_count = 0
# # # # #             error_count = 0
            
# # # # #             if os.path.isfile(folder) and folder.endswith(".docx"):
# # # # #                 # Single file mode
# # # # #                 try:
# # # # #                     extract_and_save(folder, extractor, provider_suffix)
# # # # #                     processed_count = 1
# # # # #                 except Exception as e:
# # # # #                     error_count = 1
# # # # #                     print(f"❌ Error processing file: {e}")
                    
# # # # #             elif os.path.isdir(folder):
# # # # #                 # Folder mode
# # # # #                 for filename in os.listdir(folder):
# # # # #                     if filename.lower().endswith(".docx"):
# # # # #                         filepath = os.path.join(folder, filename)
# # # # #                         try:
# # # # #                             extract_and_save(filepath, extractor, provider_suffix)
# # # # #                             processed_count += 1
# # # # #                         except Exception as e:
# # # # #                             error_count += 1
# # # # #                             print(f"❌ Error processing {filename}: {e}")

# # # # #                         # Delay to respect Gemini rate limits (15 requests/min)
# # # # #                         if provider == "GoogleGemini":
# # # # #                             print("⏱️ Waiting 10 seconds for Gemini rate limit...")
# # # # #                             time.sleep(10)
# # # # #             else:
# # # # #                 db_service.close()
# # # # #                 return {"status": "error", "message": "Invalid folder path. Must be either a .docx file or a directory."}
            
# # # # #             db_service.close()
            
# # # # #             return {
# # # # #                 "status": "success",
# # # # #                 "message": f"Processing complete. {processed_count} files processed successfully, {error_count} errors. Data saved to table: student_answers_{provider_suffix}",
# # # # #                 "details": {
# # # # #                     "processed": processed_count,
# # # # #                     "errors": error_count,
# # # # #                     "table": f"student_answers_{provider_suffix}"
# # # # #                 }
# # # # #             }
# # # # #         else:
# # # # #             db_service.close()
# # # # #             return {"status": "error", "message": "Please specify either folder for file/folder mode or use from_db=True for database mode."}
            
# # # # #     except Exception as e:
# # # # #         print(f"🐛 DEBUG: Exception in main(): {e}")
# # # # #         import traceback
# # # # #         traceback.print_exc()
# # # # #         return {"status": "error", "message": f"Unexpected error: {str(e)}"}

# # # # # if __name__ == "__main__":
# # # # #     # CRITICAL FIX: Proper argument parsing
# # # # #     parser = argparse.ArgumentParser(description="Extract and save student answers using LLMs")
# # # # #     parser.add_argument("--provider", required=True, choices=["OpenAI", "GoogleGemini"], help="LLM provider")
# # # # #     parser.add_argument("--model", required=True, help="Model name (e.g., gpt-4o, gemini-2.0-flash)")
# # # # #     parser.add_argument("--folder", help="Single DOCX file or folder containing DOCX files")
# # # # #     parser.add_argument("--from-db", action="store_true", help="Process files from database instead of folder")

# # # # #     args = parser.parse_args()
    
# # # # #     print(f"🐛 DEBUG: Command line args parsed:")
# # # # #     print(f"   args.provider = '{args.provider}'")
# # # # #     print(f"   args.model = '{args.model}'")
# # # # #     print(f"   args.from_db = {args.from_db}")
# # # # #     print(f"   args.folder = {args.folder}")
    
# # # # #     # CRITICAL FIX: Call the main function with the PARSED arguments, not defaults
# # # # #     result = main(
# # # # #         provider=args.provider,      # Use parsed provider
# # # # #         model=args.model,           # Use parsed model  
# # # # #         folder=args.folder,         # Use parsed folder
# # # # #         from_db=args.from_db        # Use parsed from_db flag
# # # # #     )
    
# # # # #     print(f"\nFinal Result: {result}")
    
# # # # #     if result["status"] == "error":
# # # # #         sys.exit(1)


# # # # # import sys
# # # # # import os
# # # # # import time
# # # # # import pathlib
# # # # # from docx import Document
# # # # # from pprint import pprint
# # # # # import psycopg2
# # # # # from psycopg2.extras import RealDictCursor
# # # # # from dotenv import load_dotenv
# # # # # import pdfplumber
# # # # # import argparse

# # # # # sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "../../")))

# # # # # from src.services.answer_extractor import AnswerExtractor
# # # # # from src.services.database_services.student_answer_db import StudentAnswerService

# # # # # load_dotenv()

# # # # # def get_provider_suffix(provider):
# # # # #     """Convert provider name to database table suffix."""
# # # # #     print(f"🐛 DEBUG: Converting provider '{provider}' to suffix...")
    
# # # # #     # Standardize the mapping to ensure consistent table naming
# # # # #     provider_map = {
# # # # #         "OpenAI": "openai",
# # # # #         "GoogleGemini": "gemini",
# # # # #         "Gemini": "gemini",  # Alternative name
# # # # #         "Google": "gemini"   # Another alternative
# # # # #     }
    
# # # # #     suffix = provider_map.get(provider, provider.lower())
# # # # #     print(f"🐛 DEBUG: Provider '{provider}' -> Suffix '{suffix}'")
# # # # #     return suffix

# # # # # def get_database_connection():
# # # # #     """Get database connection using environment variables."""
# # # # #     try:
# # # # #         conn = psycopg2.connect(
# # # # #             host=os.getenv('POSTGRES_HOST'),
# # # # #             port=os.getenv('POSTGRES_PORT'),
# # # # #             database=os.getenv('POSTGRES_DB'),
# # # # #             user=os.getenv('POSTGRES_USER'),
# # # # #             password=os.getenv('POSTGRES_PASSWORD')
# # # # #         )
# # # # #         return conn
# # # # #     except Exception as e:
# # # # #         print(f"Failed to connect to database: {e}")
# # # # #         sys.exit(1)

# # # # # def get_submissions_from_db():
# # # # #     """Retrieve all submission file paths from database using new schema."""
# # # # #     conn = get_database_connection()
# # # # #     try:
# # # # #         with conn.cursor(cursor_factory=RealDictCursor) as cur:
# # # # #             cur.execute("""
# # # # #                 SELECT submission_id, assessment_id, student_id, file_url, submission_start_at
# # # # #                 FROM "Submission" 
# # # # #                 ORDER BY assessment_id, submission_start_at ASC
# # # # #             """)
            
# # # # #             submissions = cur.fetchall()
# # # # #             print(f"Found {len(submissions)} submissions in database")
# # # # #             return submissions
            
# # # # #     except Exception as e:
# # # # #         print(f"Database error: {e}")
# # # # #         return []
# # # # #     finally:
# # # # #         conn.close()

# # # # # def resolve_file_path(file_path: str) -> pathlib.Path:
# # # # #     """Resolve database file path to actual file location."""
# # # # #     project_root = pathlib.Path.cwd()
# # # # #     # Try parent directory first (where data folder should be)
# # # # #     full_path = project_root.parent / file_path
    
# # # # #     if not full_path.exists():
# # # # #         # Try alternative path in project directory
# # # # #         alternative_path = project_root / file_path
# # # # #         if alternative_path.exists():
# # # # #             full_path = alternative_path
# # # # #         else:
# # # # #             raise FileNotFoundError(f"File not found at: {full_path} or {alternative_path}")
    
# # # # #     return full_path

# # # # # def load_docx_text(docx_path: str) -> str:
# # # # #     doc = Document(docx_path)
# # # # #     return "\n".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])

# # # # # def load_pdf_text(pdf_path: str) -> str:
# # # # #     with pdfplumber.open(pdf_path) as pdf:
# # # # #         return "\n".join(page.extract_text() or "" for page in pdf.pages)

# # # # # def load_file_text(file_path: str) -> str:
# # # # #     """Load text from file based on extension."""
# # # # #     path_obj = pathlib.Path(file_path)
# # # # #     if path_obj.suffix.lower() == '.docx':
# # # # #         return load_docx_text(file_path)
# # # # #     elif path_obj.suffix.lower() == '.pdf':
# # # # #         return load_pdf_text(file_path)
# # # # #     else:
# # # # #         raise ValueError(f"Unsupported file format: {path_obj.suffix}")

# # # # # def extract_and_save(docx_path: str, extractor: AnswerExtractor, provider_suffix: str):
# # # # #     """Extract and save answers - now uses provider_suffix directly."""
# # # # #     filename = os.path.basename(docx_path)
# # # # #     print(f"\n📄 Processing: {filename}")
# # # # #     print(f"🐛 DEBUG: Using provider_suffix: {provider_suffix}")

# # # # #     try:
# # # # #         raw_text = load_docx_text(docx_path)
        
# # # # #         # DEBUG: Check extractor provider before using
# # # # #         if hasattr(extractor, 'current_provider'):
# # # # #             print(f"🐛 DEBUG: Extractor is using provider: {extractor.current_provider}")
        
# # # # #         answers = extractor.extract_answers_with_llm(raw_text)

# # # # #         if not answers:
# # # # #             print("❌ No answers extracted.")
# # # # #             return

# # # # #         # Preview the result
# # # # #         pprint([
# # # # #             {"question": ans.full_question_id, "answer": ans.answer_text}
# # # # #             for ans in answers
# # # # #         ])

# # # # #         # Save to database - use the passed provider_suffix
# # # # #         first = answers[0]
# # # # #         print(f"🐛 DEBUG: Creating StudentAnswerService with suffix: {provider_suffix}")
# # # # #         db = StudentAnswerService(provider_suffix=provider_suffix)
# # # # #         db.initialize_table()
# # # # #         db.save_answers(
# # # # #             student_index=first.student_index,
# # # # #             module_code=first.module_code,
# # # # #             year=first.exam_year,
# # # # #             month=first.exam_month,
# # # # #             answers=answers
# # # # #         )
# # # # #         db.close()

# # # # #         print(f"✅ Saved answers for {first.student_index} | {first.module_code} | {first.exam_year}-{first.exam_month}")
# # # # #     except Exception as e:
# # # # #         print(f"❌ Failed to process {filename}: {e}")
# # # # #         raise  # Re-raise to see full traceback

# # # # # def main(provider, model, folder=None, from_db=True):  # ← FIXED: Removed default values
# # # # #     """
# # # # #     Main function that can be called from Flask API.
    
# # # # #     Args:
# # # # #         provider: LLM provider ("OpenAI" or "GoogleGemini") - REQUIRED
# # # # #         model: Model name (e.g., "gpt-4o", "gemini-2.0-flash") - REQUIRED
# # # # #         folder: Single file or folder path (optional)
# # # # #         from_db: Whether to process files from database (default: True)
    
# # # # #     Returns:
# # # # #         dict: Result status and message
# # # # #     """
# # # # #     try:
# # # # #         # DEBUG: Print received parameters
# # # # #         print(f"🐛 DEBUG: main() called with:")
# # # # #         print(f"   provider = '{provider}' (type: {type(provider)})")
# # # # #         print(f"   model = '{model}'")
# # # # #         print(f"   folder = {folder}")
# # # # #         print(f"   from_db = {from_db}")
        
# # # # #         # Convert provider to proper suffix at the start and validate
# # # # #         provider_suffix = get_provider_suffix(provider)
# # # # #         print(f"⚙️ Using provider: {provider} -> table suffix: {provider_suffix}")
        
# # # # #         # Validate the provider suffix
# # # # #         if provider_suffix not in ["openai", "gemini"]:
# # # # #             raise ValueError(f"Unsupported provider suffix: {provider_suffix}. Expected 'openai' or 'gemini'")
        
# # # # #         # DEBUG: Create extractor with explicit logging
# # # # #         print(f"🐛 DEBUG: Creating AnswerExtractor(selected_provider='{provider}', selected_model='{model}')")
# # # # #         extractor = AnswerExtractor(selected_provider=provider, selected_model=model)
        
# # # # #         # DEBUG: Inspect the created extractor
# # # # #         if hasattr(extractor, 'current_provider'):
# # # # #             print(f"🐛 DEBUG: AnswerExtractor.current_provider = {extractor.current_provider}")
# # # # #         if hasattr(extractor, 'provider'):
# # # # #             print(f"🐛 DEBUG: AnswerExtractor.provider = {extractor.provider}")
# # # # #         if hasattr(extractor, 'selected_provider'):
# # # # #             print(f"🐛 DEBUG: AnswerExtractor.selected_provider = {extractor.selected_provider}")
        
# # # # #         # Create database service once with the correct suffix
# # # # #         print(f"🐛 DEBUG: Creating StudentAnswerService with suffix: '{provider_suffix}'")
# # # # #         db_service = StudentAnswerService(provider_suffix=provider_suffix)
        
# # # # #         # Initialize the correct table (this should create student_answers_gemini for Gemini)
# # # # #         print(f"🐛 DEBUG: Initializing table for suffix: '{provider_suffix}'")
# # # # #         db_service.initialize_table()
        
# # # # #         if from_db:
# # # # #             # Database mode - process all files from database
# # # # #             submissions = get_submissions_from_db()
            
# # # # #             if not submissions:
# # # # #                 db_service.close()
# # # # #                 return {"status": "error", "message": "No submissions found in database."}
            
# # # # #             processed_count = 0
# # # # #             error_count = 0
            
# # # # #             for submission in submissions:
# # # # #                 try:
# # # # #                     file_path = submission['file_url']
# # # # #                     assessment_id = submission['assessment_id']
# # # # #                     student_id = submission['student_id']
                    
# # # # #                     print(f"\n📄 Processing: {pathlib.Path(file_path).name} (Assessment: {assessment_id}, Student: {student_id})")
                    
# # # # #                     # Resolve file path and load text
# # # # #                     full_path = resolve_file_path(file_path)
# # # # #                     raw_text = load_file_text(str(full_path))
                    
# # # # #                     # DEBUG: Confirm we're using the right provider
# # # # #                     print(f"🐛 DEBUG: About to extract with provider: {provider} (suffix: {provider_suffix})")
                    
# # # # #                     # Extract answers using LLM
# # # # #                     answers = extractor.extract_answers_with_llm(raw_text)

# # # # #                     if not answers:
# # # # #                         print("❌ No answers extracted.")
# # # # #                         error_count += 1
# # # # #                         continue

# # # # #                     # Print extracted answers for verification
# # # # #                     print("\nExtracted Answers (before saving to DB):\n")
# # # # #                     pprint([
# # # # #                         {
# # # # #                             "question": ans.full_question_id,
# # # # #                             "answer": ans.answer_text
# # # # #                         }
# # # # #                         for ans in answers
# # # # #                     ])

# # # # #                     # Save answers to DB using the shared db_service
# # # # #                     first = answers[0]
                    
# # # # #                     # DEBUG: Confirm table being used
# # # # #                     table_name = f"student_answers_{provider_suffix}"
# # # # #                     print(f"🐛 DEBUG: Saving to table: {table_name}")
                    
# # # # #                     db_service.save_answers(
# # # # #                         student_index=first.student_index,
# # # # #                         module_code=first.module_code,
# # # # #                         year=first.exam_year,
# # # # #                         month=first.exam_month,
# # # # #                         answers=answers
# # # # #                     )

# # # # #                     print(f"✅ Saved answers for {first.student_index} | {first.module_code} | {first.exam_year}-{first.exam_month} to {table_name}")
# # # # #                     processed_count += 1
                    
# # # # #                     # Delay to respect Gemini rate limits (15 requests/min)
# # # # #                     if provider == "GoogleGemini":
# # # # #                         print("⏱️ Waiting 10 seconds for Gemini rate limit...")
# # # # #                         time.sleep(10)
                    
# # # # #                 except Exception as e:
# # # # #                     print(f"❌ Error processing {pathlib.Path(submission['file_url']).name}: {e}")
# # # # #                     error_count += 1
# # # # #                     continue
            
# # # # #             db_service.close()
            
# # # # #             return {
# # # # #                 "status": "success",
# # # # #                 "message": f"Processing complete. {processed_count} files processed successfully, {error_count} errors. Data saved to table: student_answers_{provider_suffix}",
# # # # #                 "details": {
# # # # #                     "processed": processed_count,
# # # # #                     "errors": error_count,
# # # # #                     "table": f"student_answers_{provider_suffix}"
# # # # #                 }
# # # # #             }
            
# # # # #         elif folder:
# # # # #             # Original folder/file mode
# # # # #             processed_count = 0
# # # # #             error_count = 0
            
# # # # #             if os.path.isfile(folder) and folder.endswith(".docx"):
# # # # #                 # Single file mode
# # # # #                 try:
# # # # #                     extract_and_save(folder, extractor, provider_suffix)
# # # # #                     processed_count = 1
# # # # #                 except Exception as e:
# # # # #                     error_count = 1
# # # # #                     print(f"❌ Error processing file: {e}")
                    
# # # # #             elif os.path.isdir(folder):
# # # # #                 # Folder mode
# # # # #                 for filename in os.listdir(folder):
# # # # #                     if filename.lower().endswith(".docx"):
# # # # #                         filepath = os.path.join(folder, filename)
# # # # #                         try:
# # # # #                             extract_and_save(filepath, extractor, provider_suffix)
# # # # #                             processed_count += 1
# # # # #                         except Exception as e:
# # # # #                             error_count += 1
# # # # #                             print(f"❌ Error processing {filename}: {e}")

# # # # #                         # Delay to respect Gemini rate limits (15 requests/min)
# # # # #                         if provider == "GoogleGemini":
# # # # #                             print("⏱️ Waiting 10 seconds for Gemini rate limit...")
# # # # #                             time.sleep(10)
# # # # #             else:
# # # # #                 db_service.close()
# # # # #                 return {"status": "error", "message": "Invalid folder path. Must be either a .docx file or a directory."}
            
# # # # #             db_service.close()
            
# # # # #             return {
# # # # #                 "status": "success",
# # # # #                 "message": f"Processing complete. {processed_count} files processed successfully, {error_count} errors. Data saved to table: student_answers_{provider_suffix}",
# # # # #                 "details": {
# # # # #                     "processed": processed_count,
# # # # #                     "errors": error_count,
# # # # #                     "table": f"student_answers_{provider_suffix}"
# # # # #                 }
# # # # #             }
# # # # #         else:
# # # # #             db_service.close()
# # # # #             return {"status": "error", "message": "Please specify either folder for file/folder mode or use from_db=True for database mode."}
            
# # # # #     except Exception as e:
# # # # #         print(f"🐛 DEBUG: Exception in main(): {e}")
# # # # #         import traceback
# # # # #         traceback.print_exc()
# # # # #         return {"status": "error", "message": f"Unexpected error: {str(e)}"}

# # # # # if __name__ == "__main__":
# # # # #     # CRITICAL FIX: Proper argument parsing
# # # # #     parser = argparse.ArgumentParser(description="Extract and save student answers using LLMs")
# # # # #     parser.add_argument("--provider", required=True, choices=["OpenAI", "GoogleGemini"], help="LLM provider")
# # # # #     parser.add_argument("--model", required=True, help="Model name (e.g., gpt-4o, gemini-2.0-flash)")
# # # # #     parser.add_argument("--folder", help="Single DOCX file or folder containing DOCX files")
# # # # #     parser.add_argument("--from-db", action="store_true", help="Process files from database instead of folder")

# # # # #     args = parser.parse_args()
    
# # # # #     print(f"🐛 DEBUG: Command line args parsed:")
# # # # #     print(f"   args.provider = '{args.provider}'")
# # # # #     print(f"   args.model = '{args.model}'")
# # # # #     print(f"   args.from_db = {args.from_db}")
# # # # #     print(f"   args.folder = {args.folder}")
    
# # # # #     # CRITICAL FIX: Call the main function with the PARSED arguments, not defaults
# # # # #     result = main(
# # # # #         provider=args.provider,      # Use parsed provider
# # # # #         model=args.model,           # Use parsed model  
# # # # #         folder=args.folder,         # Use parsed folder
# # # # #         from_db=args.from_db        # Use parsed from_db flag
# # # # #     )
    
# # # # #     print(f"\nFinal Result: {result}")
    
# # # # #     if result["status"] == "error":
# # # # #         sys.exit(1)

# # # # import sys
# # # # import os
# # # # import time
# # # # import pathlib
# # # # from docx import Document
# # # # from pprint import pprint
# # # # import psycopg2
# # # # from psycopg2.extras import RealDictCursor
# # # # from dotenv import load_dotenv
# # # # import pdfplumber
# # # # import argparse

# # # # sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "../../")))

# # # # from src.services.answer_extractor import AnswerExtractor
# # # # from src.services.database_services.student_answer_db import StudentAnswerService

# # # # load_dotenv()

# # # # def get_provider_suffix(provider):
# # # #     """Convert provider name to database table suffix."""
# # # #     print(f"🐛 DEBUG: Converting provider '{provider}' to suffix...")
    
# # # #     # Standardize the mapping to ensure consistent table naming
# # # #     provider_map = {
# # # #         "OpenAI": "openai",
# # # #         "GoogleGemini": "gemini",
# # # #         "Gemini": "gemini",  # Alternative name
# # # #         "Google": "gemini"   # Another alternative
# # # #     }
    
# # # #     suffix = provider_map.get(provider, provider.lower())
# # # #     print(f"🐛 DEBUG: Provider '{provider}' -> Suffix '{suffix}'")
# # # #     return suffix

# # # # def get_database_connection():
# # # #     """Get database connection using environment variables."""
# # # #     try:
# # # #         conn = psycopg2.connect(
# # # #             host=os.getenv('POSTGRES_HOST'),
# # # #             port=os.getenv('POSTGRES_PORT'),
# # # #             database=os.getenv('POSTGRES_DB'),
# # # #             user=os.getenv('POSTGRES_USER'),
# # # #             password=os.getenv('POSTGRES_PASSWORD')
# # # #         )
# # # #         return conn
# # # #     except Exception as e:
# # # #         print(f"Failed to connect to database: {e}")
# # # #         sys.exit(1)

# # # # def get_submissions_from_db():
# # # #     """Retrieve all submission file paths from database using new schema."""
# # # #     conn = get_database_connection()
# # # #     try:
# # # #         with conn.cursor(cursor_factory=RealDictCursor) as cur:
# # # #             cur.execute("""
# # # #                 SELECT submission_id, assessment_id, student_id, file_url, submission_start_at
# # # #                 FROM "Submission" 
# # # #                 ORDER BY assessment_id, submission_start_at ASC
# # # #             """)
            
# # # #             submissions = cur.fetchall()
# # # #             print(f"Found {len(submissions)} submissions in database")
# # # #             return submissions
            
# # # #     except Exception as e:
# # # #         print(f"Database error: {e}")
# # # #         return []
# # # #     finally:
# # # #         conn.close()

# # # # def resolve_file_path(file_path: str) -> pathlib.Path:
# # # #     """Resolve database file path to actual file location."""
# # # #     project_root = pathlib.Path.cwd()
# # # #     # Try parent directory first (where data folder should be)
# # # #     full_path = project_root.parent / file_path
    
# # # #     if not full_path.exists():
# # # #         # Try alternative path in project directory
# # # #         alternative_path = project_root / file_path
# # # #         if alternative_path.exists():
# # # #             full_path = alternative_path
# # # #         else:
# # # #             raise FileNotFoundError(f"File not found at: {full_path} or {alternative_path}")
    
# # # #     return full_path

# # # # def load_docx_text(docx_path: str) -> str:
# # # #     doc = Document(docx_path)
# # # #     return "\n".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])

# # # # def load_pdf_text(pdf_path: str) -> str:
# # # #     with pdfplumber.open(pdf_path) as pdf:
# # # #         return "\n".join(page.extract_text() or "" for page in pdf.pages)

# # # # def load_file_text(file_path: str) -> str:
# # # #     """Load text from file based on extension."""
# # # #     path_obj = pathlib.Path(file_path)
# # # #     if path_obj.suffix.lower() == '.docx':
# # # #         return load_docx_text(file_path)
# # # #     elif path_obj.suffix.lower() == '.pdf':
# # # #         return load_pdf_text(file_path)
# # # #     else:
# # # #         raise ValueError(f"Unsupported file format: {path_obj.suffix}")

# # # # def extract_and_save(docx_path: str, extractor: AnswerExtractor, provider_suffix: str):
# # # #     """Extract and save answers - now uses provider_suffix directly."""
# # # #     filename = os.path.basename(docx_path)
# # # #     print(f"\n📄 Processing: {filename}")
# # # #     print(f"🐛 DEBUG: Using provider_suffix: {provider_suffix}")

# # # #     try:
# # # #         raw_text = load_docx_text(docx_path)
        
# # # #         # DEBUG: Check extractor provider before using
# # # #         if hasattr(extractor, 'current_provider'):
# # # #             print(f"🐛 DEBUG: Extractor is using provider: {extractor.current_provider}")
        
# # # #         answers = extractor.extract_answers_with_llm(raw_text)

# # # #         if not answers:
# # # #             print("❌ No answers extracted.")
# # # #             return

# # # #         # Preview the result
# # # #         pprint([
# # # #             {"question": ans.full_question_id, "answer": ans.answer_text}
# # # #             for ans in answers
# # # #         ])

# # # #         # Save to database - use the passed provider_suffix
# # # #         first = answers[0]
# # # #         print(f"🐛 DEBUG: Creating StudentAnswerService with suffix: {provider_suffix}")
# # # #         db = StudentAnswerService(provider_suffix=provider_suffix)
# # # #         db.initialize_table()
# # # #         db.save_answers(
# # # #             student_index=first.student_index,
# # # #             module_code=first.module_code,
# # # #             year=first.exam_year,
# # # #             month=first.exam_month,
# # # #             answers=answers
# # # #         )
# # # #         db.close()

# # # #         print(f"✅ Saved answers for {first.student_index} | {first.module_code} | {first.exam_year}-{first.exam_month}")
# # # #     except Exception as e:
# # # #         print(f"❌ Failed to process {filename}: {e}")
# # # #         raise  # Re-raise to see full traceback

# # # # def main(provider=None, model=None, folder=None, from_db=True, **kwargs):
# # # #     """
# # # #     Main function that can be called from Flask API or command line.
    
# # # #     Args:
# # # #         provider: LLM provider ("OpenAI" or "GoogleGemini")
# # # #         model: Model name (e.g., "gpt-4o", "gemini-2.0-flash")
# # # #         folder: Single file or folder path (optional)
# # # #         from_db: Whether to process files from database (default: True)
# # # #         **kwargs: Additional arguments from pipeline calls
    
# # # #     Returns:
# # # #         dict: Result status and message
# # # #     """
# # # #     try:
# # # #         # Handle case where provider/model might be None (pipeline calls)
# # # #         if provider is None or model is None:
# # # #             # Try to get from sys.argv if called as script
# # # #             if len(sys.argv) > 1:
# # # #                 import argparse
# # # #                 parser = argparse.ArgumentParser()
# # # #                 parser.add_argument("--provider", required=True, choices=["OpenAI", "GoogleGemini"])
# # # #                 parser.add_argument("--model", required=True)
# # # #                 parser.add_argument("--folder", default=None)
# # # #                 parser.add_argument("--from-db", action="store_true", default=True)
                
# # # #                 # Parse known args to handle cases where script is called with extra args
# # # #                 args, unknown = parser.parse_known_args()
# # # #                 provider = args.provider
# # # #                 model = args.model
# # # #                 folder = args.folder
# # # #                 from_db = args.from_db
# # # #             else:
# # # #                 raise ValueError("Provider and model must be specified either as arguments or command line parameters")
        
# # # #         # DEBUG: Print received parameters
# # # #         print(f"🐛 DEBUG: main() called with:")
# # # #         print(f"   provider = '{provider}' (type: {type(provider)})")
# # # #         print(f"   model = '{model}'")
# # # #         print(f"   folder = {folder}")
# # # #         print(f"   from_db = {from_db}")
        
# # # #         # Convert provider to proper suffix at the start and validate
# # # #         provider_suffix = get_provider_suffix(provider)
# # # #         print(f"⚙️ Using provider: {provider} -> table suffix: {provider_suffix}")
        
# # # #         # Validate the provider suffix
# # # #         if provider_suffix not in ["openai", "gemini"]:
# # # #             raise ValueError(f"Unsupported provider suffix: {provider_suffix}. Expected 'openai' or 'gemini'")
        
# # # #         # DEBUG: Create extractor with explicit logging
# # # #         print(f"🐛 DEBUG: Creating AnswerExtractor(selected_provider='{provider}', selected_model='{model}')")
# # # #         extractor = AnswerExtractor(selected_provider=provider, selected_model=model)
        
# # # #         # DEBUG: Inspect the created extractor
# # # #         if hasattr(extractor, 'current_provider'):
# # # #             print(f"🐛 DEBUG: AnswerExtractor.current_provider = {extractor.current_provider}")
# # # #         if hasattr(extractor, 'provider'):
# # # #             print(f"🐛 DEBUG: AnswerExtractor.provider = {extractor.provider}")
# # # #         if hasattr(extractor, 'selected_provider'):
# # # #             print(f"🐛 DEBUG: AnswerExtractor.selected_provider = {extractor.selected_provider}")
        
# # # #         # Create database service once with the correct suffix
# # # #         print(f"🐛 DEBUG: Creating StudentAnswerService with suffix: '{provider_suffix}'")
# # # #         db_service = StudentAnswerService(provider_suffix=provider_suffix)
        
# # # #         # Initialize the correct table (this should create student_answers_gemini for Gemini)
# # # #         print(f"🐛 DEBUG: Initializing table for suffix: '{provider_suffix}'")
# # # #         db_service.initialize_table()
        
# # # #         if from_db:
# # # #             # Database mode - process all files from database
# # # #             submissions = get_submissions_from_db()
            
# # # #             if not submissions:
# # # #                 db_service.close()
# # # #                 return {"status": "error", "message": "No submissions found in database."}
            
# # # #             processed_count = 0
# # # #             error_count = 0
            
# # # #             for submission in submissions:
# # # #                 try:
# # # #                     file_path = submission['file_url']
# # # #                     assessment_id = submission['assessment_id']
# # # #                     student_id = submission['student_id']
                    
# # # #                     print(f"\n📄 Processing: {pathlib.Path(file_path).name} (Assessment: {assessment_id}, Student: {student_id})")
                    
# # # #                     # Resolve file path and load text
# # # #                     full_path = resolve_file_path(file_path)
# # # #                     raw_text = load_file_text(str(full_path))
                    
# # # #                     # DEBUG: Confirm we're using the right provider
# # # #                     print(f"🐛 DEBUG: About to extract with provider: {provider} (suffix: {provider_suffix})")
                    
# # # #                     # Extract answers using LLM
# # # #                     answers = extractor.extract_answers_with_llm(raw_text)

# # # #                     if not answers:
# # # #                         print("❌ No answers extracted.")
# # # #                         error_count += 1
# # # #                         continue

# # # #                     # Print extracted answers for verification
# # # #                     print("\nExtracted Answers (before saving to DB):\n")
# # # #                     pprint([
# # # #                         {
# # # #                             "question": ans.full_question_id,
# # # #                             "answer": ans.answer_text
# # # #                         }
# # # #                         for ans in answers
# # # #                     ])

# # # #                     # Save answers to DB using the shared db_service
# # # #                     first = answers[0]
                    
# # # #                     # DEBUG: Confirm table being used
# # # #                     table_name = f"student_answers_{provider_suffix}"
# # # #                     print(f"🐛 DEBUG: Saving to table: {table_name}")
                    
# # # #                     db_service.save_answers(
# # # #                         student_index=first.student_index,
# # # #                         module_code=first.module_code,
# # # #                         year=first.exam_year,
# # # #                         month=first.exam_month,
# # # #                         answers=answers
# # # #                     )

# # # #                     print(f"✅ Saved answers for {first.student_index} | {first.module_code} | {first.exam_year}-{first.exam_month} to {table_name}")
# # # #                     processed_count += 1
                    
# # # #                     # Delay to respect Gemini rate limits (15 requests/min)
# # # #                     if provider == "GoogleGemini":
# # # #                         print("⏱️ Waiting 10 seconds for Gemini rate limit...")
# # # #                         time.sleep(10)
                    
# # # #                 except Exception as e:
# # # #                     print(f"❌ Error processing {pathlib.Path(submission['file_url']).name}: {e}")
# # # #                     error_count += 1
# # # #                     continue
            
# # # #             db_service.close()
            
# # # #             return {
# # # #                 "status": "success",
# # # #                 "message": f"Processing complete. {processed_count} files processed successfully, {error_count} errors. Data saved to table: student_answers_{provider_suffix}",
# # # #                 "details": {
# # # #                     "processed": processed_count,
# # # #                     "errors": error_count,
# # # #                     "table": f"student_answers_{provider_suffix}"
# # # #                 }
# # # #             }
            
# # # #         elif folder:
# # # #             # Original folder/file mode
# # # #             processed_count = 0
# # # #             error_count = 0
            
# # # #             if os.path.isfile(folder) and folder.endswith(".docx"):
# # # #                 # Single file mode
# # # #                 try:
# # # #                     extract_and_save(folder, extractor, provider_suffix)
# # # #                     processed_count = 1
# # # #                 except Exception as e:
# # # #                     error_count = 1
# # # #                     print(f"❌ Error processing file: {e}")
                    
# # # #             elif os.path.isdir(folder):
# # # #                 # Folder mode
# # # #                 for filename in os.listdir(folder):
# # # #                     if filename.lower().endswith(".docx"):
# # # #                         filepath = os.path.join(folder, filename)
# # # #                         try:
# # # #                             extract_and_save(filepath, extractor, provider_suffix)
# # # #                             processed_count += 1
# # # #                         except Exception as e:
# # # #                             error_count += 1
# # # #                             print(f"❌ Error processing {filename}: {e}")

# # # #                         # Delay to respect Gemini rate limits (15 requests/min)
# # # #                         if provider == "GoogleGemini":
# # # #                             print("⏱️ Waiting 10 seconds for Gemini rate limit...")
# # # #                             time.sleep(10)
# # # #             else:
# # # #                 db_service.close()
# # # #                 return {"status": "error", "message": "Invalid folder path. Must be either a .docx file or a directory."}
            
# # # #             db_service.close()
            
# # # #             return {
# # # #                 "status": "success",
# # # #                 "message": f"Processing complete. {processed_count} files processed successfully, {error_count} errors. Data saved to table: student_answers_{provider_suffix}",
# # # #                 "details": {
# # # #                     "processed": processed_count,
# # # #                     "errors": error_count,
# # # #                     "table": f"student_answers_{provider_suffix}"
# # # #                 }
# # # #             }
# # # #         else:
# # # #             db_service.close()
# # # #             return {"status": "error", "message": "Please specify either folder for file/folder mode or use from_db=True for database mode."}
            
# # # #     except Exception as e:
# # # #         print(f"🐛 DEBUG: Exception in main(): {e}")
# # # #         import traceback
# # # #         traceback.print_exc()
# # # #         return {"status": "error", "message": f"Unexpected error: {str(e)}"}

# # # # if __name__ == "__main__":
# # # #     # Command line argument parsing
# # # #     parser = argparse.ArgumentParser(description="Extract and save student answers using LLMs")
# # # #     parser.add_argument("--provider", required=True, choices=["OpenAI", "GoogleGemini"], help="LLM provider")
# # # #     parser.add_argument("--model", required=True, help="Model name (e.g., gpt-4o, gemini-2.0-flash)")
# # # #     parser.add_argument("--folder", help="Single DOCX file or folder containing DOCX files")
# # # #     parser.add_argument("--from-db", action="store_true", help="Process files from database instead of folder")

# # # #     args = parser.parse_args()
    
# # # #     print(f"🐛 DEBUG: Command line args parsed:")
# # # #     print(f"   args.provider = '{args.provider}'")
# # # #     print(f"   args.model = '{args.model}'")
# # # #     print(f"   args.from_db = {args.from_db}")
# # # #     print(f"   args.folder = {args.folder}")
    
# # # #     # Call the main function with the PARSED arguments
# # # #     result = main(
# # # #         provider=args.provider,      # Use parsed provider
# # # #         model=args.model,           # Use parsed model  
# # # #         folder=args.folder,         # Use parsed folder
# # # #         from_db=args.from_db        # Use parsed from_db flag
# # # #     )
    
# # # #     print(f"\nFinal Result: {result}")
    
# # # #     if result["status"] == "error":
# # # #         sys.exit(1)


# # # import sys
# # # import os
# # # import time
# # # import pathlib
# # # from docx import Document
# # # from pprint import pprint
# # # import psycopg2
# # # from psycopg2.extras import RealDictCursor
# # # from dotenv import load_dotenv
# # # import pdfplumber
# # # import argparse

# # # sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "../../")))

# # # from src.services.answer_extractor import AnswerExtractor
# # # from src.services.database_services.student_answer_db import StudentAnswerService

# # # load_dotenv()

# # # def get_provider_suffix(provider):
# # #     """Convert provider name to database table suffix."""
# # #     print(f"🐛 DEBUG: Converting provider '{provider}' to suffix...")
    
# # #     # Standardize the mapping to ensure consistent table naming
# # #     provider_map = {
# # #         "OpenAI": "openai",
# # #         "GoogleGemini": "gemini",
# # #         "Gemini": "gemini",  # Alternative name
# # #         "Google": "gemini"   # Another alternative
# # #     }
    
# # #     suffix = provider_map.get(provider, provider.lower())
# # #     print(f"🐛 DEBUG: Provider '{provider}' -> Suffix '{suffix}'")
# # #     return suffix

# # # def get_database_connection():
# # #     """Get database connection using environment variables."""
# # #     try:
# # #         conn = psycopg2.connect(
# # #             host=os.getenv('POSTGRES_HOST'),
# # #             port=os.getenv('POSTGRES_PORT'),
# # #             database=os.getenv('POSTGRES_DB'),
# # #             user=os.getenv('POSTGRES_USER'),
# # #             password=os.getenv('POSTGRES_PASSWORD')
# # #         )
# # #         return conn
# # #     except Exception as e:
# # #         print(f"Failed to connect to database: {e}")
# # #         sys.exit(1)

# # # def get_submissions_from_db():
# # #     """Retrieve all submission file paths from database using new schema."""
# # #     conn = get_database_connection()
# # #     try:
# # #         with conn.cursor(cursor_factory=RealDictCursor) as cur:
# # #             cur.execute("""
# # #                 SELECT submission_id, assessment_id, student_id, file_url, submission_start_at
# # #                 FROM "Submission" 
# # #                 ORDER BY assessment_id, submission_start_at ASC
# # #             """)
            
# # #             submissions = cur.fetchall()
# # #             print(f"Found {len(submissions)} submissions in database")
# # #             return submissions
            
# # #     except Exception as e:
# # #         print(f"Database error: {e}")
# # #         return []
# # #     finally:
# # #         conn.close()

# # # def resolve_file_path(file_path: str) -> pathlib.Path:
# # #     """Resolve database file path to actual file location."""
# # #     project_root = pathlib.Path.cwd()
# # #     # Try parent directory first (where data folder should be)
# # #     full_path = project_root.parent / file_path
    
# # #     if not full_path.exists():
# # #         # Try alternative path in project directory
# # #         alternative_path = project_root / file_path
# # #         if alternative_path.exists():
# # #             full_path = alternative_path
# # #         else:
# # #             raise FileNotFoundError(f"File not found at: {full_path} or {alternative_path}")
    
# # #     return full_path

# # # def load_docx_text(docx_path: str) -> str:
# # #     doc = Document(docx_path)
# # #     return "\n".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])

# # # def load_pdf_text(pdf_path: str) -> str:
# # #     with pdfplumber.open(pdf_path) as pdf:
# # #         return "\n".join(page.extract_text() or "" for page in pdf.pages)

# # # def load_file_text(file_path: str) -> str:
# # #     """Load text from file based on extension."""
# # #     path_obj = pathlib.Path(file_path)
# # #     if path_obj.suffix.lower() == '.docx':
# # #         return load_docx_text(file_path)
# # #     elif path_obj.suffix.lower() == '.pdf':
# # #         return load_pdf_text(file_path)
# # #     else:
# # #         raise ValueError(f"Unsupported file format: {path_obj.suffix}")

# # # def extract_and_save(docx_path: str, extractor: AnswerExtractor, provider_suffix: str):
# # #     """Extract and save answers - now uses provider_suffix directly."""
# # #     filename = os.path.basename(docx_path)
# # #     print(f"\n📄 Processing: {filename}")
# # #     print(f"🐛 DEBUG: Using provider_suffix: {provider_suffix}")

# # #     try:
# # #         raw_text = load_docx_text(docx_path)
        
# # #         # DEBUG: Check extractor provider before using
# # #         if hasattr(extractor, 'current_provider'):
# # #             print(f"🐛 DEBUG: Extractor is using provider: {extractor.current_provider}")
        
# # #         answers = extractor.extract_answers_with_llm(raw_text)

# # #         if not answers:
# # #             print("❌ No answers extracted.")
# # #             return

# # #         # Preview the result
# # #         pprint([
# # #             {"question": ans.full_question_id, "answer": ans.answer_text}
# # #             for ans in answers
# # #         ])

# # #         # Save to database - use the passed provider_suffix
# # #         first = answers[0]
# # #         print(f"🐛 DEBUG: Creating StudentAnswerService with suffix: {provider_suffix}")
# # #         db = StudentAnswerService(provider_suffix=provider_suffix)
# # #         db.initialize_table()
        
# # #         # FIX: Normalize module_code to uppercase before saving
# # #         normalized_module_code = first.module_code.upper() if first.module_code else first.module_code
# # #         print(f"🐛 DEBUG: Original module_code: '{first.module_code}' -> Normalized: '{normalized_module_code}'")
        
# # #         db.save_answers(
# # #             student_index=first.student_index,
# # #             module_code=normalized_module_code,  # Use normalized uppercase version
# # #             year=first.exam_year,
# # #             month=first.exam_month,
# # #             answers=answers
# # #         )
# # #         db.close()

# # #         print(f"✅ Saved answers for {first.student_index} | {normalized_module_code} | {first.exam_year}-{first.exam_month}")
# # #     except Exception as e:
# # #         print(f"❌ Failed to process {filename}: {e}")
# # #         raise  # Re-raise to see full traceback

# # # def main(provider=None, model=None, folder=None, from_db=True, **kwargs):
# # #     """
# # #     Main function that can be called from Flask API or command line.
    
# # #     Args:
# # #         provider: LLM provider ("OpenAI" or "GoogleGemini")
# # #         model: Model name (e.g., "gpt-4o", "gemini-2.0-flash")
# # #         folder: Single file or folder path (optional)
# # #         from_db: Whether to process files from database (default: True)
# # #         **kwargs: Additional arguments from pipeline calls
    
# # #     Returns:
# # #         dict: Result status and message
# # #     """
# # #     try:
# # #         # Handle case where provider/model might be None (pipeline calls)
# # #         if provider is None or model is None:
# # #             # Try to get from sys.argv if called as script
# # #             if len(sys.argv) > 1:
# # #                 import argparse
# # #                 parser = argparse.ArgumentParser()
# # #                 parser.add_argument("--provider", required=True, choices=["OpenAI", "GoogleGemini"])
# # #                 parser.add_argument("--model", required=True)
# # #                 parser.add_argument("--folder", default=None)
# # #                 parser.add_argument("--from-db", action="store_true", default=True)
                
# # #                 # Parse known args to handle cases where script is called with extra args
# # #                 args, unknown = parser.parse_known_args()
# # #                 provider = args.provider
# # #                 model = args.model
# # #                 folder = args.folder
# # #                 from_db = args.from_db
# # #             else:
# # #                 raise ValueError("Provider and model must be specified either as arguments or command line parameters")
        
# # #         # DEBUG: Print received parameters
# # #         print(f"🐛 DEBUG: main() called with:")
# # #         print(f"   provider = '{provider}' (type: {type(provider)})")
# # #         print(f"   model = '{model}'")
# # #         print(f"   folder = {folder}")
# # #         print(f"   from_db = {from_db}")
        
# # #         # Convert provider to proper suffix at the start and validate
# # #         provider_suffix = get_provider_suffix(provider)
# # #         print(f"⚙️ Using provider: {provider} -> table suffix: {provider_suffix}")
        
# # #         # Validate the provider suffix
# # #         if provider_suffix not in ["openai", "gemini"]:
# # #             raise ValueError(f"Unsupported provider suffix: {provider_suffix}. Expected 'openai' or 'gemini'")
        
# # #         # DEBUG: Create extractor with explicit logging
# # #         print(f"🐛 DEBUG: Creating AnswerExtractor(selected_provider='{provider}', selected_model='{model}')")
# # #         extractor = AnswerExtractor(selected_provider=provider, selected_model=model)
        
# # #         # DEBUG: Inspect the created extractor
# # #         if hasattr(extractor, 'current_provider'):
# # #             print(f"🐛 DEBUG: AnswerExtractor.current_provider = {extractor.current_provider}")
# # #         if hasattr(extractor, 'provider'):
# # #             print(f"🐛 DEBUG: AnswerExtractor.provider = {extractor.provider}")
# # #         if hasattr(extractor, 'selected_provider'):
# # #             print(f"🐛 DEBUG: AnswerExtractor.selected_provider = {extractor.selected_provider}")
        
# # #         # Create database service once with the correct suffix
# # #         print(f"🐛 DEBUG: Creating StudentAnswerService with suffix: '{provider_suffix}'")
# # #         db_service = StudentAnswerService(provider_suffix=provider_suffix)
        
# # #         # Initialize the correct table (this should create student_answers_gemini for Gemini)
# # #         print(f"🐛 DEBUG: Initializing table for suffix: '{provider_suffix}'")
# # #         db_service.initialize_table()
        
# # #         if from_db:
# # #             # Database mode - process all files from database
# # #             submissions = get_submissions_from_db()
            
# # #             if not submissions:
# # #                 db_service.close()
# # #                 return {"status": "error", "message": "No submissions found in database."}
            
# # #             processed_count = 0
# # #             error_count = 0
            
# # #             for submission in submissions:
# # #                 try:
# # #                     file_path = submission['file_url']
# # #                     assessment_id = submission['assessment_id']
# # #                     student_id = submission['student_id']
                    
# # #                     print(f"\n📄 Processing: {pathlib.Path(file_path).name} (Assessment: {assessment_id}, Student: {student_id})")
                    
# # #                     # Resolve file path and load text
# # #                     full_path = resolve_file_path(file_path)
# # #                     raw_text = load_file_text(str(full_path))
                    
# # #                     # DEBUG: Confirm we're using the right provider
# # #                     print(f"🐛 DEBUG: About to extract with provider: {provider} (suffix: {provider_suffix})")
                    
# # #                     # Extract answers using LLM
# # #                     answers = extractor.extract_answers_with_llm(raw_text)

# # #                     if not answers:
# # #                         print("❌ No answers extracted.")
# # #                         error_count += 1
# # #                         continue

# # #                     # Print extracted answers for verification
# # #                     print("\nExtracted Answers (before saving to DB):\n")
# # #                     pprint([
# # #                         {
# # #                             "question": ans.full_question_id,
# # #                             "answer": ans.answer_text
# # #                         }
# # #                         for ans in answers
# # #                     ])

# # #                     # Save answers to DB using the shared db_service
# # #                     first = answers[0]
                    
# # #                     # DEBUG: Confirm table being used
# # #                     table_name = f"student_answers_{provider_suffix}"
# # #                     print(f"🐛 DEBUG: Saving to table: {table_name}")
                    
# # #                     # FIX: Normalize module_code to uppercase before saving
# # #                     normalized_module_code = first.module_code.upper() if first.module_code else first.module_code
# # #                     print(f"🐛 DEBUG: Original module_code: '{first.module_code}' -> Normalized: '{normalized_module_code}'")
                    
# # #                     db_service.save_answers(
# # #                         student_index=first.student_index,
# # #                         module_code=normalized_module_code,  # Use normalized uppercase version
# # #                         year=first.exam_year,
# # #                         month=first.exam_month,
# # #                         answers=answers
# # #                     )

# # #                     print(f"✅ Saved answers for {first.student_index} | {normalized_module_code} | {first.exam_year}-{first.exam_month} to {table_name}")
# # #                     processed_count += 1
                    
# # #                     # Delay to respect Gemini rate limits (15 requests/min)
# # #                     if provider == "GoogleGemini":
# # #                         print("⏱️ Waiting 10 seconds for Gemini rate limit...")
# # #                         time.sleep(10)
                    
# # #                 except Exception as e:
# # #                     print(f"❌ Error processing {pathlib.Path(submission['file_url']).name}: {e}")
# # #                     error_count += 1
# # #                     continue
            
# # #             db_service.close()
            
# # #             return {
# # #                 "status": "success",
# # #                 "message": f"Processing complete. {processed_count} files processed successfully, {error_count} errors. Data saved to table: student_answers_{provider_suffix}",
# # #                 "details": {
# # #                     "processed": processed_count,
# # #                     "errors": error_count,
# # #                     "table": f"student_answers_{provider_suffix}"
# # #                 }
# # #             }
            
# # #         elif folder:
# # #             # Original folder/file mode
# # #             processed_count = 0
# # #             error_count = 0
            
# # #             if os.path.isfile(folder) and folder.endswith(".docx"):
# # #                 # Single file mode
# # #                 try:
# # #                     extract_and_save(folder, extractor, provider_suffix)
# # #                     processed_count = 1
# # #                 except Exception as e:
# # #                     error_count = 1
# # #                     print(f"❌ Error processing file: {e}")
                    
# # #             elif os.path.isdir(folder):
# # #                 # Folder mode
# # #                 for filename in os.listdir(folder):
# # #                     if filename.lower().endswith(".docx"):
# # #                         filepath = os.path.join(folder, filename)
# # #                         try:
# # #                             extract_and_save(filepath, extractor, provider_suffix)
# # #                             processed_count += 1
# # #                         except Exception as e:
# # #                             error_count += 1
# # #                             print(f"❌ Error processing {filename}: {e}")

# # #                         # Delay to respect Gemini rate limits (15 requests/min)
# # #                         if provider == "GoogleGemini":
# # #                             print("⏱️ Waiting 10 seconds for Gemini rate limit...")
# # #                             time.sleep(10)
# # #             else:
# # #                 db_service.close()
# # #                 return {"status": "error", "message": "Invalid folder path. Must be either a .docx file or a directory."}
            
# # #             db_service.close()
            
# # #             return {
# # #                 "status": "success",
# # #                 "message": f"Processing complete. {processed_count} files processed successfully, {error_count} errors. Data saved to table: student_answers_{provider_suffix}",
# # #                 "details": {
# # #                     "processed": processed_count,
# # #                     "errors": error_count,
# # #                     "table": f"student_answers_{provider_suffix}"
# # #                 }
# # #             }
# # #         else:
# # #             db_service.close()
# # #             return {"status": "error", "message": "Please specify either folder for file/folder mode or use from_db=True for database mode."}
            
# # #     except Exception as e:
# # #         print(f"🐛 DEBUG: Exception in main(): {e}")
# # #         import traceback
# # #         traceback.print_exc()
# # #         return {"status": "error", "message": f"Unexpected error: {str(e)}"}

# # # if __name__ == "__main__":
# # #     # Command line argument parsing
# # #     parser = argparse.ArgumentParser(description="Extract and save student answers using LLMs")
# # #     parser.add_argument("--provider", required=True, choices=["OpenAI", "GoogleGemini"], help="LLM provider")
# # #     parser.add_argument("--model", required=True, help="Model name (e.g., gpt-4o, gemini-2.0-flash)")
# # #     parser.add_argument("--folder", help="Single DOCX file or folder containing DOCX files")
# # #     parser.add_argument("--from-db", action="store_true", help="Process files from database instead of folder")

# # #     args = parser.parse_args()
    
# # #     print(f"🐛 DEBUG: Command line args parsed:")
# # #     print(f"   args.provider = '{args.provider}'")
# # #     print(f"   args.model = '{args.model}'")
# # #     print(f"   args.from_db = {args.from_db}")
# # #     print(f"   args.folder = {args.folder}")
    
# # #     # Call the main function with the PARSED arguments
# # #     result = main(
# # #         provider=args.provider,      # Use parsed provider
# # #         model=args.model,           # Use parsed model  
# # #         folder=args.folder,         # Use parsed folder
# # #         from_db=args.from_db        # Use parsed from_db flag
# # #     )
    
# # #     print(f"\nFinal Result: {result}")
    
# # #     if result["status"] == "error":
# # #         sys.exit(1)

# # import sys
# # import os
# # import time
# # import pathlib
# # from docx import Document
# # from pprint import pprint
# # import psycopg2
# # from psycopg2.extras import RealDictCursor
# # from dotenv import load_dotenv
# # import pdfplumber
# # import argparse

# # sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "../../")))

# # from src.services.answer_extractor import AnswerExtractor
# # from src.services.database_services.student_answer_db import StudentAnswerService

# # load_dotenv()

# # def get_provider_suffix(provider):
# #     """Convert provider name to database table suffix."""
# #     print(f"🐛 DEBUG: Converting provider '{provider}' to suffix...")
    
# #     # Standardize the mapping to ensure consistent table naming
# #     provider_map = {
# #         "OpenAI": "openai",
# #         "GoogleGemini": "gemini",
# #         "Gemini": "gemini",  # Alternative name
# #         "Google": "gemini"   # Another alternative
# #     }
    
# #     suffix = provider_map.get(provider, provider.lower())
# #     print(f"🐛 DEBUG: Provider '{provider}' -> Suffix '{suffix}'")
# #     return suffix

# # def get_database_connection():
# #     """Get database connection using environment variables."""
# #     try:
# #         conn = psycopg2.connect(
# #             host=os.getenv('POSTGRES_HOST'),
# #             port=os.getenv('POSTGRES_PORT'),
# #             database=os.getenv('POSTGRES_DB'),
# #             user=os.getenv('POSTGRES_USER'),
# #             password=os.getenv('POSTGRES_PASSWORD')
# #         )
# #         return conn
# #     except Exception as e:
# #         print(f"Failed to connect to database: {e}")
# #         sys.exit(1)

# # def get_submissions_from_db():
# #     """Retrieve all submission file paths from database using new schema."""
# #     conn = get_database_connection()
# #     try:
# #         with conn.cursor(cursor_factory=RealDictCursor) as cur:
# #             cur.execute("""
# #                 SELECT submission_id, assessment_id, student_id, file_url, submission_start_at
# #                 FROM "Submission" 
# #                 ORDER BY assessment_id, submission_start_at ASC
# #             """)
            
# #             submissions = cur.fetchall()
# #             print(f"Found {len(submissions)} submissions in database")
# #             return submissions
            
# #     except Exception as e:
# #         print(f"Database error: {e}")
# #         return []
# #     finally:
# #         conn.close()

# # def resolve_file_path(file_path: str) -> pathlib.Path:
# #     """Resolve database file path to actual file location."""
# #     project_root = pathlib.Path.cwd()
# #     # Try parent directory first (where data folder should be)
# #     full_path = project_root.parent / file_path
    
# #     if not full_path.exists():
# #         # Try alternative path in project directory
# #         alternative_path = project_root / file_path
# #         if alternative_path.exists():
# #             full_path = alternative_path
# #         else:
# #             raise FileNotFoundError(f"File not found at: {full_path} or {alternative_path}")
    
# #     return full_path

# # def load_docx_text(docx_path: str) -> str:
# #     doc = Document(docx_path)
# #     return "\n".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])

# # def load_pdf_text(pdf_path: str) -> str:
# #     with pdfplumber.open(pdf_path) as pdf:
# #         return "\n".join(page.extract_text() or "" for page in pdf.pages)

# # def load_file_text(file_path: str) -> str:
# #     """Load text from file based on extension."""
# #     path_obj = pathlib.Path(file_path)
# #     if path_obj.suffix.lower() == '.docx':
# #         return load_docx_text(file_path)
# #     elif path_obj.suffix.lower() == '.pdf':
# #         return load_pdf_text(file_path)
# #     else:
# #         raise ValueError(f"Unsupported file format: {path_obj.suffix}")

# # def extract_and_save(docx_path: str, extractor: AnswerExtractor, provider_suffix: str):
# #     """Extract and save answers - now uses provider_suffix directly."""
# #     filename = os.path.basename(docx_path)
# #     print(f"\n📄 Processing: {filename}")
# #     print(f"🐛 DEBUG: Using provider_suffix: {provider_suffix}")

# #     try:
# #         raw_text = load_docx_text(docx_path)
        
# #         # DEBUG: Check extractor provider before using
# #         if hasattr(extractor, 'current_provider'):
# #             print(f"🐛 DEBUG: Extractor is using provider: {extractor.current_provider}")
        
# #         # FIX: Try extraction with error handling for JSON parsing issues
# #         try:
# #             answers = extractor.extract_answers_with_llm(raw_text)
# #         except Exception as json_error:
# #             if "Invalid control character" in str(json_error):
# #                 print(f"⚠️ JSON parsing error detected: {json_error}")
# #                 print("🔧 This is likely due to control characters in LLM response. Retrying...")
# #                 # You might want to implement retry logic here or handle differently
# #                 answers = None
# #             else:
# #                 raise json_error

# #         if not answers:
# #             print("❌ No answers extracted.")
# #             return

# #         # Preview the result
# #         pprint([
# #             {"question": ans.full_question_id, "answer": ans.answer_text}
# #             for ans in answers
# #         ])

# #         # Save to database - use the passed provider_suffix
# #         first = answers[0]
# #         print(f"🐛 DEBUG: Creating StudentAnswerService with suffix: {provider_suffix}")
# #         db = StudentAnswerService(provider_suffix=provider_suffix)
# #         db.initialize_table()
        
# #         # FIX: Normalize module_code to uppercase before saving
# #         normalized_module_code = first.module_code.upper() if first.module_code else first.module_code
# #         print(f"🐛 DEBUG: Original module_code: '{first.module_code}' -> Normalized: '{normalized_module_code}'")
        
# #         db.save_answers(
# #             student_index=first.student_index,
# #             module_code=normalized_module_code,  # Use normalized uppercase version
# #             year=first.exam_year,
# #             month=first.exam_month,
# #             answers=answers
# #         )
# #         db.close()

# #         print(f"✅ Saved answers for {first.student_index} | {normalized_module_code} | {first.exam_year}-{first.exam_month}")
# #     except Exception as e:
# #         print(f"❌ Failed to process {filename}: {e}")
# #         raise  # Re-raise to see full traceback

# # def main(provider=None, model=None, folder=None, from_db=True, **kwargs):
# #     """
# #     Main function that can be called from Flask API or command line.
    
# #     Args:
# #         provider: LLM provider ("OpenAI" or "GoogleGemini")
# #         model: Model name (e.g., "gpt-4o", "gemini-2.0-flash")
# #         folder: Single file or folder path (optional)
# #         from_db: Whether to process files from database (default: True)
# #         **kwargs: Additional arguments from pipeline calls
    
# #     Returns:
# #         dict: Result status and message
# #     """
# #     try:
# #         # Handle case where provider/model might be None (pipeline calls)
# #         if provider is None or model is None:
# #             # Try to get from sys.argv if called as script
# #             if len(sys.argv) > 1:
# #                 import argparse
# #                 parser = argparse.ArgumentParser()
# #                 parser.add_argument("--provider", required=True, choices=["OpenAI", "GoogleGemini"])
# #                 parser.add_argument("--model", required=True)
# #                 parser.add_argument("--folder", default=None)
# #                 parser.add_argument("--from-db", action="store_true", default=True)
                
# #                 # Parse known args to handle cases where script is called with extra args
# #                 args, unknown = parser.parse_known_args()
# #                 provider = args.provider
# #                 model = args.model
# #                 folder = args.folder
# #                 from_db = args.from_db
# #             else:
# #                 raise ValueError("Provider and model must be specified either as arguments or command line parameters")
        
# #         # DEBUG: Print received parameters
# #         print(f"🐛 DEBUG: main() called with:")
# #         print(f"   provider = '{provider}' (type: {type(provider)})")
# #         print(f"   model = '{model}'")
# #         print(f"   folder = {folder}")
# #         print(f"   from_db = {from_db}")
        
# #         # Convert provider to proper suffix at the start and validate
# #         provider_suffix = get_provider_suffix(provider)
# #         print(f"⚙️ Using provider: {provider} -> table suffix: {provider_suffix}")
        
# #         # Validate the provider suffix
# #         if provider_suffix not in ["openai", "gemini"]:
# #             raise ValueError(f"Unsupported provider suffix: {provider_suffix}. Expected 'openai' or 'gemini'")
        
# #         # DEBUG: Create extractor with explicit logging
# #         print(f"🐛 DEBUG: Creating AnswerExtractor(selected_provider='{provider}', selected_model='{model}')")
# #         extractor = AnswerExtractor(selected_provider=provider, selected_model=model)
        
# #         # DEBUG: Inspect the created extractor
# #         if hasattr(extractor, 'current_provider'):
# #             print(f"🐛 DEBUG: AnswerExtractor.current_provider = {extractor.current_provider}")
# #         if hasattr(extractor, 'provider'):
# #             print(f"🐛 DEBUG: AnswerExtractor.provider = {extractor.provider}")
# #         if hasattr(extractor, 'selected_provider'):
# #             print(f"🐛 DEBUG: AnswerExtractor.selected_provider = {extractor.selected_provider}")
        
# #         # Create database service once with the correct suffix
# #         print(f"🐛 DEBUG: Creating StudentAnswerService with suffix: '{provider_suffix}'")
# #         db_service = StudentAnswerService(provider_suffix=provider_suffix)
        
# #         # Initialize the correct table (this should create student_answers_gemini for Gemini)
# #         print(f"🐛 DEBUG: Initializing table for suffix: '{provider_suffix}'")
# #         db_service.initialize_table()
        
# #         if from_db:
# #             # Database mode - process all files from database
# #             submissions = get_submissions_from_db()
            
# #             if not submissions:
# #                 db_service.close()
# #                 return {"status": "error", "message": "No submissions found in database."}
            
# #             processed_count = 0
# #             error_count = 0
            
# #             for submission in submissions:
# #                 try:
# #                     file_path = submission['file_url']
# #                     assessment_id = submission['assessment_id']
# #                     student_id = submission['student_id']
                    
# #                     print(f"\n📄 Processing: {pathlib.Path(file_path).name} (Assessment: {assessment_id}, Student: {student_id})")
                    
# #                     # Resolve file path and load text
# #                     full_path = resolve_file_path(file_path)
# #                     raw_text = load_file_text(str(full_path))
                    
# #                     # DEBUG: Confirm we're using the right provider
# #                     print(f"🐛 DEBUG: About to extract with provider: {provider} (suffix: {provider_suffix})")
                    
# #                     # FIX: Extract answers using LLM with error handling for JSON parsing issues
# #                     try:
# #                         answers = extractor.extract_answers_with_llm(raw_text)
# #                     except Exception as json_error:
# #                         if "Invalid control character" in str(json_error):
# #                             print(f"⚠️ JSON parsing error detected: {json_error}")
# #                             print("🔧 This is likely due to control characters in LLM response. Retrying...")
# #                             # You might want to implement retry logic here or handle differently
# #                             answers = None
# #                         else:
# #                             raise json_error

# #                     if not answers:
# #                         print("❌ No answers extracted.")
# #                         error_count += 1
# #                         continue

# #                     # Print extracted answers for verification
# #                     print("\nExtracted Answers (before saving to DB):\n")
# #                     pprint([
# #                         {
# #                             "question": ans.full_question_id,
# #                             "answer": ans.answer_text
# #                         }
# #                         for ans in answers
# #                     ])

# #                     # Save answers to DB using the shared db_service
# #                     first = answers[0]
                    
# #                     # DEBUG: Confirm table being used
# #                     table_name = f"student_answers_{provider_suffix}"
# #                     print(f"🐛 DEBUG: Saving to table: {table_name}")
                    
# #                     # FIX: Normalize module_code to uppercase before saving
# #                     normalized_module_code = first.module_code.upper() if first.module_code else first.module_code
# #                     print(f"🐛 DEBUG: Original module_code: '{first.module_code}' -> Normalized: '{normalized_module_code}'")
                    
# #                     db_service.save_answers(
# #                         student_index=first.student_index,
# #                         module_code=normalized_module_code,  # Use normalized uppercase version
# #                         year=first.exam_year,
# #                         month=first.exam_month,
# #                         answers=answers
# #                     )

# #                     print(f"✅ Saved answers for {first.student_index} | {normalized_module_code} | {first.exam_year}-{first.exam_month} to {table_name}")
# #                     processed_count += 1
                    
# #                     # Delay to respect Gemini rate limits (15 requests/min)
# #                     if provider == "GoogleGemini":
# #                         print("⏱️ Waiting 10 seconds for Gemini rate limit...")
# #                         time.sleep(10)
                    
# #                 except Exception as e:
# #                     print(f"❌ Error processing {pathlib.Path(submission['file_url']).name}: {e}")
# #                     error_count += 1
# #                     continue
            
# #             db_service.close()
            
# #             return {
# #                 "status": "success",
# #                 "message": f"Processing complete. {processed_count} files processed successfully, {error_count} errors. Data saved to table: student_answers_{provider_suffix}",
# #                 "details": {
# #                     "processed": processed_count,
# #                     "errors": error_count,
# #                     "table": f"student_answers_{provider_suffix}"
# #                 }
# #             }
            
# #         elif folder:
# #             # Original folder/file mode
# #             processed_count = 0
# #             error_count = 0
            
# #             if os.path.isfile(folder) and folder.endswith(".docx"):
# #                 # Single file mode
# #                 try:
# #                     extract_and_save(folder, extractor, provider_suffix)
# #                     processed_count = 1
# #                 except Exception as e:
# #                     error_count = 1
# #                     print(f"❌ Error processing file: {e}")
                    
# #             elif os.path.isdir(folder):
# #                 # Folder mode
# #                 for filename in os.listdir(folder):
# #                     if filename.lower().endswith(".docx"):
# #                         filepath = os.path.join(folder, filename)
# #                         try:
# #                             extract_and_save(filepath, extractor, provider_suffix)
# #                             processed_count += 1
# #                         except Exception as e:
# #                             error_count += 1
# #                             print(f"❌ Error processing {filename}: {e}")

# #                         # Delay to respect Gemini rate limits (15 requests/min)
# #                         if provider == "GoogleGemini":
# #                             print("⏱️ Waiting 10 seconds for Gemini rate limit...")
# #                             time.sleep(10)
# #             else:
# #                 db_service.close()
# #                 return {"status": "error", "message": "Invalid folder path. Must be either a .docx file or a directory."}
            
# #             db_service.close()
            
# #             return {
# #                 "status": "success",
# #                 "message": f"Processing complete. {processed_count} files processed successfully, {error_count} errors. Data saved to table: student_answers_{provider_suffix}",
# #                 "details": {
# #                     "processed": processed_count,
# #                     "errors": error_count,
# #                     "table": f"student_answers_{provider_suffix}"
# #                 }
# #             }
# #         else:
# #             db_service.close()
# #             return {"status": "error", "message": "Please specify either folder for file/folder mode or use from_db=True for database mode."}
            
# #     except Exception as e:
# #         print(f"🐛 DEBUG: Exception in main(): {e}")
# #         import traceback
# #         traceback.print_exc()
# #         return {"status": "error", "message": f"Unexpected error: {str(e)}"}

# # if __name__ == "__main__":
# #     # Command line argument parsing
# #     parser = argparse.ArgumentParser(description="Extract and save student answers using LLMs")
# #     parser.add_argument("--provider", required=True, choices=["OpenAI", "GoogleGemini"], help="LLM provider")
# #     parser.add_argument("--model", required=True, help="Model name (e.g., gpt-4o, gemini-2.0-flash)")
# #     parser.add_argument("--folder", help="Single DOCX file or folder containing DOCX files")
# #     parser.add_argument("--from-db", action="store_true", help="Process files from database instead of folder")

# #     args = parser.parse_args()
    
# #     print(f"🐛 DEBUG: Command line args parsed:")
# #     print(f"   args.provider = '{args.provider}'")
# #     print(f"   args.model = '{args.model}'")
# #     print(f"   args.from_db = {args.from_db}")
# #     print(f"   args.folder = {args.folder}")
    
# #     # Call the main function with the PARSED arguments
# #     result = main(
# #         provider=args.provider,      # Use parsed provider
# #         model=args.model,           # Use parsed model  
# #         folder=args.folder,         # Use parsed folder
# #         from_db=args.from_db        # Use parsed from_db flag
# #     )
    
# #     print(f"\nFinal Result: {result}")
    
# #     if result["status"] == "error":
# #         sys.exit(1)

# import sys
# import os
# import time
# import pathlib
# import json
# import re
# from docx import Document
# from pprint import pprint
# import psycopg2
# from psycopg2.extras import RealDictCursor
# from dotenv import load_dotenv
# import pdfplumber
# import argparse

# sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "../../")))

# from src.services.answer_extractor import AnswerExtractor
# from src.services.database_services.student_answer_db import StudentAnswerService

# load_dotenv()

# def clean_json_string(json_str):
#     """Clean JSON string by removing/escaping problematic control characters."""
#     if not json_str:
#         return json_str
    
#     # Remove or replace common problematic control characters
#     # Keep only allowed whitespace characters (space, tab, newline, carriage return)
#     cleaned = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', json_str)
    
#     # Additional cleaning - escape unescaped quotes within strings
#     # This is a basic approach - you might need more sophisticated handling
#     return cleaned

# def safe_json_parse(json_str, max_retries=3):
#     """Safely parse JSON with error handling and cleaning."""
#     for attempt in range(max_retries):
#         try:
#             # First attempt - try parsing as-is
#             if attempt == 0:
#                 return json.loads(json_str)
            
#             # Second attempt - clean the string
#             elif attempt == 1:
#                 cleaned_str = clean_json_string(json_str)
#                 print(f"🔧 Attempt {attempt + 1}: Trying with cleaned JSON string")
#                 return json.loads(cleaned_str)
            
#             # Third attempt - more aggressive cleaning
#             else:
#                 # Try to extract just the JSON part if it's embedded in other text
#                 json_match = re.search(r'\{.*\}|\[.*\]', json_str, re.DOTALL)
#                 if json_match:
#                     json_part = json_match.group()
#                     cleaned_part = clean_json_string(json_part)
#                     print(f"🔧 Attempt {attempt + 1}: Trying with extracted and cleaned JSON")
#                     return json.loads(cleaned_part)
#                 else:
#                     raise ValueError("No valid JSON structure found")
                    
#         except json.JSONDecodeError as e:
#             print(f"❌ JSON parse attempt {attempt + 1} failed: {e}")
#             if attempt == max_retries - 1:
#                 print("📝 Problematic JSON content (first 500 chars):")
#                 print(repr(json_str[:500]))
#                 raise e
#             continue
    
#     return None

# def get_provider_suffix(provider):
#     """Convert provider name to database table suffix."""
#     print(f"🐛 DEBUG: Converting provider '{provider}' to suffix...")
    
#     # Standardize the mapping to ensure consistent table naming
#     provider_map = {
#         "OpenAI": "openai",
#         "GoogleGemini": "gemini",
#         "Gemini": "gemini",  # Alternative name
#         "Google": "gemini"   # Another alternative
#     }
    
#     suffix = provider_map.get(provider, provider.lower())
#     print(f"🐛 DEBUG: Provider '{provider}' -> Suffix '{suffix}'")
#     return suffix

# def get_database_connection():
#     """Get database connection using environment variables."""
#     try:
#         conn = psycopg2.connect(
#             host=os.getenv('POSTGRES_HOST'),
#             port=os.getenv('POSTGRES_PORT'),
#             database=os.getenv('POSTGRES_DB'),
#             user=os.getenv('POSTGRES_USER'),
#             password=os.getenv('POSTGRES_PASSWORD')
#         )
#         return conn
#     except Exception as e:
#         print(f"Failed to connect to database: {e}")
#         sys.exit(1)

# # def get_submissions_from_db():
# #     """Retrieve all submission file paths from database using new schema."""
# #     conn = get_database_connection()
# #     try:
# #         with conn.cursor(cursor_factory=RealDictCursor) as cur:
# #             cur.execute("""
# #                 SELECT submission_id, assessment_id, student_id, file_url, submission_start_at
# #                 FROM "Submission" 
# #                 ORDER BY assessment_id, submission_start_at ASC
# #             """)
            
# #             submissions = cur.fetchall()
# #             print(f"Found {len(submissions)} submissions in database")
# #             return submissions
            
# #     except Exception as e:
# #         print(f"Database error: {e}")
# #         return []
# #     finally:
# #         conn.close()


# def get_submissions_from_db(assessment_id=None, selected_ids=None):
#     """Retrieve submission file paths filtered by assessment and optionally by selected IDs."""
#     conn = get_database_connection()
#     try:
#         with conn.cursor(cursor_factory=RealDictCursor) as cur:
#             query = """
#                 SELECT submission_id, assessment_id, student_id, file_url, submission_start_at
#                 FROM "Submission"
#                 WHERE 1=1
#             """
#             params = []

#             if assessment_id:
#                 query += " AND assessment_id = %s"
#                 params.append(assessment_id)

#             if selected_ids:
#                 query += " AND submission_id = ANY(%s)"
#                 params.append(selected_ids)

#             query += " ORDER BY submission_start_at ASC"

#             cur.execute(query, params)
#             submissions = cur.fetchall()
#             print(f"Found {len(submissions)} submissions in database for assessment {assessment_id}")
#             return submissions
#     except Exception as e:
#         print(f"Database error: {e}")
#         return []
#     finally:
#         conn.close()

# def resolve_file_path(file_path: str) -> pathlib.Path:
#     """Resolve database file path to actual file location."""
#     project_root = pathlib.Path.cwd()
#     # Try parent directory first (where data folder should be)
#     full_path = project_root.parent / file_path
    
#     if not full_path.exists():
#         # Try alternative path in project directory
#         alternative_path = project_root / file_path
#         if alternative_path.exists():
#             full_path = alternative_path
#         else:
#             raise FileNotFoundError(f"File not found at: {full_path} or {alternative_path}")
    
#     return full_path

# def load_docx_text(docx_path: str) -> str:
#     doc = Document(docx_path)
#     return "\n".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])

# def load_pdf_text(pdf_path: str) -> str:
#     with pdfplumber.open(pdf_path) as pdf:
#         return "\n".join(page.extract_text() or "" for page in pdf.pages)

# def load_file_text(file_path: str) -> str:
#     """Load text from file based on extension."""
#     path_obj = pathlib.Path(file_path)
#     if path_obj.suffix.lower() == '.docx':
#         return load_docx_text(file_path)
#     elif path_obj.suffix.lower() == '.pdf':
#         return load_pdf_text(file_path)
#     else:
#         raise ValueError(f"Unsupported file format: {path_obj.suffix}")

# def extract_and_save(docx_path: str, extractor: AnswerExtractor, provider_suffix: str):
#     """Extract and save answers - now uses provider_suffix directly."""
#     filename = os.path.basename(docx_path)
#     print(f"\n📄 Processing: {filename}")
#     print(f"🐛 DEBUG: Using provider_suffix: {provider_suffix}")

#     try:
#         raw_text = load_docx_text(docx_path)
        
#         # DEBUG: Check extractor provider before using
#         if hasattr(extractor, 'current_provider'):
#             print(f"🐛 DEBUG: Extractor is using provider: {extractor.current_provider}")
        
#         # Enhanced extraction with robust error handling
#         answers = None
#         max_extraction_attempts = 3
        
#         for attempt in range(max_extraction_attempts):
#             try:
#                 print(f"🔄 Extraction attempt {attempt + 1}/{max_extraction_attempts}")
#                 answers = extractor.extract_answers_with_llm(raw_text)
                
#                 if answers:
#                     print(f"✅ Successfully extracted {len(answers)} answers on attempt {attempt + 1}")
#                     break
#                 else:
#                     print(f"⚠️ No answers extracted on attempt {attempt + 1}")
                    
#             except json.JSONDecodeError as json_error:
#                 print(f"❌ JSON parsing error on attempt {attempt + 1}: {json_error}")
                
#                 if "Invalid control character" in str(json_error):
#                     print("🔧 Detected control character issue - this might be due to LLM response formatting")
                    
#                     # If this is not the last attempt, wait and retry
#                     if attempt < max_extraction_attempts - 1:
#                         print(f"⏳ Waiting 5 seconds before retry...")
#                         time.sleep(5)
#                         continue
#                     else:
#                         print("💥 All extraction attempts failed due to JSON parsing errors")
#                         answers = None
#                         break
#                 else:
#                     # For other JSON errors, re-raise immediately
#                     raise json_error
                    
#             except Exception as other_error:
#                 print(f"❌ Other error on attempt {attempt + 1}: {other_error}")
#                 if attempt == max_extraction_attempts - 1:
#                     raise other_error
#                 else:
#                     print(f"⏳ Waiting 5 seconds before retry...")
#                     time.sleep(5)
#                     continue

#         if not answers:
#             print("❌ No answers extracted after all attempts.")
#             return

#         # Preview the result
#         pprint([
#             {"question": ans.full_question_id, "answer": ans.answer_text}
#             for ans in answers
#         ])

#         # Save to database - use the passed provider_suffix
#         first = answers[0]
#         print(f"🐛 DEBUG: Creating StudentAnswerService with suffix: {provider_suffix}")
#         db = StudentAnswerService(provider_suffix=provider_suffix)
#         db.initialize_table()
        
#         # FIX: Normalize module_code to uppercase before saving
#         normalized_module_code = first.module_code.upper() if first.module_code else first.module_code
#         print(f"🐛 DEBUG: Original module_code: '{first.module_code}' -> Normalized: '{normalized_module_code}'")
        
#         db.save_answers(
#             student_index=first.student_index,
#             module_code=normalized_module_code,  # Use normalized uppercase version
#             year=first.exam_year,
#             month=first.exam_month,
#             answers=answers,
#             assessment_id=submission["assessment_id"] 
#         )
#         db.close()

#         print(f"✅ Saved answers for {first.student_index} | {normalized_module_code} | {first.exam_year}-{first.exam_month}")
#     except Exception as e:
#         print(f"❌ Failed to process {filename}: {e}")
#         raise  # Re-raise to see full traceback

# def main(provider=None, model=None, folder=None, from_db=True, **kwargs):
#     """
#     Main function that can be called from Flask API or command line.
    
#     Args:
#         provider: LLM provider ("OpenAI" or "GoogleGemini")
#         model: Model name (e.g., "gpt-4o", "gemini-2.0-flash")
#         folder: Single file or folder path (optional)
#         from_db: Whether to process files from database (default: True)
#         **kwargs: Additional arguments from pipeline calls
    
#     Returns:
#         dict: Result status and message
#     """
#     try:
#         # Handle case where provider/model might be None (pipeline calls)
#         if provider is None or model is None:
#             # Try to get from sys.argv if called as script
#             if len(sys.argv) > 1:
#                 import argparse
#                 parser = argparse.ArgumentParser()
#                 parser.add_argument("--provider", required=True, choices=["OpenAI", "GoogleGemini"])
#                 parser.add_argument("--model", required=True)
#                 parser.add_argument("--folder", default=None)
#                 parser.add_argument("--from-db", action="store_true", default=True)
                
#                 # Parse known args to handle cases where script is called with extra args
#                 args, unknown = parser.parse_known_args()
#                 provider = args.provider
#                 model = args.model
#                 folder = args.folder
#                 from_db = args.from_db
#             else:
#                 raise ValueError("Provider and model must be specified either as arguments or command line parameters")
        
#         # DEBUG: Print received parameters
#         print(f"🐛 DEBUG: main() called with:")
#         print(f"   provider = '{provider}' (type: {type(provider)})")
#         print(f"   model = '{model}'")
#         print(f"   folder = {folder}")
#         print(f"   from_db = {from_db}")
        
#         # Convert provider to proper suffix at the start and validate
#         provider_suffix = get_provider_suffix(provider)
#         print(f"⚙️ Using provider: {provider} -> table suffix: {provider_suffix}")
        
#         # Validate the provider suffix
#         if provider_suffix not in ["openai", "gemini"]:
#             raise ValueError(f"Unsupported provider suffix: {provider_suffix}. Expected 'openai' or 'gemini'")
        
#         # DEBUG: Create extractor with explicit logging
#         print(f"🐛 DEBUG: Creating AnswerExtractor(selected_provider='{provider}', selected_model='{model}')")
#         extractor = AnswerExtractor(selected_provider=provider, selected_model=model)
        
#         # DEBUG: Inspect the created extractor
#         if hasattr(extractor, 'current_provider'):
#             print(f"🐛 DEBUG: AnswerExtractor.current_provider = {extractor.current_provider}")
#         if hasattr(extractor, 'provider'):
#             print(f"🐛 DEBUG: AnswerExtractor.provider = {extractor.provider}")
#         if hasattr(extractor, 'selected_provider'):
#             print(f"🐛 DEBUG: AnswerExtractor.selected_provider = {extractor.selected_provider}")
        
#         # Create database service once with the correct suffix
#         print(f"🐛 DEBUG: Creating StudentAnswerService with suffix: '{provider_suffix}'")
#         db_service = StudentAnswerService(provider_suffix=provider_suffix)
        
#         # Initialize the correct table (this should create student_answers_gemini for Gemini)
#         print(f"🐛 DEBUG: Initializing table for suffix: '{provider_suffix}'")
#         db_service.initialize_table()
        
#         # if from_db:
#         #     # Database mode - process all files from database
#         #     submissions = get_submissions_from_db()
#         if from_db:
#             submissions = get_submissions_from_db(
#                 assessment_id=kwargs.get("assessment_id"),
#                 selected_ids=kwargs.get("selected_submissions")
#             )

#             if not submissions:
#                 db_service.close()
#                 return {"status": "error", "message": "No submissions found in database."}
            
#             processed_count = 0
#             error_count = 0
            
#             for submission in submissions:
#                 try:
#                     file_path = submission['file_url']
#                     assessment_id = submission['assessment_id']
#                     student_id = submission['student_id']
                    
#                     print(f"\n📄 Processing: {pathlib.Path(file_path).name} (Assessment: {assessment_id}, Student: {student_id})")
                    
#                     # Resolve file path and load text
#                     full_path = resolve_file_path(file_path)
#                     raw_text = load_file_text(str(full_path))
                    
#                     # DEBUG: Confirm we're using the right provider
#                     print(f"🐛 DEBUG: About to extract with provider: {provider} (suffix: {provider_suffix})")
                    
#                     # Enhanced extraction with robust error handling
#                     answers = None
#                     max_extraction_attempts = 3
                    
#                     for attempt in range(max_extraction_attempts):
#                         try:
#                             print(f"🔄 Extraction attempt {attempt + 1}/{max_extraction_attempts}")
#                             answers = extractor.extract_answers_with_llm(raw_text)
                            
#                             if answers:
#                                 print(f"✅ Successfully extracted {len(answers)} answers on attempt {attempt + 1}")
#                                 break
#                             else:
#                                 print(f"⚠️ No answers extracted on attempt {attempt + 1}")
                                
#                         except json.JSONDecodeError as json_error:
#                             print(f"❌ JSON parsing error on attempt {attempt + 1}: {json_error}")
                            
#                             if "Invalid control character" in str(json_error):
#                                 print("🔧 Detected control character issue - this might be due to LLM response formatting")
                                
#                                 # If this is not the last attempt, wait and retry
#                                 if attempt < max_extraction_attempts - 1:
#                                     print(f"⏳ Waiting 10 seconds before retry...")
#                                     time.sleep(10)
#                                     continue
#                                 else:
#                                     print("💥 All extraction attempts failed due to JSON parsing errors")
#                                     answers = None
#                                     break
#                             else:
#                                 # For other JSON errors, re-raise immediately
#                                 raise json_error
                                
#                         except Exception as other_error:
#                             print(f"❌ Other error on attempt {attempt + 1}: {other_error}")
#                             if attempt == max_extraction_attempts - 1:
#                                 raise other_error
#                             else:
#                                 print(f"⏳ Waiting 10 seconds before retry...")
#                                 time.sleep(10)
#                                 continue

#                     if not answers:
#                         print("❌ No answers extracted after all attempts.")
#                         error_count += 1
#                         continue

#                     # Print extracted answers for verification
#                     print("\nExtracted Answers (before saving to DB):\n")
#                     pprint([
#                         {
#                             "question": ans.full_question_id,
#                             "answer": ans.answer_text
#                         }
#                         for ans in answers
#                     ])

#                     # Save answers to DB using the shared db_service
#                     first = answers[0]
                    
#                     # DEBUG: Confirm table being used
#                     table_name = f"student_answers_{provider_suffix}"
#                     print(f"🐛 DEBUG: Saving to table: {table_name}")
                    
#                     # FIX: Normalize module_code to uppercase before saving
#                     normalized_module_code = first.module_code.upper() if first.module_code else first.module_code
#                     print(f"🐛 DEBUG: Original module_code: '{first.module_code}' -> Normalized: '{normalized_module_code}'")
                    
#                     db_service.save_answers(
#                         student_index=first.student_index,
#                         module_code=normalized_module_code,  # Use normalized uppercase version
#                         year=first.exam_year,
#                         month=first.exam_month,
#                         answers=answers,
#                         assessment_id=submission["assessment_id"] 
#                     )

#                     print(f"✅ Saved answers for {first.student_index} | {normalized_module_code} | {first.exam_year}-{first.exam_month} to {table_name}")
#                     processed_count += 1
                    
#                     # Delay to respect Gemini rate limits (15 requests/min)
#                     if provider == "GoogleGemini":
#                         print("⏱️ Waiting 10 seconds for Gemini rate limit...")
#                         time.sleep(10)
                    
#                 except Exception as e:
#                     print(f"❌ Error processing {pathlib.Path(submission['file_url']).name}: {e}")
#                     error_count += 1
#                     continue
            
#             db_service.close()
            
#             return {
#                 "status": "success",
#                 "message": f"Processing complete. {processed_count} files processed successfully, {error_count} errors. Data saved to table: student_answers_{provider_suffix}",
#                 "details": {
#                     "processed": processed_count,
#                     "errors": error_count,
#                     "table": f"student_answers_{provider_suffix}"
#                 }
#             }
            
#         elif folder:
#             # Original folder/file mode
#             processed_count = 0
#             error_count = 0
            
#             if os.path.isfile(folder) and folder.endswith(".docx"):
#                 # Single file mode
#                 try:
#                     extract_and_save(folder, extractor, provider_suffix)
#                     processed_count = 1
#                 except Exception as e:
#                     error_count = 1
#                     print(f"❌ Error processing file: {e}")
                    
#             elif os.path.isdir(folder):
#                 # Folder mode
#                 for filename in os.listdir(folder):
#                     if filename.lower().endswith(".docx"):
#                         filepath = os.path.join(folder, filename)
#                         try:
#                             extract_and_save(filepath, extractor, provider_suffix)
#                             processed_count += 1
#                         except Exception as e:
#                             error_count += 1
#                             print(f"❌ Error processing {filename}: {e}")

#                         # Delay to respect Gemini rate limits (15 requests/min)
#                         if provider == "GoogleGemini":
#                             print("⏱️ Waiting 10 seconds for Gemini rate limit...")
#                             time.sleep(10)
#             else:
#                 db_service.close()
#                 return {"status": "error", "message": "Invalid folder path. Must be either a .docx file or a directory."}
            
#             db_service.close()
            
#             return {
#                 "status": "success",
#                 "message": f"Processing complete. {processed_count} files processed successfully, {error_count} errors. Data saved to table: student_answers_{provider_suffix}",
#                 "details": {
#                     "processed": processed_count,
#                     "errors": error_count,
#                     "table": f"student_answers_{provider_suffix}"
#                 }
#             }
#         else:
#             db_service.close()
#             return {"status": "error", "message": "Please specify either folder for file/folder mode or use from_db=True for database mode."}
            
#     except Exception as e:
#         print(f"🐛 DEBUG: Exception in main(): {e}")
#         import traceback
#         traceback.print_exc()
#         return {"status": "error", "message": f"Unexpected error: {str(e)}"}

# if __name__ == "__main__":
#     # Command line argument parsing
#     parser = argparse.ArgumentParser(description="Extract and save student answers using LLMs")
#     parser.add_argument("--provider", required=True, choices=["OpenAI", "GoogleGemini"], help="LLM provider")
#     parser.add_argument("--model", required=True, help="Model name (e.g., gpt-4o, gemini-2.0-flash)")
#     parser.add_argument("--folder", help="Single DOCX file or folder containing DOCX files")
#     parser.add_argument("--from-db", action="store_true", help="Process files from database instead of folder")

#     args = parser.parse_args()
    
#     print(f"🐛 DEBUG: Command line args parsed:")
#     print(f"   args.provider = '{args.provider}'")
#     print(f"   args.model = '{args.model}'")
#     print(f"   args.from_db = {args.from_db}")
#     print(f"   args.folder = {args.folder}")
    
#     # Call the main function with the PARSED arguments
#     result = main(
#         provider=args.provider,      # Use parsed provider
#         model=args.model,           # Use parsed model  
#         folder=args.folder,         # Use parsed folder
#         from_db=args.from_db        # Use parsed from_db flag
#     )
    
#     print(f"\nFinal Result: {result}")
    
#     if result["status"] == "error":
#         sys.exit(1)


# import sys
# import os
# import time
# import pathlib
# import json
# import re
# from docx import Document
# from pprint import pprint
# import psycopg2
# from psycopg2.extras import RealDictCursor
# from dotenv import load_dotenv
# import pdfplumber
# import argparse

# sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "../../")))

# from src.services.answer_extractor import AnswerExtractor
# from src.services.database_services.student_answer_db import StudentAnswerService

# load_dotenv()

# def clean_json_string(json_str):
#     """Clean JSON string by removing/escaping problematic control characters."""
#     if not json_str:
#         return json_str
    
#     cleaned = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', json_str)
#     return cleaned

# def safe_json_parse(json_str, max_retries=3):
#     """Safely parse JSON with error handling and cleaning."""
#     for attempt in range(max_retries):
#         try:
#             if attempt == 0:
#                 return json.loads(json_str)
#             elif attempt == 1:
#                 cleaned_str = clean_json_string(json_str)
#                 print(f"🔧 Attempt {attempt + 1}: Trying with cleaned JSON string")
#                 return json.loads(cleaned_str)
#             else:
#                 json_match = re.search(r'\{.*\}|\[.*\]', json_str, re.DOTALL)
#                 if json_match:
#                     json_part = json_match.group()
#                     cleaned_part = clean_json_string(json_part)
#                     print(f"🔧 Attempt {attempt + 1}: Trying with extracted and cleaned JSON")
#                     return json.loads(cleaned_part)
#                 else:
#                     raise ValueError("No valid JSON structure found")
                    
#         except json.JSONDecodeError as e:
#             print(f"❌ JSON parse attempt {attempt + 1} failed: {e}")
#             if attempt == max_retries - 1:
#                 print("📝 Problematic JSON content (first 500 chars):")
#                 print(repr(json_str[:500]))
#                 raise e
#             continue
    
#     return None

# def get_provider_suffix(provider):
#     """Convert provider name to database table suffix."""
#     print(f"🐛 DEBUG: Converting provider '{provider}' to suffix...")
    
#     provider_map = {
#         "OpenAI": "openai",
#         "GoogleGemini": "gemini",
#         "Gemini": "gemini",
#         "Google": "gemini"
#     }
    
#     suffix = provider_map.get(provider, provider.lower())
#     print(f"🐛 DEBUG: Provider '{provider}' -> Suffix '{suffix}'")
#     return suffix

# def get_database_connection():
#     """Get database connection using environment variables."""
#     try:
#         conn = psycopg2.connect(
#             host=os.getenv('POSTGRES_HOST'),
#             port=os.getenv('POSTGRES_PORT'),
#             database=os.getenv('POSTGRES_DB'),
#             user=os.getenv('POSTGRES_USER'),
#             password=os.getenv('POSTGRES_PASSWORD')
#         )
#         return conn
#     except Exception as e:
#         print(f"Failed to connect to database: {e}")
#         sys.exit(1)

# def get_submissions_from_db():
#     """Retrieve all submission file paths from database using new schema."""
#     conn = get_database_connection()
#     try:
#         with conn.cursor(cursor_factory=RealDictCursor) as cur:
#             cur.execute("""
#                 SELECT s.submission_id, s.assessment_id, s.student_id, s.file_url, s.submission_start_at,
#                        st.registration_number, a.created_on, m.module_code
#                 FROM "Submission" s
#                 JOIN "Student" st ON s.student_id = st.user_id
#                 JOIN "Assessment" a ON s.assessment_id = a.assessment_id
#                 JOIN "Module" m ON a.module_id = m.module_id
#                 ORDER BY s.assessment_id, s.submission_start_at ASC
#             """)
            
#             submissions = cur.fetchall()
#             print(f"Found {len(submissions)} submissions in database")
#             return submissions
            
#     except Exception as e:
#         print(f"Database error: {e}")
#         return []
#     finally:
#         conn.close()

# def resolve_file_path(file_path: str) -> pathlib.Path:
#     """Resolve database file path to actual file location."""
#     project_root = pathlib.Path.cwd()
#     full_path = project_root.parent / file_path
    
#     if not full_path.exists():
#         alternative_path = project_root / file_path
#         if alternative_path.exists():
#             full_path = alternative_path
#         else:
#             raise FileNotFoundError(f"File not found at: {full_path} or {alternative_path}")
    
#     return full_path

# def load_docx_text(docx_path: str) -> str:
#     doc = Document(docx_path)
#     return "\n".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])

# def load_pdf_text(pdf_path: str) -> str:
#     with pdfplumber.open(pdf_path) as pdf:
#         return "\n".join(page.extract_text() or "" for page in pdf.pages)

# def load_file_text(file_path: str) -> str:
#     """Load text from file based on extension."""
#     path_obj = pathlib.Path(file_path)
#     if path_obj.suffix.lower() == '.docx':
#         return load_docx_text(file_path)
#     elif path_obj.suffix.lower() == '.pdf':
#         return load_pdf_text(file_path)
#     else:
#         raise ValueError(f"Unsupported file format: {path_obj.suffix}")

# def process_submission(submission, extractor, provider_suffix, db_service):
#     """Process a single submission and save to database with proper mapping."""
#     try:
#         file_path = submission['file_url']
#         assessment_id = submission['assessment_id']
#         student_id = submission['student_id']
#         registration_number = submission['registration_number']
#         module_code = submission['module_code']
        
#         # Get year and month from assessment creation date
#         created_date = submission['created_on']
#         year = created_date.year
#         month = created_date.strftime('%B')
        
#         print(f"\n📄 Processing: {pathlib.Path(file_path).name}")
#         print(f"   Assessment: {assessment_id}")
#         print(f"   Student: {registration_number} (ID: {student_id})")
#         print(f"   Module: {module_code}")
#         print(f"   Date: {year}-{month}")
        
#         # Resolve file path and load text
#         full_path = resolve_file_path(file_path)
#         raw_text = load_file_text(str(full_path))
        
#         # Extract answers with retry logic
#         answers = None
#         max_extraction_attempts = 3
        
#         for attempt in range(max_extraction_attempts):
#             try:
#                 print(f"🔄 Extraction attempt {attempt + 1}/{max_extraction_attempts}")
#                 answers = extractor.extract_answers_with_llm(raw_text)
                
#                 if answers:
#                     print(f"✅ Successfully extracted {len(answers)} answers on attempt {attempt + 1}")
#                     break
#                 else:
#                     print(f"⚠️ No answers extracted on attempt {attempt + 1}")
                    
#             except json.JSONDecodeError as json_error:
#                 print(f"❌ JSON parsing error on attempt {attempt + 1}: {json_error}")
                
#                 if "Invalid control character" in str(json_error):
#                     print("🔧 Detected control character issue - this might be due to LLM response formatting")
                    
#                     if attempt < max_extraction_attempts - 1:
#                         print(f"⏳ Waiting 10 seconds before retry...")
#                         time.sleep(10)
#                         continue
#                     else:
#                         print("💥 All extraction attempts failed due to JSON parsing errors")
#                         answers = None
#                         break
#                 else:
#                     raise json_error
                    
#             except Exception as other_error:
#                 print(f"❌ Other error on attempt {attempt + 1}: {other_error}")
#                 if attempt == max_extraction_attempts - 1:
#                     raise other_error
#                 else:
#                     print(f"⏳ Waiting 10 seconds before retry...")
#                     time.sleep(10)
#                     continue

#         if not answers:
#             print("❌ No answers extracted after all attempts.")
#             return False

#         # Override extracted data with database-mapped data
#         for answer in answers:
#             answer.student_index = registration_number  # Use registration_number as student_index
#             answer.module_code = module_code  # Use module_code from database
#             answer.exam_year = year  # Use year from assessment creation date
#             answer.exam_month = month  # Use month from assessment creation date

#         print(f"🔄 Mapped data: Student {registration_number}, Module {module_code}, {year}-{month}")

#         # Print extracted answers for verification
#         print("\nExtracted Answers (after database mapping):\n")
#         pprint([
#             {
#                 "question": ans.full_question_id,
#                 "answer": ans.answer_text,
#                 "student_index": ans.student_index,
#                 "module_code": ans.module_code,
#                 "exam_year": ans.exam_year,
#                 "exam_month": ans.exam_month
#             }
#             for ans in answers[:3]  # Show first 3 for brevity
#         ])

#         # Save to database with mapped data
#         db_service.save_answers(
#             student_index=registration_number,  # Use registration_number
#             module_code=module_code,  # Use database module_code
#             year=year,  # Use assessment creation year
#             month=month,  # Use assessment creation month
#             answers=answers,
#             assessment_id=assessment_id  # Pass assessment_id for tracking
#         )

#         print(f"✅ Saved answers for {registration_number} | {module_code} | {year}-{month}")
#         return True
        
#     except Exception as e:
#         print(f"❌ Error processing {pathlib.Path(submission['file_url']).name}: {e}")
#         import traceback
#         traceback.print_exc()
#         return False

# def main(provider=None, model=None, folder=None, from_db=True, selected_submissions=None, assessment_id=None, **kwargs):
#     """
#     Main function that can be called from Flask API or command line.
    
#     Args:
#         provider: LLM provider ("OpenAI" or "GoogleGemini")
#         model: Model name (e.g., "gpt-4o", "gemini-2.0-flash")
#         folder: Single file or folder path (optional)
#         from_db: Whether to process files from database (default: True)
#         selected_submissions: List of submission IDs to process (optional)
#         assessment_id: Filter by specific assessment ID (optional)
#         **kwargs: Additional arguments from pipeline calls
    
#     Returns:
#         dict: Result status and message
#     """
#     try:
#         # Handle command line arguments if called as script
#         if provider is None or model is None:
#             if len(sys.argv) > 1:
#                 import argparse
#                 parser = argparse.ArgumentParser()
#                 parser.add_argument("--provider", required=True, choices=["OpenAI", "GoogleGemini"])
#                 parser.add_argument("--model", required=True)
#                 parser.add_argument("--folder", default=None)
#                 parser.add_argument("--from-db", action="store_true", default=True)
                
#                 args, unknown = parser.parse_known_args()
#                 provider = args.provider
#                 model = args.model
#                 folder = args.folder
#                 from_db = args.from_db
#             else:
#                 raise ValueError("Provider and model must be specified")
        
#         print(f"🐛 DEBUG: main() called with:")
#         print(f"   provider = '{provider}'")
#         print(f"   model = '{model}'")
#         print(f"   assessment_id = {assessment_id}")
#         print(f"   selected_submissions = {selected_submissions}")
        
#         provider_suffix = get_provider_suffix(provider)
#         print(f"⚙️ Using provider: {provider} -> table suffix: {provider_suffix}")
        
#         if provider_suffix not in ["openai", "gemini"]:
#             raise ValueError(f"Unsupported provider suffix: {provider_suffix}")
        
#         extractor = AnswerExtractor(selected_provider=provider, selected_model=model)
#         db_service = StudentAnswerService(provider_suffix=provider_suffix)
#         db_service.initialize_table()
        
#         if from_db:
#             # Get submissions from database
#             all_submissions = get_submissions_from_db()
            
#             if not all_submissions:
#                 db_service.close()
#                 return {"status": "error", "message": "No submissions found in database."}
            
#             # Filter by assessment_id if provided
#             if assessment_id:
#                 submissions = [s for s in all_submissions if s['assessment_id'] == assessment_id]
#                 print(f"Filtered to {len(submissions)} submissions for assessment {assessment_id}")
#             else:
#                 submissions = all_submissions
            
#             # Filter by selected submissions if provided
#             if selected_submissions:
#                 submissions = [s for s in submissions if s['submission_id'] in selected_submissions]
#                 print(f"Filtered to {len(submissions)} selected submissions")
            
#             if not submissions:
#                 db_service.close()
#                 return {"status": "error", "message": "No submissions match the filtering criteria."}
            
#             processed_count = 0
#             error_count = 0
            
#             for submission in submissions:
#                 success = process_submission(submission, extractor, provider_suffix, db_service)
#                 if success:
#                     processed_count += 1
#                 else:
#                     error_count += 1
                
#                 # Rate limiting for Gemini
#                 if provider == "GoogleGemini":
#                     print("⏱️ Waiting 10 seconds for Gemini rate limit...")
#                     time.sleep(10)
            
#             db_service.close()
            
#             return {
#                 "status": "success",
#                 "message": f"Processing complete. {processed_count} files processed successfully, {error_count} errors.",
#                 "details": {
#                     "processed": processed_count,
#                     "errors": error_count,
#                     "table": f"student_answers_{provider_suffix}",
#                     "assessment_id": assessment_id,
#                     "selected_submissions": len(selected_submissions) if selected_submissions else "all"
#                 }
#             }
            
#         elif folder:
#             # Original folder/file mode (unchanged)
#             processed_count = 0
#             error_count = 0
            
#             if os.path.isfile(folder) and folder.endswith(".docx"):
#                 # Single file processing would go here
#                 pass
#             elif os.path.isdir(folder):
#                 # Directory processing would go here
#                 pass
#             else:
#                 db_service.close()
#                 return {"status": "error", "message": "Invalid folder path."}
            
#             db_service.close()
#             return {
#                 "status": "success",
#                 "message": f"Processing complete. {processed_count} files processed successfully, {error_count} errors.",
#                 "details": {"processed": processed_count, "errors": error_count}
#             }
#         else:
#             db_service.close()
#             return {"status": "error", "message": "Please specify either folder or use from_db=True."}
            
#     except Exception as e:
#         print(f"🐛 DEBUG: Exception in main(): {e}")
#         import traceback
#         traceback.print_exc()
#         return {"status": "error", "message": f"Unexpected error: {str(e)}"}

# if __name__ == "__main__":
#     parser = argparse.ArgumentParser(description="Extract and save student answers using LLMs")
#     parser.add_argument("--provider", required=True, choices=["OpenAI", "GoogleGemini"], help="LLM provider")
#     parser.add_argument("--model", required=True, help="Model name")
#     parser.add_argument("--folder", help="Single DOCX file or folder containing DOCX files")
#     parser.add_argument("--from-db", action="store_true", help="Process files from database")

#     args = parser.parse_args()
    
#     result = main(
#         provider=args.provider,
#         model=args.model,
#         folder=args.folder,
#         from_db=args.from_db
#     )
    
#     print(f"\nFinal Result: {result}")
    
#     if result["status"] == "error":
#         sys.exit(1)


import sys
import os
import time
import pathlib
import json
import re
from docx import Document
from pprint import pprint
import psycopg2
from psycopg2.extras import RealDictCursor
from dotenv import load_dotenv
import pdfplumber
import argparse

sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "../../")))

from src.services.answer_extractor import AnswerExtractor
from src.services.database_services.student_answer_db import StudentAnswerService

load_dotenv()

def clean_json_string(json_str):
    """Clean JSON string by removing/escaping problematic control characters."""
    if not json_str:
        return json_str
    
    # Remove or replace common problematic control characters
    # Keep only allowed whitespace characters (space, tab, newline, carriage return)
    cleaned = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', json_str)
    
    # Additional cleaning - escape unescaped quotes within strings
    # This is a basic approach - you might need more sophisticated handling
    return cleaned

def safe_json_parse(json_str, max_retries=3):
    """Safely parse JSON with error handling and cleaning."""
    for attempt in range(max_retries):
        try:
            # First attempt - try parsing as-is
            if attempt == 0:
                return json.loads(json_str)
            
            # Second attempt - clean the string
            elif attempt == 1:
                cleaned_str = clean_json_string(json_str)
                print(f"🔧 Attempt {attempt + 1}: Trying with cleaned JSON string")
                return json.loads(cleaned_str)
            
            # Third attempt - more aggressive cleaning
            else:
                # Try to extract just the JSON part if it's embedded in other text
                json_match = re.search(r'\{.*\}|\[.*\]', json_str, re.DOTALL)
                if json_match:
                    json_part = json_match.group()
                    cleaned_part = clean_json_string(json_part)
                    print(f"🔧 Attempt {attempt + 1}: Trying with extracted and cleaned JSON")
                    return json.loads(cleaned_part)
                else:
                    raise ValueError("No valid JSON structure found")
                    
        except json.JSONDecodeError as e:
            print(f"❌ JSON parse attempt {attempt + 1} failed: {e}")
            if attempt == max_retries - 1:
                print("📝 Problematic JSON content (first 500 chars):")
                print(repr(json_str[:500]))
                raise e
            continue
    
    return None

def get_provider_suffix(provider):
    """Convert provider name to database table suffix."""
    print(f"🐛 DEBUG: Converting provider '{provider}' to suffix...")
    
    # Standardize the mapping to ensure consistent table naming
    provider_map = {
        "OpenAI": "openai",
        "GoogleGemini": "gemini",
        "Gemini": "gemini",  # Alternative name
        "Google": "gemini"   # Another alternative
    }
    
    suffix = provider_map.get(provider, provider.lower())
    print(f"🐛 DEBUG: Provider '{provider}' -> Suffix '{suffix}'")
    return suffix

def get_database_connection():
    """Get database connection using environment variables."""
    try:
        conn = psycopg2.connect(
            host=os.getenv('POSTGRES_HOST'),
            port=os.getenv('POSTGRES_PORT'),
            database=os.getenv('POSTGRES_DB'),
            user=os.getenv('POSTGRES_USER'),
            password=os.getenv('POSTGRES_PASSWORD')
        )
        return conn
    except Exception as e:
        print(f"Failed to connect to database: {e}")
        sys.exit(1)

def get_selected_submissions_from_db(selected_submission_ids=None, assessment_id=None):
    """
    Retrieve specific submissions from database with full mapping data.
    Filter by submission IDs and/or assessment ID.
    """
    conn = get_database_connection()
    try:
        with conn.cursor(cursor_factory=RealDictCursor) as cur:
            if selected_submission_ids and assessment_id:
                # Filter by both submission IDs and assessment ID
                submission_ids_tuple = tuple(selected_submission_ids)
                placeholders = ','.join(['%s'] * len(submission_ids_tuple))
                
                cur.execute(f"""
                    SELECT s.submission_id, s.assessment_id, s.student_id, s.file_url, 
                           s.submission_start_at,
                           st.registration_number, st.user_id as student_user_id,
                           a.assessment_id, a.created_on as assessment_created_on,
                           m.module_code, m.module_name
                    FROM "Submission" s
                    JOIN "Student" st ON s.student_id = st.user_id
                    JOIN "Assessment" a ON s.assessment_id = a.assessment_id
                    JOIN "Module" m ON a.module_id = m.module_id
                    WHERE s.submission_id IN ({placeholders}) AND s.assessment_id = %s
                    ORDER BY s.submission_start_at ASC
                """, submission_ids_tuple + (assessment_id,))
                
            elif selected_submission_ids:
                # Filter by submission IDs only
                submission_ids_tuple = tuple(selected_submission_ids)
                placeholders = ','.join(['%s'] * len(submission_ids_tuple))
                
                cur.execute(f"""
                    SELECT s.submission_id, s.assessment_id, s.student_id, s.file_url, 
                           s.submission_start_at,
                           st.registration_number, st.user_id as student_user_id,
                           a.assessment_id, a.created_on as assessment_created_on,
                           m.module_code, m.module_name
                    FROM "Submission" s
                    JOIN "Student" st ON s.student_id = st.user_id
                    JOIN "Assessment" a ON s.assessment_id = a.assessment_id
                    JOIN "Module" m ON a.module_id = m.module_id
                    WHERE s.submission_id IN ({placeholders})
                    ORDER BY s.submission_start_at ASC
                """, submission_ids_tuple)
                
            elif assessment_id:
                # Filter by assessment ID only
                cur.execute("""
                    SELECT s.submission_id, s.assessment_id, s.student_id, s.file_url, 
                           s.submission_start_at,
                           st.registration_number, st.user_id as student_user_id,
                           a.assessment_id, a.created_on as assessment_created_on,
                           m.module_code, m.module_name
                    FROM "Submission" s
                    JOIN "Student" st ON s.student_id = st.user_id
                    JOIN "Assessment" a ON s.assessment_id = a.assessment_id
                    JOIN "Module" m ON a.module_id = m.module_id
                    WHERE s.assessment_id = %s
                    ORDER BY s.submission_start_at ASC
                """, (assessment_id,))
            else:
                # Get all submissions (fallback)
                cur.execute("""
                    SELECT s.submission_id, s.assessment_id, s.student_id, s.file_url, 
                           s.submission_start_at,
                           st.registration_number, st.user_id as student_user_id,
                           a.assessment_id, a.created_on as assessment_created_on,
                           m.module_code, m.module_name
                    FROM "Submission" s
                    JOIN "Student" st ON s.student_id = st.user_id
                    JOIN "Assessment" a ON s.assessment_id = a.assessment_id
                    JOIN "Module" m ON a.module_id = m.module_id
                    ORDER BY s.submission_start_at ASC
                """)
            
            submissions = cur.fetchall()
            print(f"Found {len(submissions)} submissions matching criteria")
            return [dict(sub) for sub in submissions]
            
    except Exception as e:
        print(f"Database error: {e}")
        return []
    finally:
        conn.close()

def check_submission_already_extracted(submission_id, provider_suffix):
    """Check if a specific submission has already been extracted."""
    conn = get_database_connection()
    try:
        with conn.cursor() as cur:
            answers_table = f"student_answers_{provider_suffix}"
            
            # Check if the table exists first
            cur.execute("""
                SELECT EXISTS (
                    SELECT FROM information_schema.tables 
                    WHERE table_name = %s
                );
            """, (answers_table,))
            
            table_exists = cur.fetchone()[0]
            if not table_exists:
                return False
            
            # Check if submission_id column exists
            cur.execute("""
                SELECT EXISTS (
                    SELECT FROM information_schema.columns 
                    WHERE table_name = %s AND column_name = 'submission_id'
                );
            """, (answers_table,))
            
            column_exists = cur.fetchone()[0]
            if not column_exists:
                return False
            
            # Check if this submission has been processed
            cur.execute(f"""
                SELECT EXISTS (
                    SELECT 1 FROM "{answers_table}" 
                    WHERE submission_id = %s
                );
            """, (submission_id,))
            
            return cur.fetchone()[0]
            
    except Exception as e:
        print(f"Error checking extraction status for submission {submission_id}: {e}")
        return False
    finally:
        conn.close()

def resolve_file_path(file_path: str) -> pathlib.Path:
    """Resolve database file path to actual file location."""
    project_root = pathlib.Path.cwd()
    # Try parent directory first (where data folder should be)
    full_path = project_root.parent / file_path
    
    if not full_path.exists():
        # Try alternative path in project directory
        alternative_path = project_root / file_path
        if alternative_path.exists():
            full_path = alternative_path
        else:
            raise FileNotFoundError(f"File not found at: {full_path} or {alternative_path}")
    
    return full_path

def load_docx_text(docx_path: str) -> str:
    doc = Document(docx_path)
    return "\n".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])

def load_pdf_text(pdf_path: str) -> str:
    with pdfplumber.open(pdf_path) as pdf:
        return "\n".join(page.extract_text() or "" for page in pdf.pages)

def load_file_text(file_path: str) -> str:
    """Load text from file based on extension."""
    path_obj = pathlib.Path(file_path)
    if path_obj.suffix.lower() == '.docx':
        return load_docx_text(file_path)
    elif path_obj.suffix.lower() == '.pdf':
        return load_pdf_text(file_path)
    else:
        raise ValueError(f"Unsupported file format: {path_obj.suffix}")

def extract_and_save_with_db_mapping(submission_data: dict, extractor: AnswerExtractor, provider_suffix: str):
    """
    Extract and save answers using database-mapped data instead of file-extracted metadata.
    """
    submission_id = submission_data['submission_id']
    file_path = submission_data['file_url']
    student_registration_number = submission_data['registration_number']
    module_code = submission_data['module_code']
    assessment_created_on = submission_data['assessment_created_on']
    assessment_id = submission_data['assessment_id']
    
    filename = os.path.basename(file_path)
    print(f"\n📄 Processing: {filename}")
    print(f"🔗 Submission ID: {submission_id}")
    print(f"👤 Student: {student_registration_number}")
    print(f"📚 Module: {module_code}")
    print(f"🆔 Assessment: {assessment_id}")

    try:
        # Check if already extracted
        if check_submission_already_extracted(submission_id, provider_suffix):
            print(f"⚠️ Submission {submission_id} already extracted, skipping...")
            return True

        # Load file content
        full_path = resolve_file_path(file_path)
        raw_text = load_file_text(str(full_path))
        
        # Enhanced extraction with robust error handling
        answers = None
        max_extraction_attempts = 3
        
        for attempt in range(max_extraction_attempts):
            try:
                print(f"🔄 Extraction attempt {attempt + 1}/{max_extraction_attempts}")
                answers = extractor.extract_answers_with_llm(raw_text)
                
                if answers:
                    print(f"✅ Successfully extracted {len(answers)} answers on attempt {attempt + 1}")
                    break
                else:
                    print(f"⚠️ No answers extracted on attempt {attempt + 1}")
                    
            except json.JSONDecodeError as json_error:
                print(f"❌ JSON parsing error on attempt {attempt + 1}: {json_error}")
                
                if "Invalid control character" in str(json_error):
                    print("🔧 Detected control character issue - this might be due to LLM response formatting")
                    
                    # If this is not the last attempt, wait and retry
                    if attempt < max_extraction_attempts - 1:
                        print(f"⏳ Waiting 10 seconds before retry...")
                        time.sleep(10)
                        continue
                    else:
                        print("💥 All extraction attempts failed due to JSON parsing errors")
                        answers = None
                        break
                else:
                    # For other JSON errors, re-raise immediately
                    raise json_error
                    
            except Exception as other_error:
                print(f"❌ Other error on attempt {attempt + 1}: {other_error}")
                if attempt == max_extraction_attempts - 1:
                    raise other_error
                else:
                    print(f"⏳ Waiting 10 seconds before retry...")
                    time.sleep(10)
                    continue

        if not answers:
            print("❌ No answers extracted after all attempts.")
            return False

        # Override extracted metadata with database-mapped data
        year = assessment_created_on.year
        month = assessment_created_on.strftime('%B')
        
        print(f"🔄 Overriding extracted metadata with database values:")
        print(f"   Student Index: {student_registration_number}")
        print(f"   Module Code: {module_code}")
        print(f"   Year: {year}")
        print(f"   Month: {month}")
        
        # Update all answers with database-mapped metadata
        for answer in answers:
            answer.student_index = student_registration_number
            answer.module_code = module_code.upper()  # Normalize to uppercase
            answer.exam_year = year
            answer.exam_month = month

        # Preview the result with database-mapped data
        print("\nAnswers with database-mapped metadata:")
        pprint([
            {
                "question": ans.full_question_id, 
                "answer": ans.answer_text[:100] + "..." if len(ans.answer_text) > 100 else ans.answer_text,
                "student_index": ans.student_index,
                "module_code": ans.module_code,
                "year": ans.exam_year,
                "month": ans.exam_month
            }
            for ans in answers[:3]  # Show first 3 for preview
        ])

        # Save to database with enhanced tracking
        db = StudentAnswerService(provider_suffix=provider_suffix)
        db.initialize_table()
        
        # Save answers with submission tracking
        db.save_answers_with_submission_tracking(
            student_index=student_registration_number,
            module_code=module_code.upper(),
            year=year,
            month=month,
            answers=answers,
            submission_id=submission_id,
            assessment_id=assessment_id
        )
        db.close()

        print(f"✅ Saved answers for {student_registration_number} | {module_code} | {year}-{month} | Submission: {submission_id}")
        return True
        
    except Exception as e:
        print(f"❌ Failed to process {filename}: {e}")
        import traceback
        traceback.print_exc()
        return False

def main(provider=None, model=None, folder=None, from_db=True, selected_submission_ids=None, 
         assessment_id=None, submissions_data=None, **kwargs):
    """
    Enhanced main function with database mapping and submission filtering.
    
    Args:
        provider: LLM provider ("OpenAI" or "GoogleGemini")
        model: Model name (e.g., "gpt-4o", "gemini-2.0-flash")
        folder: Single file or folder path (optional, for backward compatibility)
        from_db: Whether to process files from database (default: True)
        selected_submission_ids: List of specific submission IDs to process
        assessment_id: Assessment ID for filtering
        submissions_data: Pre-fetched submission data from Flask API
        **kwargs: Additional arguments from pipeline calls
    
    Returns:
        dict: Result status and message
    """
    try:
        # Handle case where provider/model might be None (pipeline calls)
        if provider is None or model is None:
            # Try to get from sys.argv if called as script
            if len(sys.argv) > 1:
                parser = argparse.ArgumentParser()
                parser.add_argument("--provider", required=True, choices=["OpenAI", "GoogleGemini"])
                parser.add_argument("--model", required=True)
                parser.add_argument("--folder", default=None)
                parser.add_argument("--from-db", action="store_true", default=True)
                parser.add_argument("--selected-submission-ids", nargs="*", default=None)
                parser.add_argument("--assessment-id", default=None)
                
                # Parse known args to handle cases where script is called with extra args
                args, unknown = parser.parse_known_args()
                provider = args.provider
                model = args.model
                folder = args.folder
                from_db = args.from_db
                selected_submission_ids = args.selected_submission_ids
                assessment_id = args.assessment_id
            else:
                raise ValueError("Provider and model must be specified either as arguments or command line parameters")
        
        # DEBUG: Print received parameters
        print(f"🐛 DEBUG: main() called with:")
        print(f"   provider = '{provider}' (type: {type(provider)})")
        print(f"   model = '{model}'")
        print(f"   selected_submission_ids = {selected_submission_ids}")
        print(f"   assessment_id = {assessment_id}")
        print(f"   from_db = {from_db}")
        
        # Convert provider to proper suffix at the start and validate
        provider_suffix = get_provider_suffix(provider)
        print(f"⚙️ Using provider: {provider} -> table suffix: {provider_suffix}")
        
        # Validate the provider suffix
        if provider_suffix not in ["openai", "gemini"]:
            raise ValueError(f"Unsupported provider suffix: {provider_suffix}. Expected 'openai' or 'gemini'")
        
        # Create extractor with explicit logging
        print(f"🐛 DEBUG: Creating AnswerExtractor(selected_provider='{provider}', selected_model='{model}')")
        extractor = AnswerExtractor(selected_provider=provider, selected_model=model)
        
        if from_db:
            # Database mode - process selected submissions with database mapping
            
            # Use pre-fetched data if available (from Flask API)
            if submissions_data:
                submissions = submissions_data
                print(f"📥 Using pre-fetched submissions data: {len(submissions)} submissions")
            else:
                # Fetch submissions from database based on criteria
                submissions = get_selected_submissions_from_db(
                    selected_submission_ids=selected_submission_ids,
                    assessment_id=assessment_id
                )
            
            if not submissions:
                return {
                    "status": "success", 
                    "message": "No submissions found matching criteria.",
                    "details": {"processed": 0, "errors": 0, "skipped": 0}
                }
            
            processed_count = 0
            error_count = 0
            skipped_count = 0
            
            for submission in submissions:
                try:
                    # Process with database mapping
                    success = extract_and_save_with_db_mapping(submission, extractor, provider_suffix)
                    
                    if success is True:
                        processed_count += 1
                    elif success is False:
                        error_count += 1
                    else:
                        skipped_count += 1  # Already processed
                    
                    # Delay to respect API rate limits
                    if provider == "GoogleGemini":
                        print("⏱️ Waiting 10 seconds for Gemini rate limit...")
                        time.sleep(10)
                    elif provider == "OpenAI":
                        time.sleep(2)  # Shorter delay for OpenAI
                    
                except Exception as e:
                    print(f"❌ Error processing submission {submission['submission_id']}: {e}")
                    error_count += 1
                    continue
            
            return {
                "status": "success",
                "message": f"Processing complete. {processed_count} files processed, {skipped_count} skipped (already processed), {error_count} errors. Data saved to table: student_answers_{provider_suffix}",
                "details": {
                    "processed": processed_count,
                    "errors": error_count,
                    "skipped": skipped_count,
                    "table": f"student_answers_{provider_suffix}"
                }
            }
            
        elif folder:
            # Legacy folder mode for backward compatibility
            print("⚠️ Using legacy folder mode - database mapping not available")
            processed_count = 0
            error_count = 0
            
            # Create database service for legacy mode
            db_service = StudentAnswerService(provider_suffix=provider_suffix)
            db_service.initialize_table()
            
            if os.path.isfile(folder) and folder.endswith(".docx"):
                # Single file mode
                try:
                    raw_text = load_docx_text(folder)
                    answers = extractor.extract_answers_with_llm(raw_text)
                    if answers:
                        first = answers[0]
                        db_service.save_answers(
                            student_index=first.student_index,
                            module_code=first.module_code.upper() if first.module_code else None,
                            year=first.exam_year,
                            month=first.exam_month,
                            answers=answers
                        )
                        processed_count = 1
                    else:
                        error_count = 1
                except Exception as e:
                    error_count = 1
                    print(f"❌ Error processing file: {e}")
                    
            elif os.path.isdir(folder):
                # Folder mode
                for filename in os.listdir(folder):
                    if filename.lower().endswith(".docx"):
                        filepath = os.path.join(folder, filename)
                        try:
                            raw_text = load_docx_text(filepath)
                            answers = extractor.extract_answers_with_llm(raw_text)
                            if answers:
                                first = answers[0]
                                db_service.save_answers(
                                    student_index=first.student_index,
                                    module_code=first.module_code.upper() if first.module_code else None,
                                    year=first.exam_year,
                                    month=first.exam_month,
                                    answers=answers
                                )
                                processed_count += 1
                            else:
                                error_count += 1
                        except Exception as e:
                            error_count += 1
                            print(f"❌ Error processing {filename}: {e}")

                        # Delay to respect rate limits
                        if provider == "GoogleGemini":
                            print("⏱️ Waiting 10 seconds for Gemini rate limit...")
                            time.sleep(10)
            else:
                db_service.close()
                return {"status": "error", "message": "Invalid folder path. Must be either a .docx file or a directory."}
            
            db_service.close()
            
            return {
                "status": "success",
                "message": f"Legacy mode processing complete. {processed_count} files processed, {error_count} errors. Data saved to table: student_answers_{provider_suffix}",
                "details": {
                    "processed": processed_count,
                    "errors": error_count,
                    "skipped": 0,
                    "table": f"student_answers_{provider_suffix}"
                }
            }
        else:
            return {"status": "error", "message": "Please specify either folder for file/folder mode or use from_db=True for database mode."}
            
    except Exception as e:
        print(f"🐛 DEBUG: Exception in main(): {e}")
        import traceback
        traceback.print_exc()
        return {"status": "error", "message": f"Unexpected error: {str(e)}"}

if __name__ == "__main__":
    # Command line argument parsing
    parser = argparse.ArgumentParser(description="Extract and save student answers using LLMs with database mapping")
    parser.add_argument("--provider", required=True, choices=["OpenAI", "GoogleGemini"], help="LLM provider")
    parser.add_argument("--model", required=True, help="Model name (e.g., gpt-4o, gemini-2.0-flash)")
    parser.add_argument("--folder", help="Single DOCX file or folder containing DOCX files (legacy mode)")
    parser.add_argument("--from-db", action="store_true", help="Process files from database with proper mapping")
    parser.add_argument("--selected-submission-ids", nargs="*", help="Specific submission IDs to process")
    parser.add_argument("--assessment-id", help="Assessment ID for filtering submissions")

    args = parser.parse_args()
    
    print(f"🐛 DEBUG: Command line args parsed:")
    print(f"   args.provider = '{args.provider}'")
    print(f"   args.model = '{args.model}'")
    print(f"   args.from_db = {args.from_db}")
    print(f"   args.selected_submission_ids = {args.selected_submission_ids}")
    print(f"   args.assessment_id = {args.assessment_id}")
    
    # Call the main function with the PARSED arguments
    result = main(
        provider=args.provider,
        model=args.model,
        folder=args.folder,
        from_db=args.from_db,
        selected_submission_ids=args.selected_submission_ids,
        assessment_id=args.assessment_id
    )
    
    print(f"\nFinal Result: {result}")
    
    if result["status"] == "error":
        sys.exit(1)