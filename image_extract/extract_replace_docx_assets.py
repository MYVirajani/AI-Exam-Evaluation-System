import os
from io import BytesIO
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import qn

# For DOCX → PDF conversion
from docx2pdf import convert  # pip install docx2pdf

# -----------------------------
# CONFIGURATION
# -----------------------------
INPUT_DIR = "content"  # Folder with .docx files
EXTRACT_DIR = os.path.join(INPUT_DIR, "extracted")
OUTPUT_DIR = os.path.join(INPUT_DIR, "updated_docs")
PDF_DIR = os.path.join(INPUT_DIR, "pdfs")

os.makedirs(EXTRACT_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)
os.makedirs(PDF_DIR, exist_ok=True)


# -----------------------------
# HELPER FUNCTIONS
# -----------------------------
def save_image(image_blob, doc_basename, counter):
    """Save image bytes to file and return relative path."""
    image_name = f"{doc_basename}_img_{counter}.png"
    image_path = os.path.join(EXTRACT_DIR, image_name)
    image = Image.open(BytesIO(image_blob))
    image.save(image_path)
    return image_path


def save_table_as_image(table, doc_basename, counter):
    """Save table content as an image with text (simple visual snapshot)."""
    rows = []
    for row in table.rows:
        cols = [cell.text.strip() for cell in row.cells]
        rows.append(" | ".join(cols))
    text = "\n".join(rows) or "(Empty Table)"

    img_name = f"{doc_basename}_tbl_{counter}.png"
    img_path = os.path.join(EXTRACT_DIR, img_name)

    font = ImageFont.load_default()
    lines = text.splitlines()
    width = int(max(font.getlength(line) for line in lines) + 20)
    height = int(15 * len(lines) + 20)
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    y = 10
    for line in lines:
        draw.text((10, y), line, fill="black", font=font)
        y += 15
    img.save(img_path)
    return img_path


# -----------------------------
# CORE PROCESSING FUNCTION
# -----------------------------
def process_docx(docx_path):
    """Extract images & tables in correct order, replace with URLs, and convert to PDF."""
    doc = Document(docx_path)
    basename = os.path.splitext(os.path.basename(docx_path))[0]
    print(f"\n📄 Processing: {basename}.docx")

    img_counter = 0
    tbl_counter = 0

    # Iterate through all document elements in reading order
    body_elements = list(doc.element.body)

    for element in body_elements:
        tag = element.tag
        if tag.endswith("tbl"):  # Table
            tbl_counter += 1
            for tbl in doc.tables:
                if tbl._element is element:
                    tbl_path = save_table_as_image(tbl, basename, tbl_counter)
                    parent = element.getparent()
                    idx = parent.index(element)
                    placeholder = parse_xml(
                        f'<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                        f'<w:r><w:t>[Table: {tbl_path}]</w:t></w:r></w:p>'
                    )
                    parent.insert(idx + 1, placeholder)
                    parent.remove(element)
                    break

        elif tag.endswith("p"):  # Paragraph (might contain image)
            paragraph = None
            for p in doc.paragraphs:
                if p._element is element:
                    paragraph = p
                    break
            if paragraph:
                for run in paragraph.runs:
                    blips = run.element.xpath(".//a:blip")
                    if blips:
                        img_counter += 1
                        embed_rid = blips[0].get(qn("r:embed"))
                        image_part = doc.part.related_parts[embed_rid]
                        image_data = image_part.blob
                        img_path = save_image(image_data, basename, img_counter)
                        run.text = f"[Image: {img_path}]"
                        for blip in blips:
                            blip.getparent().remove(blip)

    # Save updated document
    updated_path = os.path.join(OUTPUT_DIR, f"updated_{basename}.docx")
    doc.save(updated_path)
    print(f"✅ Saved updated Word: {updated_path}")
    print(f"   Extracted {img_counter} images, {tbl_counter} tables.")

    # Convert updated Word to PDF
    pdf_path = os.path.join(PDF_DIR, f"{basename}.pdf")
    try:
        convert(updated_path, pdf_path)
        print(f"📘 Converted to PDF: {pdf_path}")
    except Exception as e:
        print(f"⚠️ PDF conversion failed for {basename}: {e}")

    return updated_path, pdf_path


# -----------------------------
# MAIN SCRIPT
# -----------------------------
def main():
    docx_files = [f for f in os.listdir(INPUT_DIR) if f.lower().endswith(".docx")]
    if not docx_files:
        print("⚠️ No .docx files found in 'content/' folder.")
        return

    for file in docx_files:
        process_docx(os.path.join(INPUT_DIR, file))

    print("\n🎉 All documents processed successfully!")
    print(f"📂 Extracted files in: {EXTRACT_DIR}")
    print(f"📄 Updated .docx files saved in: {OUTPUT_DIR}")
    print(f"📘 PDFs saved in: {PDF_DIR}")


if __name__ == "__main__":
    main()
