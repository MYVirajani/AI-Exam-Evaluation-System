

<<<<<<< HEAD
# # # import argparse
# # # import logging
# # # import pathlib
# # # from typing import Iterable
# # # from docx import Document
# # # import pdfplumber
# # # from rich import print

# # # from src.services.model_answer_extractor import ModelAnswerExtractor
# # # from src.services.embedding.openai_embedder import OpenAIEmbedder
# # # from src.services.embedding.gemini_embedder import GeminiEmbedder
# # # from src.services.database_services.model_answer_embedding_db import ModelAnswerEmbeddingDB

# # # logging.basicConfig(level=logging.INFO)
# # # logger = logging.getLogger(__name__)

# # # def read_text(path: pathlib.Path) -> str:
# # #     if path.suffix.lower() == ".docx":
# # #         doc = Document(path)
# # #         return "\n".join(p.text for p in doc.paragraphs)
# # #     elif path.suffix.lower() == ".pdf":
# # #         with pdfplumber.open(path) as pdf:
# # #             return "\n".join(p.extract_text() or "" for p in pdf.pages)
# # #     raise ValueError(f"Unsupported file type: {path.name}")

# # # def iter_files(root: pathlib.Path, patterns: Iterable[str]) -> Iterable[pathlib.Path]:
# # #     for file in root.rglob("*"):
# # #         if file.suffix.lower() in patterns and file.is_file():
# # #             yield file

# # # def main() -> None:
# # #     ap = argparse.ArgumentParser(description="Extract & embed model answers")
# # #     ap.add_argument("--provider", required=True, choices=["OpenAI", "GoogleGemini"])
# # #     ap.add_argument("--model", required=True, help="LLM model name (e.g. gpt-4o, gemini-1.5-pro)")
# # #     ap.add_argument("--embedder", default="text-embedding-3-small", help="Embedding model name")
# # #     ap.add_argument("--root", default="data/Model_Answers", help="Root folder of model answers")
# # #     ap.add_argument("--ext", nargs="*", default=[".pdf", ".docx"], help="File extensions (default: .pdf .docx)")
# # #     args = ap.parse_args()

# # #     root = pathlib.Path(args.root)
# # #     if not root.exists():
# # #         logger.error("Root folder does not exist: %s", root)
# # #         return

# # #     print(f"[bold]⏳ Scanning [cyan]{root}[/] …[/]")

# # #     extractor = ModelAnswerExtractor(args.provider, args.model)

# # #     # Choose embedder
# # #     embedder = OpenAIEmbedder(args.embedder) if args.provider == "OpenAI" else GeminiEmbedder(model_name=args.embedder)

# # #     # Vector DB
# # #     vec_db = ModelAnswerEmbeddingDB(embedder)
# # #     file_patterns = {e.lower() if e.startswith(".") else f".{e.lower()}" for e in args.ext}

# # #     processed = 0
# # #     for file in iter_files(root, file_patterns):
# # #         processed += 1
# # #         print(f"→ {file.relative_to(root)}")
# # #         raw_text = read_text(file)

# # #         answers = extractor.extract(raw_text)
# # #         vec_db.save_embeddings(answers)

# # #     vec_db.close()
# # #     msg = "No model-answer files found." if processed == 0 else f"✅ Done. {processed} file(s) processed."
# # #     print(f"[green]{msg}[/]")

# # # if __name__ == "__main__":
# # #     main()


# # """
# # Run from the project root:

# #   python -m src.scripts.embed_model_answers \
# #          --provider OpenAI \
# #          --model gpt-4o \
# #          --embedder text-embedding-3-small
# # """

# # import argparse
# # import logging
# # import pathlib
# # import sys
# # import os
# # from typing import Iterable, List, Dict, Any

# # from docx import Document
# # import pdfplumber
# # from rich import print
# # import psycopg2
# # from psycopg2.extras import RealDictCursor
# # from dotenv import load_dotenv

# # # project imports
# # from src.services.model_answer_extractor import ModelAnswerExtractor
# # from src.services.embedding.openai_embedder import OpenAIEmbedder
# # from src.services.embedding.gemini_embedder import GeminiEmbedder
# # from src.services.database_services.model_answer_embedding_db import (
# #     ModelAnswerEmbeddingDB,
# # )

# # # Load environment variables
# # load_dotenv()

# # logging.basicConfig(level=logging.INFO)
# # logger = logging.getLogger(__name__)


# # # --------------------------------------------------------------------------- #
# # #                            Database helpers                                 #
# # # --------------------------------------------------------------------------- #
# # def get_database_connection():
# #     """Get database connection using environment variables."""
# #     try:
# #         conn = psycopg2.connect(
# #             host=os.getenv('POSTGRES_HOST'),
# #             port=os.getenv('POSTGRES_PORT'),
# #             database=os.getenv('POSTGRES_DB'),
# #             user=os.getenv('POSTGRES_USER'),
# #             password=os.getenv('POSTGRES_PASSWORD')
# #         )
# #         return conn
# #     except Exception as e:
# #         logger.error(f"Failed to connect to database: {e}")
# #         sys.exit(1)


# # def get_model_answer_files_from_db() -> List[Dict[str, Any]]:
# #     """Retrieve all model answer paper file paths from database."""
# #     conn = get_database_connection()
# #     try:
# #         with conn.cursor(cursor_factory=RealDictCursor) as cur:
# #             cur.execute("""
# #                 SELECT model_answer_paper_id, assessment_id, file_url, created_on 
# #                 FROM "Model_Answer_Paper" 
# #                 ORDER BY created_on ASC
# #             """)
            
# #             files = cur.fetchall()
# #             logger.info(f"Found {len(files)} model answer files in database")
# #             return files
            
# #     except Exception as e:
# #         logger.error(f"Database error: {e}")
# #         return []
# #     finally:
# #         conn.close()


# # def get_project_root() -> pathlib.Path:
# #     """Get the project root directory (current working directory)."""
# #     return pathlib.Path.cwd()


# # def resolve_file_path(file_path: str) -> pathlib.Path:
# #     """Resolve database file path to actual file location."""
# #     project_root = get_project_root()
# #     # Try parent directory first (where data folder should be)
# #     full_path = project_root.parent / file_path
    
# #     logger.info(f"Looking for file at: {full_path}")
    
# #     if not full_path.exists():
# #         # Try alternative path in project directory
# #         alternative_path = project_root / file_path
# #         logger.info(f"Alternative path: {alternative_path}")
        
# #         if alternative_path.exists():
# #             full_path = alternative_path
# #         else:
# #             # List what's actually in the expected directory
# #             expected_dir = full_path.parent
# #             if expected_dir.exists():
# #                 files_in_dir = list(expected_dir.iterdir())
# #                 logger.error(f"Expected directory exists but file not found. Contents: {[f.name for f in files_in_dir]}")
# #             else:
# #                 logger.error(f"Expected directory does not exist: {expected_dir}")
# #             raise FileNotFoundError(f"File not found at: {full_path} or {alternative_path}")
    
# #     logger.info(f"Successfully found file at: {full_path}")
# #     return full_path


# # # --------------------------------------------------------------------------- #
# # #                              File helpers                                   #
# # # --------------------------------------------------------------------------- #
# # def read_text(path: pathlib.Path) -> str:
# #     """Return plain text from .docx or .pdf (raise on others)."""
# #     suffix = path.suffix.lower()

# #     if suffix == ".docx":
# #         doc = Document(path)
# #         return "\n".join(p.text for p in doc.paragraphs)

# #     if suffix == ".pdf":
# #         with pdfplumber.open(path) as pdf:
# #             return "\n".join(p.extract_text() or "" for p in pdf.pages)

# #     raise ValueError(f"Unsupported file type: {path.name}")


# # def iter_files(root: pathlib.Path, patterns: Iterable[str]) -> Iterable[pathlib.Path]:
# #     """Yield files under *root* whose suffix matches one of *patterns*."""
# #     for file in root.rglob("*"):
# #         if file.suffix.lower() in patterns and file.is_file():
# #             yield file


# # # --------------------------------------------------------------------------- #
# # #                               Main routine                                  #
# # # --------------------------------------------------------------------------- #
# # def main() -> None:
# #     ap = argparse.ArgumentParser(description="Extract & embed model answers")
# #     ap.add_argument("--provider", required=True, choices=["OpenAI", "GoogleGemini"])
# #     ap.add_argument(
# #         "--model",
# #         required=True,
# #         help="LLM used for extraction (e.g. gpt-4o, gemini-1.5-pro)",
# #     )
# #     ap.add_argument(
# #         "--embedder",
# #         default="text-embedding-3-small",
# #         help="Embedding model name (OpenAI or Gemini)",
# #     )
# #     args = ap.parse_args()

# #     logger.info("Processing all model answer files from database")

# #     extractor = ModelAnswerExtractor(args.provider, args.model)

# #     # Choose embedding backend
# #     if args.provider == "OpenAI":
# #         embedder = OpenAIEmbedder(args.embedder)
# #     else:  # GoogleGemini
# #         embedder = GeminiEmbedder(model_name=args.embedder)

# #     vec_db = ModelAnswerEmbeddingDB(embedder)
# #     processed = 0

# #     print(f"[bold]⏳ Fetching model answer files from database…[/]")
    
# #     # Get files from database
# #     db_files = get_model_answer_files_from_db()
    
# #     if not db_files:
# #         print("[yellow]No model answer files found in database.[/]")
# #         vec_db.close()
# #         return

# #     # Process each file from database
# #     for db_file in db_files:
# #         try:
# #             file_url = db_file['file_url']
# #             file_id = db_file['model_answer_paper_id']
# #             assessment_id = db_file['assessment_id']
            
# #             print(f"→ Processing: {pathlib.Path(file_url).name} (ID: {file_id}, Assessment: {assessment_id})")
            
# #             # Resolve file path and load text
# #             full_path = resolve_file_path(file_url)
            
# #             # Check if file extension is supported
# #             if full_path.suffix.lower() not in ['.pdf', '.docx']:
# #                 logger.warning(f"Skipping unsupported file type: {full_path.name}")
# #                 continue
            
# #             raw_text = read_text(full_path)
            
# #             if not raw_text.strip():
# #                 logger.warning(f"No text extracted from {full_path.name}")
# #                 continue
            
# #             # Extract and embed answers
# #             answers = extractor.extract(raw_text)
# #             vec_db.save_embeddings(answers)
            
# #             processed += 1
            
# #         except Exception as e:
# #             logger.error(f"Error processing {pathlib.Path(db_file['file_url']).name}: {e}")
# #             continue

# #     vec_db.close()

# #     if processed == 0:
# #         msg = "[yellow]No model-answer files found or processed.[/]"
# #     else:
# #         msg = f"[green]✅ Done. {processed} file(s) processed and embedded.[/]"
    
# #     print(f"{msg}")


# # if __name__ == "__main__":
# #     main()

# """
# Enhanced model answer embedding script with assessment-specific filtering.
# Now supports database mapping and assessment tracking.

# Run from the project root:

#   python -m src.scripts.embed_model_answers \
#          --provider OpenAI \
#          --model gpt-4o \
#          --embedder text-embedding-3-small \
#          --assessment-id ASSESSMENT_ID
# """

# import argparse
# import logging
# import pathlib
# import sys
# import os
# from typing import Iterable, List, Dict, Any

# from docx import Document
# import pdfplumber
# from rich import print
# import psycopg2
# from psycopg2.extras import RealDictCursor
# from dotenv import load_dotenv

# # project imports
# from src.services.model_answer_extractor import ModelAnswerExtractor
# from src.services.embedding.openai_embedder import OpenAIEmbedder
# from src.services.embedding.gemini_embedder import GeminiEmbedder
# from src.services.database_services.model_answer_embedding_db import (
#     ModelAnswerEmbeddingDB,
# )

# # Load environment variables
# load_dotenv()
=======
# import argparse
# import logging
# import pathlib
# from typing import Iterable
# from docx import Document
# import pdfplumber
# from rich import print

# from src.services.model_answer_extractor import ModelAnswerExtractor
# from src.services.embedding.openai_embedder import OpenAIEmbedder
# from src.services.embedding.gemini_embedder import GeminiEmbedder
# from src.services.database_services.model_answer_embedding_db import ModelAnswerEmbeddingDB
>>>>>>> ea07d59eb89948e42998ca858c699570ea3da626

# logging.basicConfig(level=logging.INFO)
# logger = logging.getLogger(__name__)

<<<<<<< HEAD

# # --------------------------------------------------------------------------- #
# #                            Database helpers                                 #
# # --------------------------------------------------------------------------- #
# def get_database_connection():
#     """Get database connection using environment variables."""
#     try:
#         conn = psycopg2.connect(
#             host=os.getenv('POSTGRES_HOST'),
#             port=os.getenv('POSTGRES_PORT'),
#             database=os.getenv('POSTGRES_DB'),
#             user=os.getenv('POSTGRES_USER'),
#             password=os.getenv('POSTGRES_PASSWORD')
#         )
#         return conn
#     except Exception as e:
#         logger.error(f"Failed to connect to database: {e}")
#         sys.exit(1)


# def get_model_answer_files_from_db(assessment_id: str = None) -> List[Dict[str, Any]]:
#     """Retrieve model answer paper file paths from database with optional assessment filtering."""
#     conn = get_database_connection()
#     try:
#         with conn.cursor(cursor_factory=RealDictCursor) as cur:
#             if assessment_id:
#                 # Filter by specific assessment
#                 cur.execute("""
#                     SELECT map.model_answer_paper_id, map.assessment_id, map.file_url, map.created_on,
#                            a.created_on as assessment_created_on, m.module_code
#                     FROM "Model_Answer_Paper" map 
#                     JOIN "Assessment" a ON map.assessment_id = a.assessment_id
#                     JOIN "Module" m ON a.module_id = m.module_id
#                     WHERE map.assessment_id = %s
#                     ORDER BY map.created_on ASC
#                 """, (assessment_id,))
#             else:
#                 # Get all model answer files
#                 cur.execute("""
#                     SELECT map.model_answer_paper_id, map.assessment_id, map.file_url, map.created_on,
#                            a.created_on as assessment_created_on, m.module_code
#                     FROM "Model_Answer_Paper" map 
#                     JOIN "Assessment" a ON map.assessment_id = a.assessment_id
#                     JOIN "Module" m ON a.module_id = m.module_id
#                     ORDER BY map.created_on ASC
#                 """)
            
#             files = cur.fetchall()
#             logger.info(f"Found {len(files)} model answer files in database")
#             return files
            
#     except Exception as e:
#         logger.error(f"Database error: {e}")
#         return []
#     finally:
#         conn.close()


# def get_project_root() -> pathlib.Path:
#     """Get the project root directory (current working directory)."""
#     return pathlib.Path.cwd()


# def resolve_file_path(file_path: str) -> pathlib.Path:
#     """Resolve database file path to actual file location."""
#     project_root = get_project_root()
#     # Try parent directory first (where data folder should be)
#     full_path = project_root.parent / file_path
    
#     logger.info(f"Looking for file at: {full_path}")
    
#     if not full_path.exists():
#         # Try alternative path in project directory
#         alternative_path = project_root / file_path
#         logger.info(f"Alternative path: {alternative_path}")
        
#         if alternative_path.exists():
#             full_path = alternative_path
#         else:
#             # List what's actually in the expected directory
#             expected_dir = full_path.parent
#             if expected_dir.exists():
#                 files_in_dir = list(expected_dir.iterdir())
#                 logger.error(f"Expected directory exists but file not found. Contents: {[f.name for f in files_in_dir]}")
#             else:
#                 logger.error(f"Expected directory does not exist: {expected_dir}")
#             raise FileNotFoundError(f"File not found at: {full_path} or {alternative_path}")
    
#     logger.info(f"Successfully found file at: {full_path}")
#     return full_path


# # --------------------------------------------------------------------------- #
# #                              File helpers                                   #
# # --------------------------------------------------------------------------- #
# def read_text(path: pathlib.Path) -> str:
#     """Return plain text from .docx or .pdf (raise on others)."""
#     suffix = path.suffix.lower()

#     if suffix == ".docx":
#         doc = Document(path)
#         return "\n".join(p.text for p in doc.paragraphs)

#     if suffix == ".pdf":
#         with pdfplumber.open(path) as pdf:
#             return "\n".join(p.extract_text() or "" for p in pdf.pages)

#     raise ValueError(f"Unsupported file type: {path.name}")


# def iter_files(root: pathlib.Path, patterns: Iterable[str]) -> Iterable[pathlib.Path]:
#     """Yield files under *root* whose suffix matches one of *patterns*."""
=======
# def read_text(path: pathlib.Path) -> str:
#     if path.suffix.lower() == ".docx":
#         doc = Document(path)
#         return "\n".join(p.text for p in doc.paragraphs)
#     elif path.suffix.lower() == ".pdf":
#         with pdfplumber.open(path) as pdf:
#             return "\n".join(p.extract_text() or "" for p in pdf.pages)
#     raise ValueError(f"Unsupported file type: {path.name}")

# def iter_files(root: pathlib.Path, patterns: Iterable[str]) -> Iterable[pathlib.Path]:
>>>>>>> ea07d59eb89948e42998ca858c699570ea3da626
#     for file in root.rglob("*"):
#         if file.suffix.lower() in patterns and file.is_file():
#             yield file

<<<<<<< HEAD

# # --------------------------------------------------------------------------- #
# #                               Main routine                                  #
# # --------------------------------------------------------------------------- #
# def main(provider: str = None, model: str = None, embedder_name: str = None, 
#          assessment_id: str = None, **kwargs) -> None:
#     """
#     Main function that can be called from Flask API or command line.
    
#     Args:
#         provider: LLM provider ("OpenAI" or "GoogleGemini")
#         model: Model name for extraction
#         embedder_name: Embedding model name
#         assessment_id: Filter by specific assessment ID
#         **kwargs: Additional arguments from pipeline calls
#     """
    
#     # Handle command line arguments if called as script
#     if provider is None:
#         ap = argparse.ArgumentParser(description="Extract & embed model answers")
#         ap.add_argument("--provider", required=True, choices=["OpenAI", "GoogleGemini"])
#         ap.add_argument(
#             "--model",
#             required=True,
#             help="LLM used for extraction (e.g. gpt-4o, gemini-1.5-pro)",
#         )
#         ap.add_argument(
#             "--embedder",
#             default="text-embedding-3-small",
#             help="Embedding model name (OpenAI or Gemini)",
#         )
#         ap.add_argument(
#             "--assessment-id",
#             help="Filter by specific assessment ID",
#         )
#         args = ap.parse_args()
        
#         provider = args.provider
#         model = args.model
#         embedder_name = args.embedder
#         assessment_id = args.assessment_id

#     # Set defaults if not provided
#     if not embedder_name:
#         embedder_name = 'text-embedding-3-small' if provider == 'OpenAI' else 'models/embedding-001'

#     logger.info(f"Processing model answer files from database with provider: {provider}")
#     if assessment_id:
#         logger.info(f"Filtering by assessment ID: {assessment_id}")

#     extractor = ModelAnswerExtractor(provider, model)

#     # Choose embedding backend
#     if provider == "OpenAI":
#         embedder = OpenAIEmbedder(embedder_name)
#     else:  # GoogleGemini
#         embedder = GeminiEmbedder(model_name=embedder_name)

#     vec_db = ModelAnswerEmbeddingDB(embedder)
#     processed = 0

#     print(f"[bold]⏳ Fetching model answer files from database…[/]")
    
#     # Get files from database with optional assessment filtering
#     db_files = get_model_answer_files_from_db(assessment_id)
    
#     if not db_files:
#         print("[yellow]No model answer files found in database.[/]")
#         vec_db.close()
#         return

#     # Process each file from database
#     for db_file in db_files:
#         try:
#             file_url = db_file['file_url']
#             file_id = db_file['model_answer_paper_id']
#             assessment_id_from_db = db_file['assessment_id']
#             module_code = db_file['module_code']
            
#             # Get year and month from assessment creation date
#             created_date = db_file['assessment_created_on']
#             year = created_date.year
#             month = created_date.strftime('%B')
            
#             print(f"→ Processing: {pathlib.Path(file_url).name}")
#             print(f"  Assessment: {assessment_id_from_db}, Module: {module_code}")
#             print(f"  Date: {year}-{month}")
            
#             # Resolve file path and load text
#             full_path = resolve_file_path(file_url)
            
#             # Check if file extension is supported
#             if full_path.suffix.lower() not in ['.pdf', '.docx']:
#                 logger.warning(f"Skipping unsupported file type: {full_path.name}")
#                 continue
            
#             raw_text = read_text(full_path)
            
#             if not raw_text.strip():
#                 logger.warning(f"No text extracted from {full_path.name}")
#                 continue
            
#             # Extract answers
#             answers = extractor.extract(raw_text)
            
#             if not answers:
#                 logger.warning(f"No answers extracted from {full_path.name}")
#                 continue
            
#             # Override extracted metadata with database mappings
#             for answer in answers:
#                 answer.module_code = module_code  # Use database module_code
#                 answer.exam_year = year  # Use assessment creation year
#                 answer.exam_month = month  # Use assessment creation month
#                 # Add assessment_id if the model supports it
#                 if hasattr(answer, 'assessment_id'):
#                     answer.assessment_id = assessment_id_from_db
            
#             logger.info(f"Mapped metadata: Module {module_code}, {year}-{month}, Assessment {assessment_id_from_db}")
            
#             # Save embeddings with mapped metadata
#             vec_db.save_embeddings(answers, assessment_id=assessment_id_from_db)
            
#             processed += 1
            
#         except Exception as e:
#             logger.error(f"Error processing {pathlib.Path(db_file['file_url']).name}: {e}")
#             continue

#     vec_db.close()

#     if processed == 0:
#         msg = "[yellow]No model-answer files found or processed.[/]"
#     else:
#         msg = f"[green]✅ Done. {processed} file(s) processed and embedded.[/]"
    
#     print(f"{msg}")
    
#     return {
#         "status": "success",
#         "message": f"Processed {processed} model answer files",
#         "details": {
#             "processed": processed,
#             "assessment_id": assessment_id,
#             "provider": provider
#         }
#     }

=======
# def main() -> None:
#     ap = argparse.ArgumentParser(description="Extract & embed model answers")
#     ap.add_argument("--provider", required=True, choices=["OpenAI", "GoogleGemini"])
#     ap.add_argument("--model", required=True, help="LLM model name (e.g. gpt-4o, gemini-1.5-pro)")
#     ap.add_argument("--embedder", default="text-embedding-3-small", help="Embedding model name")
#     ap.add_argument("--root", default="data/Model_Answers", help="Root folder of model answers")
#     ap.add_argument("--ext", nargs="*", default=[".pdf", ".docx"], help="File extensions (default: .pdf .docx)")
#     args = ap.parse_args()

#     root = pathlib.Path(args.root)
#     if not root.exists():
#         logger.error("Root folder does not exist: %s", root)
#         return

#     print(f"[bold]⏳ Scanning [cyan]{root}[/] …[/]")

#     extractor = ModelAnswerExtractor(args.provider, args.model)

#     # Choose embedder
#     embedder = OpenAIEmbedder(args.embedder) if args.provider == "OpenAI" else GeminiEmbedder(model_name=args.embedder)

#     # Vector DB
#     vec_db = ModelAnswerEmbeddingDB(embedder)
#     file_patterns = {e.lower() if e.startswith(".") else f".{e.lower()}" for e in args.ext}

#     processed = 0
#     for file in iter_files(root, file_patterns):
#         processed += 1
#         print(f"→ {file.relative_to(root)}")
#         raw_text = read_text(file)

#         answers = extractor.extract(raw_text)
#         vec_db.save_embeddings(answers)

#     vec_db.close()
#     msg = "No model-answer files found." if processed == 0 else f"✅ Done. {processed} file(s) processed."
#     print(f"[green]{msg}[/]")
>>>>>>> ea07d59eb89948e42998ca858c699570ea3da626

# if __name__ == "__main__":
#     main()

<<<<<<< HEAD
"""
Run from the project root:

  python -m src.scripts.embed_model_answers \
         --provider OpenAI \
         --model gpt-4o \
         --embedder text-embedding-3-small \
         --assessment-id ASSESSMENT_ID
"""

=======
>>>>>>> ea07d59eb89948e42998ca858c699570ea3da626
import argparse
import logging
import pathlib
import sys
import os
from typing import Iterable, List, Dict, Any, Optional

from docx import Document
import pdfplumber
from rich import print
import psycopg2
from psycopg2.extras import RealDictCursor
from dotenv import load_dotenv

# project imports
from src.services.model_answer_extractor import ModelAnswerExtractor
from src.services.embedding.openai_embedder import OpenAIEmbedder
from src.services.embedding.gemini_embedder import GeminiEmbedder
from src.services.database_services.model_answer_embedding_db import (
    ModelAnswerEmbeddingDB,
)

# Load environment variables
load_dotenv()

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


<<<<<<< HEAD
# --------------------------------------------------------------------------- #
#                            Database helpers                                 #
# --------------------------------------------------------------------------- #
def get_database_connection():
    """Get database connection using environment variables."""
    try:
        conn = psycopg2.connect(
            host=os.getenv('POSTGRES_HOST'),
            port=os.getenv('POSTGRES_PORT'),
            database=os.getenv('POSTGRES_DB'),
            user=os.getenv('POSTGRES_USER'),
            password=os.getenv('POSTGRES_PASSWORD')
        )
        return conn
    except Exception as e:
        logger.error(f"Failed to connect to database: {e}")
        sys.exit(1)


def get_assessment_data(assessment_id: str) -> Optional[Dict[str, Any]]:
    """Get assessment data with module information."""
    conn = get_database_connection()
    try:
        with conn.cursor(cursor_factory=RealDictCursor) as cur:
            cur.execute("""
                SELECT a.assessment_id, a.created_on, a.module_id,
                       m.module_code, m.module_name
                FROM "Assessment" a
                JOIN "Module" m ON a.module_id = m.module_id
                WHERE a.assessment_id = %s
            """, (assessment_id,))
            
            result = cur.fetchone()
            if result:
                created_date = result['created_on']
                return {
                    'assessment_id': result['assessment_id'],
                    'module_code': result['module_code'],
                    'module_name': result['module_name'],
                    'exam_year': created_date.year,
                    'exam_month': created_date.strftime('%B')
                }
            return None
            
    except Exception as e:
        logger.error(f"Database error getting assessment data: {e}")
        return None
    finally:
        conn.close()


def get_model_answer_file_for_assessment(assessment_id: str) -> Optional[Dict[str, Any]]:
    """Retrieve model answer paper file for specific assessment."""
    conn = get_database_connection()
    try:
        with conn.cursor(cursor_factory=RealDictCursor) as cur:
            cur.execute("""
                SELECT model_answer_paper_id, assessment_id, file_url, created_on 
                FROM "Model_Answer_Paper" 
                WHERE assessment_id = %s
            """, (assessment_id,))
            
            result = cur.fetchone()
            if result:
                logger.info(f"Found model answer file for assessment {assessment_id}: {result['file_url']}")
                return dict(result)
            else:
                logger.warning(f"No model answer file found for assessment {assessment_id}")
                return None
            
    except Exception as e:
        logger.error(f"Database error: {e}")
        return None
    finally:
        conn.close()


def get_project_root() -> pathlib.Path:
    """Get the project root directory (current working directory)."""
    return pathlib.Path.cwd()


def resolve_file_path(file_path: str) -> pathlib.Path:
    """Resolve database file path to actual file location."""
    project_root = get_project_root()
    # Try parent directory first (where data folder should be)
    full_path = project_root.parent / file_path
    
    logger.info(f"Looking for file at: {full_path}")
    
    if not full_path.exists():
        # Try alternative path in project directory
        alternative_path = project_root / file_path
        logger.info(f"Alternative path: {alternative_path}")
        
        if alternative_path.exists():
            full_path = alternative_path
        else:
            # List what's actually in the expected directory
            expected_dir = full_path.parent
            if expected_dir.exists():
                files_in_dir = list(expected_dir.iterdir())
                logger.error(f"Expected directory exists but file not found. Contents: {[f.name for f in files_in_dir]}")
            else:
                logger.error(f"Expected directory does not exist: {expected_dir}")
            raise FileNotFoundError(f"File not found at: {full_path} or {alternative_path}")
    
    logger.info(f"Successfully found file at: {full_path}")
    return full_path


# --------------------------------------------------------------------------- #
#                              File helpers                                   #
# --------------------------------------------------------------------------- #
=======
>>>>>>> ea07d59eb89948e42998ca858c699570ea3da626
def read_text(path: pathlib.Path) -> str:
    """Return plain text from .docx or .pdf (raise on others)."""
    suffix = path.suffix.lower()

    if suffix == ".docx":
        doc = Document(path)
        return "\n".join(p.text for p in doc.paragraphs)

    if suffix == ".pdf":
        with pdfplumber.open(path) as pdf:
            return "\n".join(p.extract_text() or "" for p in pdf.pages)

    raise ValueError(f"Unsupported file type: {path.name}")

<<<<<<< HEAD

# --------------------------------------------------------------------------- #
#                               Main routine                                  #
# --------------------------------------------------------------------------- #
def main(assessment_id: Optional[str] = None, provider: Optional[str] = None, 
         model: Optional[str] = None, embedder: Optional[str] = None) -> None:
    """Main function that can be called directly or via command line."""
    
    # Handle command line arguments if called via CLI
    if assessment_id is None:
        ap = argparse.ArgumentParser(description="Extract & embed model answers for specific assessment")
        ap.add_argument("--provider", required=True, choices=["OpenAI", "GoogleGemini"])
        ap.add_argument(
            "--model",
            required=True,
            help="LLM used for extraction (e.g. gpt-4o, gemini-1.5-pro)",
        )
        ap.add_argument(
            "--embedder",
            default="text-embedding-3-small",
            help="Embedding model name (OpenAI or Gemini)",
        )
        ap.add_argument(
            "--assessment-id",
            required=True,
            help="Assessment ID to process model answer for"
        )
        args = ap.parse_args()
        
        assessment_id = args.assessment_id
        provider = args.provider
        model = args.model
        embedder = args.embedder
=======

def iter_files(root: pathlib.Path, patterns: Iterable[str]) -> Iterable[pathlib.Path]:
    for file in root.rglob("*"):
        if file.suffix.lower() in patterns and file.is_file():
            yield file


def main() -> None:
    ap = argparse.ArgumentParser(description="Extract & embed model answers")
    ap.add_argument("--provider", required=True, choices=["OpenAI", "GoogleGemini", "DeepSeek"])
    ap.add_argument("--model", required=True, help="LLM model name (e.g. gpt-4o, gemini-1.5-pro, deepseek-chat)")
    ap.add_argument("--embedder", default="text-embedding-3-small", help="Embedding model name")
    ap.add_argument("--root", default="data/Model_Answers", help="Root folder of model answers")
    ap.add_argument("--ext", nargs="*", default=[".pdf", ".docx"], help="File extensions (default: .pdf .docx)")
    args = ap.parse_args()
>>>>>>> ea07d59eb89948e42998ca858c699570ea3da626

    logger.info(f"Processing model answer for assessment {assessment_id}")

    # Get assessment data from database
    assessment_data = get_assessment_data(assessment_id)
    if not assessment_data:
        print(f"[red]Assessment {assessment_id} not found in database.[/]")
        return

    # Get model answer file for this specific assessment
    model_answer_file = get_model_answer_file_for_assessment(assessment_id)
    if not model_answer_file:
        print(f"[yellow]No model answer file found for assessment {assessment_id}.[/]")
        return

    # Create extractor - but we'll override the extracted metadata with database data
    extractor = ModelAnswerExtractor(provider, model)

<<<<<<< HEAD
    # Choose embedding backend
    if provider == "OpenAI":
        embedder_instance = OpenAIEmbedder(embedder)
    else:  # GoogleGemini
        embedder_instance = GeminiEmbedder(model_name=embedder)

    vec_db = ModelAnswerEmbeddingDB(embedder_instance)
=======
    # Choose embedder
    provider_override = None
    if args.provider == "OpenAI":
        embedder = OpenAIEmbedder(args.embedder)
    elif args.provider == "GoogleGemini":
        embedder = GeminiEmbedder(model_name=args.embedder)
    elif args.provider == "DeepSeek":
        print("⚠️ DeepSeek selected: using OpenAI embeddings but saving in DeepSeek table")
        embedder = OpenAIEmbedder("text-embedding-3-small")
        provider_override = "deepseek"
    else:
        raise ValueError(f"Unsupported provider: {args.provider}")

    # Vector DB
    vec_db = ModelAnswerEmbeddingDB(embedder, provider_override=provider_override)
    file_patterns = {e.lower() if e.startswith(".") else f".{e.lower()}" for e in args.ext}
>>>>>>> ea07d59eb89948e42998ca858c699570ea3da626

    try:
        file_url = model_answer_file['file_url']
        file_id = model_answer_file['model_answer_paper_id']
        
        print(f"→ Processing: {pathlib.Path(file_url).name} (ID: {file_id}, Assessment: {assessment_id})")
        
        # Resolve file path and load text
        full_path = resolve_file_path(file_url)
        
        # Check if file extension is supported
        if full_path.suffix.lower() not in ['.pdf', '.docx']:
            logger.warning(f"Skipping unsupported file type: {full_path.name}")
            return

        raw_text = read_text(full_path)
        
        if not raw_text.strip():
            logger.warning(f"No text extracted from {full_path.name}")
            return
        
        # Extract answers using LLM
        answers = extractor.extract(raw_text)
        
        # Override extracted metadata with database-mapped data
        for answer in answers:
            answer.module_code = assessment_data['module_code']
            answer.exam_year = assessment_data['exam_year']
            answer.exam_month = assessment_data['exam_month']
        
        # Save embeddings with assessment context
        vec_db.save_embeddings(answers, assessment_id)
        
        print(f"[green]✅ Successfully processed and embedded model answer for assessment {assessment_id}[/]")
        print(f"[green]   Module: {assessment_data['module_code']}, Year: {assessment_data['exam_year']}, Month: {assessment_data['exam_month']}[/]")

    except Exception as e:
        logger.error(f"Error processing model answer for assessment {assessment_id}: {e}")
        print(f"[red]Error processing model answer: {e}[/]")
    finally:
        vec_db.close()



if __name__ == "__main__":
<<<<<<< HEAD
    main()
=======
    main()

>>>>>>> ea07d59eb89948e42998ca858c699570ea3da626
